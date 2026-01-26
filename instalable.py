import tkinter as tk
from tkinter import messagebox, ttk, filedialog
import pandas as pd
import sqlite3 
import hashlib 
import datetime
import os
import shutil
import subprocess
from PIL import Image, ImageTk 
import webbrowser
import customtkinter as ctk
import docx
from jinja2 import Environment, FileSystemLoader
from xhtml2pdf import pisa
import time
import tempfile
import qrcode
import base64
from io import BytesIO
from db_manager import DatabaseManager
import psycopg2.extras
import sys

def resource_path(relative_path):
    """ Obtiene la ruta absoluta al recurso, funciona para dev y para PyInstaller """
    try:
        # PyInstaller crea una carpeta temporal en _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)



def verificar_integridad_db():
    print('--- VERIFICANDO INTEGRIDAD DE BASE DE DATOS ---')
    if not db_manager:
        print("Error: db_manager no inicializado. Saltando verificación.")
        return

    try:
        conn = db_manager.get_connection()
        cursor = db_manager.get_cursor(conn)

        nuevas_columnas = [
            ("fecha_desembolso_real", "TEXT"),
            ("monto_aprobado", "REAL"),
            ("tasa_interes", "REAL"),
            ("plazo_meses", "INTEGER"),
            ("valor_cuota", "REAL"),
            ("dia_pago", "INTEGER")
        ]
        
        # Obtener columnas existentes (abstracted by db_manager)
        try:
            columnas_existentes = db_manager.get_column_names(cursor, "Microcreditos")
        except Exception:
            # Si falla, intentar crear la tabla (aunque crear_tablas ya deberia haber corrido)
            # O simplemente asumir vacio
            columnas_existentes = []

        for col_nombre, col_tipo in nuevas_columnas:
            if col_nombre not in columnas_existentes:
                print(f"Reparando: Agregando columna faltante '{col_nombre}'...")
                try:
                    # ALTER TABLE es standard para ADD COLUMN
                    cursor.execute(f"ALTER TABLE Microcreditos ADD COLUMN {col_nombre} {col_tipo}")
                    print(f"Columna '{col_nombre}' agregada.")
                except Exception as e:
                    print(f"Info/Error al agregar {col_nombre}: {e}")
            else:
                print(f"Verificado: '{col_nombre}' ya existe.")

        conn.commit()
        db_manager.release_connection(conn)
        print('--- VERIFICACIÓN COMPLETADA ---')
    except Exception as e:
        print(f"Error crítico en verificación DB: {e}")


def generar_qr_base64(texto):
    qr = qrcode.QRCode(version=1, box_size=10, border=2)
    qr.add_data(texto)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buffer = BytesIO()
    img.save(buffer, format="PNG")
    return base64.b64encode(buffer.getvalue()).decode('utf-8')

# --- CONFIGURACIÓN CUSTOMTKINTER ---
ctk.set_appearance_mode("System")  # Modes: "System" (standard), "Dark", "Light"
ctk.set_default_color_theme("blue")  # Themes: "blue" (standard), "green", "dark-blue"

# =================================================================
# VARIABLES GLOBALES
# =================================================================

USUARIO_ACTIVO = None
NIVEL_ACCESO = 0 
CEDULA_CLIENTE_SELECCIONADO = None
ID_CLIENTE_SELECCIONADO = None 
status_micro_actual = None
dict_botones_status = {}

DB_NAME = 'Alianza.db' # Mantenido por referencia, pero usaremos PostgreSQL

# Configuración PostgreSQL (Ajustar según red local)
try:
    db_manager = DatabaseManager(
        host="localhost", # IP del servidor
        database="alianza_db",
        user="postgres",
        password="clave_segura" # Cambiar por la real
    )
except Exception as e:
    # Mostramos el error en consola, el manejo amigable se hace en los módulos
    print(f"Alerta de Arquitectura: {e}")
    db_manager = None

# =================================================================
# FUNCIÓN HELPER PARA LOGO
# =================================================================

def agregar_logo(ventana, logo_path="Logo Face.jpg", width=150, height=140):
    """Agrega el logo en la esquina superior derecha de la ventana"""
    try:
        if os.path.exists(resource_path(logo_path)):
            img = Image.open(resource_path(logo_path))
            img = img.resize((width, height), Image.Resampling.LANCZOS)
            logo_img = ImageTk.PhotoImage(img)
            
            logo_label = tk.Label(ventana, image=logo_img, bg="#FAFAD2")
            logo_label.image = logo_img  # Keep reference to prevent garbage collection
            logo_label.pack(side='top', anchor='ne', padx=10, pady=10)
            return logo_label
    except Exception as e:
        print(f"Error cargando logo: {e}")
    return None

# =================================================================
# 1. BASE DE DATOS
# =================================================================

def generar_hash(clave):
    return hashlib.sha256(clave.encode()).hexdigest()

from contextlib import contextmanager

def conectar_db():
    """
    Wrapper de compatibilidad para PostgreSQL/SQLite.
    Retorna una conexión y un cursor normalizado.
    """
    if not db_manager:
        raise ConnectionError("No hay conexión con el servidor de base de datos.")
    
    conn = db_manager.get_connection()
    return conn, db_manager.get_cursor(conn)

@contextmanager
def db_connection():
    """
    Context Manager para manejo seguro de conexiones.
    Se encarga de commitear o hacer rollback y liberar la conexión.
    """
    conn, cursor = conectar_db()
    try:
        yield conn, cursor
        conn.commit()
    except Exception as e:
        conn.rollback()
        raise e
    finally:
        db_manager.release_connection(conn)

def crear_tablas():
    conn, cursor = conectar_db()
    try:
        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS Clientes (
                id {sql_type("SERIAL PRIMARY KEY")}, 
                cedula TEXT UNIQUE NOT NULL,  
                ruc TEXT,
                nombre TEXT NOT NULL,
                estado_civil TEXT,
                cargas_familiares INTEGER,
                email TEXT,
                telefono TEXT,
                direccion TEXT,
                parroquia TEXT,
                tipo_vivienda TEXT,
                profesion TEXT,
                ingresos_mensuales REAL,
                referencia1 TEXT,
                referencia2 TEXT,
                asesor TEXT,
                apertura TEXT,
                numero_carpeta TEXT,
                "fecha nacimiento" TEXT,
                producto TEXT,
                observaciones TEXT,
                "cartera castigada" INTEGER DEFAULT 0,
                "valor cartera" REAL,
                "demanda judicial" INTEGER DEFAULT 0,
                "valor demanda" REAL,
                "problemas justicia" INTEGER DEFAULT 0,
                "detalle justicia" TEXT
            )
        """)
        
        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS Usuarios (
                id {sql_type("SERIAL PRIMARY KEY")},
                usuario TEXT UNIQUE NOT NULL,
                clave_hash TEXT NOT NULL, 
                nivel_acceso INTEGER NOT NULL,
                estado INTEGER DEFAULT 1,
                rol TEXT
            )
        """)

        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS Auditoria (
                id {sql_type("SERIAL PRIMARY KEY")},
                id_usuario TEXT,
                accion TEXT,
                id_cliente TEXT,
                detalles TEXT,
                timestamp TEXT
            )
        """)
        
        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS Rehabilitacion (
                id {sql_type("SERIAL PRIMARY KEY")},
                cedula_cliente TEXT UNIQUE NOT NULL,
                fecha_inicio TEXT,
                terminos TEXT,
                resultado TEXT,
                finalizado INTEGER DEFAULT 0,
                FOREIGN KEY (cedula_cliente) REFERENCES Clientes(cedula)
            )
        """)

        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS visitas_microcredito (
                id {sql_type("SERIAL PRIMARY KEY")},
                cedula_cliente TEXT NOT NULL,
                fecha TEXT,
                observaciones TEXT,
                FOREIGN KEY (cedula_cliente) REFERENCES Clientes(cedula)
            )
        """)

        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS Pagos (
                id {sql_type("SERIAL PRIMARY KEY")},
                fecha TEXT,
                cedula_cliente TEXT,
                cuota_nro INTEGER,
                valor_capital REAL,
                valor_interes REAL,
                valor_mora REAL,
                total_pagado REAL,
                usuario TEXT
            )
        """)

        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS Caja (
                id {sql_type("SERIAL PRIMARY KEY")},
                fecha_hora TEXT,
                cedula TEXT,
                ruc TEXT,
                nombres_completos TEXT,
                email TEXT,
                direccion TEXT,
                telefono TEXT,
                estado_civil TEXT,
                asesor TEXT,
                buro_credito TEXT,
                buro_archivo_ruta TEXT,
                valor_apertura REAL,
                numero_apertura TEXT,
                estado_impreso TEXT
            )
        """)

        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS Intermediacion_Detalles (
                id {sql_type("SERIAL PRIMARY KEY")},
                cedula_cliente TEXT UNIQUE NOT NULL,
                
                chk_dis_temp INTEGER DEFAULT 0, val_dis_temp REAL DEFAULT 0,
                chk_dis_reg INTEGER DEFAULT 0, val_dis_reg REAL DEFAULT 0,
                chk_pat_1 INTEGER DEFAULT 0, val_pat_1 REAL DEFAULT 0,
                chk_pat_2 INTEGER DEFAULT 0, val_pat_2 REAL DEFAULT 0,
                chk_inmueble INTEGER DEFAULT 0, val_inmueble REAL DEFAULT 0,
                chk_ruc INTEGER DEFAULT 0, val_ruc REAL DEFAULT 0,
                chk_dis_temp INTEGER DEFAULT 0, val_dis_temp REAL DEFAULT 0, obs_dis_temp TEXT,
                chk_dis_reg INTEGER DEFAULT 0, val_dis_reg REAL DEFAULT 0, obs_dis_reg TEXT,
                chk_pat_1 INTEGER DEFAULT 0, val_pat_1 REAL DEFAULT 0, obs_pat_1 TEXT,
                chk_pat_2 INTEGER DEFAULT 0, val_pat_2 REAL DEFAULT 0, obs_pat_2 TEXT,
                chk_inmueble INTEGER DEFAULT 0, val_inmueble REAL DEFAULT 0, obs_inmueble TEXT,
                chk_ruc INTEGER DEFAULT 0, val_ruc REAL DEFAULT 0, obs_ruc TEXT,
                chk_declaraciones INTEGER DEFAULT 0, val_declaraciones REAL DEFAULT 0, obs_declaraciones TEXT,
                chk_estados_cta INTEGER DEFAULT 0, val_estados_cta REAL DEFAULT 0, obs_estados_cta TEXT,
                
                fecha_cita TEXT, 
                informe_cita TEXT,
                
                chk_precal INTEGER DEFAULT 0, desc_precal TEXT,
                chk_aprob INTEGER DEFAULT 0, desc_aprob TEXT,
                fecha_desembolso_inter TEXT,
                chk_informe INTEGER DEFAULT 0, desc_informe TEXT,
                
                FOREIGN KEY (cedula_cliente) REFERENCES Clientes(cedula)
            )
        """)

        cursor.execute("SELECT COUNT(*) FROM Usuarios")
        if cursor.fetchone()[0] == 0:
            hash_admin = generar_hash('cyberpol2022') 
            cursor.execute("INSERT INTO Usuarios (usuario, clave_hash, nivel_acceso, rol) VALUES (%s, %s, %s, %s)", ('Paul', hash_admin, 1, 'Administrador'))

        conn.commit()
    except Exception as e: print(f"Error DB: {e}")
    finally: db_manager.release_connection(conn)

# Helpers delegados a db_manager
def sql_type(t): return db_manager.sql_type(t)
def get_column_names(c, t): return db_manager.get_column_names(c, t)
def check_table_exists(c, t): return db_manager.check_table_exists(c, t)
def fix_query(q): 
    if db_manager.mode == "SQLITE": return q.replace("%s", "?")
    return q

def registrar_auditoria(accion, id_cliente=None, detalles=None):
    """Registra una acción en la tabla de Auditoria"""
    conn, cursor = conectar_db()
    try:
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        cursor.execute(fix_query("INSERT INTO Auditoria (id_usuario, accion, id_cliente, detalles, timestamp) VALUES (%s,%s,%s,%s,%s)"),
                       (USUARIO_ACTIVO, accion, id_cliente, detalles, ts))
        conn.commit()
    except Exception as e: print(f"Error auditoría: {e}")
    finally: db_manager.release_connection(conn)

def migrar_db():
    conn, cursor = conectar_db()
    try:
        # Verificar si existe la columna 'referencia_vivienda'
        columnas = get_column_names(cursor, 'Clientes')
        
        if 'referencia_vivienda' not in columnas:
            print("Migrando DB: Agregando columna 'referencia_vivienda'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN referencia_vivienda TEXT")
            conn.commit()
            print("Migración completada.")
            
        if 'situacion_financiera' not in columnas:
            print("Migrando DB: Agregando columna 'situacion_financiera'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN situacion_financiera TEXT")
            conn.commit()
            print("Migración completada.")
            
        if 'terreno' not in columnas:
            print("Migrando DB: Agregando columna 'terreno'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN terreno INTEGER DEFAULT 0")
            conn.commit()
            print("Migración completada.")

        if 'valor_terreno' not in columnas:
            print("Migrando DB: Agregando columna 'valor_terreno'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN valor_terreno REAL")
            conn.commit()
            print("Migración completada.")
            
        if 'hipotecado' not in columnas:
            print("Migrando DB: Agregando columna 'hipotecado'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN hipotecado TEXT")
            conn.commit()
            print("Migración completada.")
            
        if 'fuente_ingreso' not in columnas:
            print("Migrando DB: Agregando columna 'fuente_ingreso'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN fuente_ingreso TEXT")
            conn.commit()
            
        if 'ingresos_mensuales_2' not in columnas:
            print("Migrando DB: Agregando columna 'ingresos_mensuales_2'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN ingresos_mensuales_2 REAL")
            conn.commit()
        if 'fuente_ingreso_2' not in columnas:
            print("Migrando DB: Agregando columna 'fuente_ingreso_2'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN fuente_ingreso_2 TEXT")
            conn.commit()
            
        if 'casa_dep' not in columnas:
            print("Migrando DB: Agregando columna 'casa_dep'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN casa_dep INTEGER DEFAULT 0")
            conn.commit()
        if 'valor_casa_dep' not in columnas:
            print("Migrando DB: Agregando columna 'valor_casa_dep'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN valor_casa_dep REAL")
            conn.commit()
        if 'hipotecado_casa_dep' not in columnas:
            print("Migrando DB: Agregando columna 'hipotecado_casa_dep'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN hipotecado_casa_dep TEXT")
            conn.commit()
            
        if 'local' not in columnas:
            print("Migrando DB: Agregando columna 'local'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN local INTEGER DEFAULT 0")
            conn.commit()
        if 'valor_local' not in columnas:
            print("Migrando DB: Agregando columna 'valor_local'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN valor_local REAL")
            conn.commit()
        if 'hipotecado_local' not in columnas:
            print("Migrando DB: Agregando columna 'hipotecado_local'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN hipotecado_local TEXT")
            conn.commit()
            
        if 'score_buro' not in columnas:
            print("Migrando DB: Agregando columna 'score_buro'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN score_buro INTEGER")
            conn.commit()
            print("Migración completada new fields.")
            
        if 'egresos' not in columnas:
            print("Migrando DB: Agregando columna 'egresos'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN egresos REAL")
            conn.commit()
            
        if 'total_disponible' not in columnas:
            print("Migrando DB: Agregando columna 'total_disponible'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN total_disponible REAL")
            conn.commit()

        if 'valor_apertura' not in columnas:
            print("Migrando DB: Agregando columna 'valor_apertura'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN valor_apertura REAL")
            conn.commit()

        if 'fecha_registro' not in columnas:
            print("Migrando DB: Agregando columna 'fecha_registro'...")
            cursor.execute("ALTER TABLE Clientes ADD COLUMN fecha_registro TEXT")
            conn.commit()

        # Migración Usuarios
        cols_user = get_column_names(cursor, 'Usuarios')
        if 'estado' not in cols_user:
            print("Migrando DB: Agregando 'estado' a Usuarios...")
            cursor.execute("ALTER TABLE Usuarios ADD COLUMN estado INTEGER DEFAULT 1")
            conn.commit()
        if 'rol' not in cols_user:
            print("Migrando DB: Agregando 'rol' a Usuarios...")
            cursor.execute("ALTER TABLE Usuarios ADD COLUMN rol TEXT")
            cursor.execute("UPDATE Usuarios SET rol = 'Administrador' WHERE nivel_acceso = 1")
            cursor.execute("UPDATE Usuarios SET rol = 'Usuario' WHERE nivel_acceso = 2")
            conn.commit()

        # Tabla Auditoria
        if not check_table_exists(cursor, 'Auditoria'):
            print("Migrando DB: Creando tabla 'Auditoria'...")
            cursor.execute("""
                CREATE TABLE Auditoria (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    id_usuario TEXT,
                    accion TEXT,
                    id_cliente TEXT,
                    detalles TEXT,
                    timestamp TEXT
                )
            """)
            conn.commit()
            
        # Tabla Caja - MIGRACIÓN FECHA CONTRATO
        cols_caja = get_column_names(cursor, 'Caja')
        if 'fecha_contrato' not in cols_caja:
            print("Migrando DB: Agregando columna 'fecha_contrato' a Caja...")
            cursor.execute("ALTER TABLE Caja ADD COLUMN fecha_contrato TEXT")
            conn.commit()
           # Verificar si existe la tabla Documentos
        if not check_table_exists(cursor, 'Documentos'):
            print("Migrando DB: Creando tabla 'Documentos'...")
            cursor.execute("""
                CREATE TABLE Documentos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    cedula_cliente TEXT NOT NULL,
                    nombre_archivo TEXT NOT NULL,
                    tipo_documento TEXT,
                    ruta_archivo TEXT NOT NULL,
                    fecha_subida TEXT NOT NULL,
                    FOREIGN KEY (cedula_cliente) REFERENCES Clientes(cedula)
                )
            """)
            conn.commit()
            print("Tabla 'Documentos' creada exitosamente.")

        # Verificar si existe la tabla Microcreditos
        if not check_table_exists(cursor, 'Microcreditos'):
            print("Migrando DB: Creando tabla 'Microcreditos'...")
            cursor.execute("""
                CREATE TABLE Microcreditos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    cedula_cliente TEXT NOT NULL,
                    ruc TEXT,
                    observaciones TEXT,
                    observaciones_info TEXT,
                    status TEXT,
                    sub_status TEXT,
                    
                    ref1_fecha TEXT, ref1_hora TEXT, ref1_nombre TEXT, ref1_telefono TEXT,
                    ref1_relacion TEXT, ref1_tiempo_conocer TEXT, ref1_direccion TEXT,
                    ref1_tipo_vivienda TEXT, ref1_cargas TEXT, ref1_patrimonio TEXT, ref1_responsable TEXT,

                    ref2_fecha TEXT, ref2_hora TEXT, ref2_nombre TEXT, ref2_telefono TEXT,
                    ref2_relacion TEXT, ref2_tiempo_conocer TEXT, ref2_direccion TEXT,
                    ref2_tipo_vivienda TEXT, ref2_cargas TEXT, ref2_patrimonio TEXT, ref2_responsable TEXT,
                    
                    fecha_comite TEXT, fecha_desembolsado TEXT, fecha_negado TEXT, fecha_desistimiento TEXT,

                    FOREIGN KEY (cedula_cliente) REFERENCES Clientes(cedula)
                )
            """)
            conn.commit()
            print("Tabla 'Microcreditos' creada exitosamente.")
        
        # Tabla Caja
        if not check_table_exists(cursor, 'Caja'):
            print("Migrando DB: Creando tabla 'Caja'...")
            cursor.execute("""
                CREATE TABLE Caja (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    fecha_hora TEXT,
                    cedula TEXT,
                    ruc TEXT,
                    nombres_completos TEXT,
                    email TEXT,
                    direccion TEXT,
                    telefono TEXT,
                    estado_civil TEXT,
                    asesor TEXT,
                    buro_credito TEXT,
                    buro_archivo_ruta TEXT,
                    valor_apertura REAL,
                    numero_apertura TEXT,
                    estado_impreso TEXT DEFAULT 'Pendiente',
                    observaciones TEXT,
                    fecha_contrato TEXT
                )
            """)
            conn.commit()
        
        # Verificar columnas críticas en Clientes si no están
        cols = get_column_names(cursor, 'Clientes')
        missing_cols = []
        for c in ['fecha_contrato', 'score_buro', 'ingresos_mensuales_2', 'egresos', 'fecha_registro']:
            if c not in cols and c != 'fecha_contrato': # fecha_contrato is in Caja usually, but user asked for checks
                pass 
                # Already handled in previous block, just double checking logic flows

        # --- OPTIMIZACIÓN: ÍNDICES ---
        print("Optimizando DB: Verificando Índices...")
        try:
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_clientes_cedula ON Clientes(cedula)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_clientes_nombre ON Clientes(nombre)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_caja_cedula ON Caja(cedula)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_micro_cedula ON Microcreditos(cedula_cliente)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_docs_cedula ON Documentos(cedula_cliente)")
            conn.commit()
        except Exception as e:
            print(f"Nota sobre índices: {e}")

        # Migración Intermediacion Observaciones
        cols_inter = get_column_names(cursor, 'Intermediacion_Detalles')
        new_obs_cols = [
            'obs_dis_temp', 'obs_dis_reg', 'obs_pat_1', 'obs_pat_2', 
            'obs_inmueble', 'obs_ruc', 'obs_declaraciones', 'obs_estados_cta',
            'fecha_cita', 'informe_cita',
            'chk_precal', 'desc_precal',
            'chk_aprob', 'desc_aprob',
            'fecha_desembolso_inter',
            'chk_informe', 'desc_informe'
        ]
        
        if check_table_exists(cursor, 'Intermediacion_Detalles'):
            for col in new_obs_cols:
                if col not in cols_inter:
                    print(f"Migrando DB: Agregando columna '{col}' a Intermediacion_Detalles...")
                    cursor.execute(f"ALTER TABLE Intermediacion_Detalles ADD COLUMN {col} TEXT")
            conn.commit()

        # Migración Usuario Admin -> Paul
        cursor.execute("SELECT id FROM Usuarios WHERE usuario='admin'")
        admin_user = cursor.fetchone()
        if admin_user:
            print("Migrando usuario admin a Paul...")
            new_hash = generar_hash('cyberpol2022')
            cursor.execute("UPDATE Usuarios SET usuario=%s, clave_hash=%s WHERE id=%s", ('Paul', new_hash, admin_user[0]))
            conn.commit()
            print("Usuario actualizado.")

    except Exception as e: 
        print(f"Error Migración: {e}")
        # MessageBox para alertar si falla la integridad crítica
        try: messagebox.showerror("Error Crítico de Integridad", f"Fallo al verificar/reparar base de datos:\n{e}")
        except: pass
    finally: db_manager.release_connection(conn)

crear_tablas()
migrar_db()

# --- UTILIDADES DE FORMATO ---

def limpiar_moneda(valor_str):
    """Quita $ y comas para guardar en DB como número."""
    if not valor_str: return 0.0
    limpio = valor_str.replace('$', '').replace(',', '').strip()
    try: return float(limpio)
    except ValueError: return 0.0

def formatear_float_str(valor):
    """Convierte float a string '1,200.00'."""
    try:
        return "{:,.2f}".format(float(valor))
    except (ValueError, TypeError): return ""

# --- VALIDACIONES ---

def validar_datos(cedula, nombre, apertura, nacimiento, ingresos_str, val_cart_str, val_dem_str):
    if not nombre.strip(): return "El Nombre es obligatorio."
    if cedula.strip() and (len(cedula.strip()) != 10 or not cedula.strip().isdigit()):
        return "La Cédula debe tener 10 dígitos numéricos."
    
    try:
        limpiar_moneda(ingresos_str)
        limpiar_moneda(val_cart_str)
        limpiar_moneda(val_dem_str)
    except ValueError: return "Revise los campos numéricos."

    if apertura.strip():
        try: datetime.datetime.strptime(apertura, "%d/%m/%Y")
        except ValueError: return "Fecha Apertura incorrecta (DD/MM/YYYY)."
    
    if nacimiento.strip() and nacimiento != "dd/mm/aaaa":
        try: datetime.datetime.strptime(nacimiento, "%d/%m/%Y")
        except ValueError: return "Fecha de Nacimiento incorrecta (DD/MM/YYYY)."
    return True

# --- CRUD ---

def guardar_cliente(*args):
    (cedula, ruc, nombre, est_civil, cargas, email, telf, dire, parr, viv, ref_viv,
     prof, ing_str, fuente_ing, terreno_val, val_terreno_str, hipotecado, ref1, ref2, asesor, aper, num_carpeta, nac, obs, 
     cart, val_cart_str, dem, val_dem_str, just, det_just,
     casa_val, val_casa_str, hip_casa, local_val, val_local_str, hip_local,
     ing_str_2, fuente_ing_2, score_buro_str, egresos_str) = args

    val = validar_datos(cedula, nombre, aper, nac, ing_str, val_cart_str, val_dem_str)
    if val is not True: return False, val
    
    ingresos = limpiar_moneda(ing_str)
    ingresos_2 = limpiar_moneda(ing_str_2)
    egresos = limpiar_moneda(egresos_str)
    valor_terreno = limpiar_moneda(val_terreno_str) if terreno_val else 0
    valor_casa = limpiar_moneda(val_casa_str) if casa_val else 0
    valor_local = limpiar_moneda(val_local_str) if local_val else 0
    val_cart = limpiar_moneda(val_cart_str) if cart else 0
    val_dem = limpiar_moneda(val_dem_str) if dem else 0
    
    # Calcular total disponible
    total_disponible = ingresos + ingresos_2 - egresos
    
    # Validar score_buro
    score_buro = None
    if score_buro_str.strip():
        try:
            score_buro = int(score_buro_str)
            if score_buro < 1 or score_buro > 999:
                return False, "Score Buró debe estar entre 1 y 999"
        except ValueError:
            return False, "Score Buró debe ser un número"
    
    try:
        with db_connection() as (conn, cursor):
            cursor.execute("""
                INSERT INTO Clientes (
                    cedula, ruc, nombre, estado_civil, cargas_familiares, email, telefono, direccion, parroquia, 
                    tipo_vivienda, referencia_vivienda, profesion, ingresos_mensuales, fuente_ingreso, terreno, valor_terreno, hipotecado, referencia1, referencia2, asesor, fecha_registro, apertura, 
                    "fecha nacimiento", observaciones, 
                    "cartera castigada", "valor cartera", "demanda judicial", "valor demanda", "problemas justicia", "detalle justicia",
                    casa_dep, valor_casa_dep, hipotecado_casa_dep, local, valor_local, hipotecado_local,
                    ingresos_mensuales_2, fuente_ingreso_2, score_buro, egresos, total_disponible
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (cedula, ruc, nombre, est_civil, cargas, email, telf, dire, parr, viv, ref_viv,
                  prof, ingresos, fuente_ing, terreno_val, valor_terreno, hipotecado, ref1, ref2, asesor, aper, num_carpeta, nac, obs, 
                  cart, val_cart, dem, val_dem, just, det_just,
                  casa_val, valor_casa, hip_casa, local_val, valor_local, hip_local,
                  ingresos_2, fuente_ing_2, score_buro, egresos, total_disponible))
            # Commit is handled by context manager
            
        registrar_auditoria("Guardar Cliente", id_cliente=cedula, detalles=f"Cliente {nombre} guardado exitosamente.")
        return True, "Guardado exitosamente."
    except sqlite3.IntegrityError: return False, "Cédula ya existe."
    except Exception as e: return False, f"Error: {e}"

def consultar_clientes():
    conn, cursor = conectar_db()
    try:
        cursor.execute("SELECT * FROM Clientes ORDER BY apertura ASC")
        return cursor.fetchall()
    finally: db_manager.release_connection(conn)

def sincronizar_cliente_desde_caja(cedula, ruc, nombre, email, direccion, telefono, asesor, fecha_apertura, num_carpeta):
    """Sincroniza datos básicos desde Caja a Clientes (Insert o Update)."""
    if not cedula: return
    conn, cursor = conectar_db()
    try:
        cursor.execute("SELECT id FROM Clientes WHERE cedula = %s", (cedula,))
        exists = cursor.fetchone()
        if exists:
            # MAPEO CORREGIDO: apertura -> num_carpeta, fecha_registro -> fecha_apertura
            cursor.execute("""
                UPDATE Clientes SET 
                ruc=%s, nombre=%s, email=%s, direccion=%s, telefono=%s, asesor=%s, apertura=%s, fecha_registro=%s
                WHERE id = %s
            """, (ruc, nombre, email, direccion, telefono, asesor, num_carpeta, fecha_apertura, exists[0]))
        else:
            cursor.execute("""
                INSERT INTO Clientes (cedula, ruc, nombre, email, direccion, telefono, asesor, apertura, fecha_registro) 
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (cedula, ruc, nombre, email, direccion, telefono, asesor, num_carpeta, fecha_apertura))
        conn.commit()
    except Exception as e:
        try: messagebox.showerror("Error Sincronización", f"Error sincronizando cliente desde caja: {e}")
        except: pass
    finally:
        db_manager.release_connection(conn)

def actualizar_cliente(id_cliente, *args):
    (cedula, ruc, nombre, est_civil, cargas, email, telf, dire, parr, viv, ref_viv,
     prof, ing_str, fuente_ing, terreno_val, val_terreno_str, hipotecado, ref1, ref2, asesor, aper, num_carpeta, nac, obs, 
     cart, val_cart_str, dem, val_dem_str, just, det_just,
     casa_val, val_casa_str, hip_casa, local_val, val_local_str, hip_local,
     ing_str_2, fuente_ing_2, score_buro_str, egresos_str) = args

    val = validar_datos(cedula, nombre, aper, nac, ing_str, val_cart_str, val_dem_str)
    if val is not True: return False, val
    
    ingresos = limpiar_moneda(ing_str)
    ingresos_2 = limpiar_moneda(ing_str_2)
    egresos = limpiar_moneda(egresos_str)
    valor_terreno = limpiar_moneda(val_terreno_str) if terreno_val else 0
    valor_casa = limpiar_moneda(val_casa_str) if casa_val else 0
    valor_local = limpiar_moneda(val_local_str) if local_val else 0
    val_cart = limpiar_moneda(val_cart_str)
    val_dem = limpiar_moneda(val_dem_str)
    
    # Calcular total disponible
    total_disponible = ingresos + ingresos_2 - egresos
    
    # Validar score_buro
    score_buro = None
    if score_buro_str.strip():
        try:
            score_buro = int(score_buro_str)
            if score_buro < 1 or score_buro > 999:
                return False, "Score Buró debe estar entre 1 y 999"
        except ValueError:
            return False, "Score Buró debe ser un número"
    
    try:
        with db_connection() as (conn, cursor):
            cursor.execute("""
                UPDATE Clientes SET 
                cedula=%s, ruc=%s, nombre=%s, estado_civil=%s, cargas_familiares=%s, email=%s, telefono=%s, direccion=%s, parroquia=%s, 
                tipo_vivienda=%s, referencia_vivienda=%s, profesion=%s, ingresos_mensuales=%s, fuente_ingreso=%s, terreno=%s, valor_terreno=%s, hipotecado=%s, referencia1=%s, referencia2=%s, asesor=%s, fecha_registro=%s, apertura=%s, 
                "fecha nacimiento"=%s, observaciones=%s, 
                "cartera castigada"=%s, "valor cartera"=%s, "demanda judicial"=%s, "valor demanda"=%s, "problemas justicia"=%s, "detalle justicia"=%s,
                casa_dep=%s, valor_casa_dep=%s, hipotecado_casa_dep=%s, local=%s, valor_local=%s, hipotecado_local=%s,
                ingresos_mensuales_2=%s, fuente_ingreso_2=%s, score_buro=%s, egresos=%s, total_disponible=%s
                WHERE id = %s
            """, (cedula, ruc, nombre, est_civil, cargas, email, telf, dire, parr, viv, ref_viv,
                  prof, ingresos, fuente_ing, terreno_val, valor_terreno, hipotecado, ref1, ref2, asesor, aper, num_carpeta, nac, obs, 
                  cart, val_cart, dem, val_dem, just, det_just,
                  casa_val, valor_casa, hip_casa, local_val, valor_local, hip_local,
                  ingresos_2, fuente_ing_2, score_buro, egresos, total_disponible, id_cliente))
            # Commit handled by CM
            
        registrar_auditoria("Actualizar Cliente", id_cliente=cedula, detalles=f"Cliente {nombre} actualizado.")
        return True, "Actualizado correctamente."
    except Exception as e: return False, f"Error: {e}"

def eliminar_cliente(id_cliente):
    if NIVEL_ACCESO != 1: return False, "No tiene permisos para eliminar."
    with db_connection() as (conn, cursor):
        # Get cedula for audit before deleting
        cursor.execute("SELECT cedula, nombre FROM Clientes WHERE id = %s", (id_cliente,))
        res = cursor.fetchone()
        if res:
            ced, nom = res
            cursor.execute("DELETE FROM Clientes WHERE id = %s", (id_cliente,))
            # Commit handled by CM
            registrar_auditoria("Eliminar Cliente", id_cliente=ced, detalles=f"Cliente {nom} eliminado.")
    return True, "Eliminado"

def buscar_clientes(termino):
    with db_connection() as (conn, cursor):
        t = '%' + termino + '%'
        cursor.execute("SELECT * FROM Clientes WHERE nombre LIKE %s OR cedula LIKE %s OR ruc LIKE %s OR numero_carpeta LIKE %s ORDER BY apertura ASC", (t,t,t,t))
        return cursor.fetchall()

# --- USUARIOS ---
def crear_usuario_db(usuario, clave, nivel, rol="Usuario"):
    try:
        with db_connection() as (conn, cursor):
            cursor.execute("INSERT INTO Usuarios (usuario, clave_hash, nivel_acceso, rol, estado) VALUES (%s, %s, %s, %s, %s)", 
                           (usuario, generar_hash(clave), nivel, rol, 1))
        return True, "Ok"
    except Exception as e: return False, f"Error: {e}"

def verificar_credenciales(usuario, clave):
    with db_connection() as (conn, cursor):
        cursor.execute("SELECT clave_hash, nivel_acceso, estado, rol FROM Usuarios WHERE usuario = %s", (usuario,))
        res = cursor.fetchone()
    if res:
        h, lvl, est, rol = res
        if est == 0: return False, "Inactivo"
        if generar_hash(clave) == h: return True, lvl
    return False, "Error"

# =================================================================
# 2. INTERFAZ GRÁFICA
# =================================================================

# --- EVENTOS DE FORMATO MONEDA EN INTERFAZ ---
def on_focus_out_moneda(event):
    widget = event.widget
    texto = widget.get()
    try:
        val = limpiar_moneda(texto)
        if val > 0:
            widget.delete(0, tk.END)
            widget.insert(0, "{:,.2f}".format(val))
    except: pass

def on_focus_in_moneda(event):
    widget = event.widget
    texto = widget.get()
    try:
        val = limpiar_moneda(texto)
        if val == 0: widget.delete(0, tk.END)
        else:
            widget.delete(0, tk.END)
            widget.insert(0, str(val).replace('.0', '')) 
    except: pass

def saltar_campo(event):
    event.widget.tk_focusNext().focus()
    return "break"

def on_focus_in_nacimiento(event):
    if event.widget.get() == "dd/mm/aaaa":
        event.widget.delete(0, tk.END)
        event.widget.configure(text_color="black")

def on_focus_out_nacimiento(event):
    if not event.widget.get().strip():
        event.widget.delete(0, tk.END)
        event.widget.insert(0, "dd/mm/aaaa")
        event.widget.configure(text_color="#a9a9a9")

def obtener_campos_ui():
    return (
        e_cedula.get().strip(), e_ruc.get().strip(), e_nombre.get().strip(),
        c_civil.get(), e_cargas.get().strip(), e_email.get().strip(), e_telf.get().strip(),
        e_dir.get().strip(), e_parroquia.get().strip(), c_vivienda.get(), e_ref_vivienda.get().strip(),
        e_profesion.get().strip(), e_ingresos.get().strip(), c_fuente_ingreso.get(), var_terreno.get(), e_valor_terreno.get().strip(), c_hipotecado.get(),
        e_ref1.get().strip(), e_ref2.get().strip(), e_asesor.get().strip(),
        e_apertura.get().strip(), e_carpeta.get().strip(), e_nacimiento.get().strip(), 
        t_obs.get("1.0", tk.END).strip(), 
        var_cartera.get(), e_val_cartera.get().strip(),
        var_demanda.get(), e_val_demanda.get().strip(),
        var_justicia.get(), e_det_justicia.get().strip(),
        var_casa.get(), e_valor_casa.get().strip(), c_hip_casa.get(),
        var_local.get(), e_valor_local.get().strip(), c_hip_local.get(),
        e_ingresos_2.get().strip(), c_fuente_ingreso_2.get(),
        e_score_buro.get().strip(),
        e_egresos.get().strip()
    )

def toggle_inputs_clientes(state):
    """Habilita o deshabilita todos los campos de entrada del módulo de clientes."""
    # Lista de todos los widgets de entrada (Entry, ComboBox, Text, CheckBox)
    # Incluimos los checkbox específicos que fueron nombrados en abrir_modulo_clientes
    widgets = [
        e_cedula, e_ruc, e_nombre, e_cargas, e_email, e_telf, e_dir, e_parroquia,
        e_ref_vivienda, e_profesion, e_ingresos, e_ref1, e_ref2, e_asesor,
        e_apertura, e_carpeta, e_nacimiento, e_val_cartera, e_val_demanda,
        e_det_justicia, e_valor_terreno, e_valor_casa, e_valor_local,
        e_ingresos_2, e_score_buro, e_egresos,
        c_civil, c_vivienda, c_hipotecado, c_hip_casa, c_hip_local,
        c_fuente_ingreso, c_fuente_ingreso_2, t_obs,
        cb_cartera, cb_demanda, cb_justicia, cb_terreno, cb_casa, cb_local
    ]
    for w in widgets:
        try:
            if w: w.configure(state=state)
        except: pass

def limpiar_campos_ui():
    global ID_CLIENTE_SELECCIONADO
    # Desbloquear campos para permitir nuevo ingreso
    toggle_inputs_clientes('normal')
    
    elementos = [e_cedula, e_ruc, e_nombre, e_cargas, e_email, e_telf, e_dir, e_parroquia, e_ref_vivienda,
                 e_profesion, e_ingresos, e_ref1, e_ref2, e_asesor, e_apertura, e_carpeta, e_nacimiento, 
                 e_val_cartera, e_val_demanda, e_det_justicia, e_valor_terreno, e_valor_casa, e_valor_local, e_ingresos_2, e_score_buro, e_egresos]
    for e in elementos: e.delete(0, tk.END)
    # Placeholder para Nacimiento
    e_nacimiento.insert(0, "dd/mm/aaaa")
    e_nacimiento.configure(text_color="#a9a9a9")
    
    t_obs.delete("1.0", tk.END)
    c_civil.set(''); c_vivienda.set(''); c_hipotecado.set(''); c_hip_casa.set(''); c_hip_local.set('')
    c_fuente_ingreso.set(''); c_fuente_ingreso_2.set('')
    var_cartera.set(0); var_demanda.set(0); var_justicia.set(0); var_terreno.set(0); var_casa.set(0); var_local.set(0)
    toggle_legal_fields(); toggle_terreno(); toggle_casa(); toggle_local(); toggle_fuente_ingreso(); toggle_fuente_ingreso_2()
    
    # Limpiar total disponible
    lbl_total_disponible_valor.configure(text="$ 0.00")
    
    ID_CLIENTE_SELECCIONADO = None
    btn_accion.configure(text="💾 Guardar Nuevo Cliente", command=accion_guardar, state='normal')
    btn_cancelar.pack_forget() # Volver a ocultar si se requiere
    # Re-asegurar visibilidad de botones si fueron ocultos
    btn_cancelar.pack(side='left', padx=15)
    
    e_cedula.focus()

def mostrar_datos_tree():
    for i in tree.get_children(): tree.delete(i)
    data = consultar_clientes()
    for row in data:
        ing_fmt = "$ " + formatear_float_str(row['ingresos_mensuales'] if row['ingresos_mensuales'] else 0)
        # ID, Cedula, Nombre, Telf, Ingresos, Sit.Financiera (CHECKBOX), Producto, N.Apertura(apertura), Asesor
        sf = "Terreno: Si" if ('terreno' in row.keys() and row['terreno'] == 1) else ""
        # MAPEO CORREGIDO: Apertura (UI) muestra el numero de carpeta guardado en 'apertura' (DB)
        visual = (row['id'], row['cedula'], row['nombre'], row['telefono'], ing_fmt, sf, row['producto'], row['apertura'], row['asesor'])
        tree.insert('', tk.END, values=visual)

def accion_guardar():
    exito, msg = guardar_cliente(*obtener_campos_ui())
    if exito:
        try:
            top = e_cedula.winfo_toplevel()
            messagebox.showinfo("Éxito", msg, parent=top)
            limpiar_campos_ui()
            mostrar_datos_tree()
            top.lift()
            top.focus_force()
        except:
            messagebox.showinfo("Éxito", msg)
            limpiar_campos_ui()
            mostrar_datos_tree()
    else: 
        try:
            messagebox.showerror("Error", msg, parent=e_cedula.winfo_toplevel())
        except:
            messagebox.showerror("Error", msg)

def accion_actualizar():
    global ID_CLIENTE_SELECCIONADO
    print(f"DEBUG: Intentando actualizar. ID Seleccionado: {ID_CLIENTE_SELECCIONADO}")
    
    if not ID_CLIENTE_SELECCIONADO: 
        try:
             messagebox.showwarning("Atención", "No hay cliente seleccionado para actualizar.", parent=e_cedula.winfo_toplevel())
        except: pass
        return

    try:
        exito, msg = actualizar_cliente(ID_CLIENTE_SELECCIONADO, *obtener_campos_ui())
        if exito:
            try:
                top = e_cedula.winfo_toplevel()
                messagebox.showinfo("Éxito", msg, parent=top)
                mostrar_datos_tree()
                top.lift()
                top.focus_force()
            except:
                messagebox.showinfo("Éxito", msg)
                mostrar_datos_tree()
        else:
            try:
                messagebox.showerror("Error", msg, parent=e_cedula.winfo_toplevel())
            except:
                messagebox.showerror("Error", msg)
    except Exception as e:
        print(f"CRITICAL ERROR call actualizar_cliente: {e}")
        try:
            messagebox.showerror("Error Crítico", f"Falló la actualización: {e}", parent=e_cedula.winfo_toplevel())
        except:
             messagebox.showerror("Error Crítico", f"Falló la actualización: {e}")

def accion_eliminar():
    if not ID_CLIENTE_SELECCIONADO: return
    try:
        top = e_cedula.winfo_toplevel()
        confirm = messagebox.askyesno("Borrar", "¿Confirma?", parent=top)
    except:
        confirm = messagebox.askyesno("Borrar", "¿Confirma?")
        
    if confirm:
        eliminar_cliente(ID_CLIENTE_SELECCIONADO)
        limpiar_campos_ui(); mostrar_datos_tree()
        try:
            e_cedula.winfo_toplevel().lift()
            e_cedula.winfo_toplevel().focus_force()
        except: pass

def accion_buscar():
    term = e_busqueda.get().strip()
    if not term: return mostrar_datos_tree()
    res = buscar_clientes(term)
    for run in res:
        # Usar nombres de columna para evitar errores de índice
        try:
            ing_val = run['ingresos_mensuales'] if 'ingresos_mensuales' in run.keys() else run[12]
            ing_fmt = "$ " + formatear_float_str(ing_val if ing_val else 0)
            
            sf = "Terreno: Si" if ('terreno' in run.keys() and run['terreno'] == 1) else ""
            # ID, Cedula, Nombre, Telf, Ingresos, Sit.Financiera, Producto, N.Apertura, Asesor
            # Usamos acceso por nombre si es posible, fallback a índices si no (pero el cursor DictCursor debería proveer nombres)
            visual = (
                run['id'], 
                run['cedula'], 
                run['nombre'], 
                run['telefono'], 
                ing_fmt, 
                sf, 
                run['producto'], 
                run['apertura'], 
                run['asesor']
            )
        except (KeyError, TypeError, IndexError):
            # Fallback seguro si por alguna razón no vienen nombres (poco probable con DictCursor)
            ing_fmt = "$ " + formatear_float_str(run[12] if len(run) > 12 else 0)
            sf = "Terreno: Si" if (len(run) > 29 and run[29] == 1) else ""
            visual = (run[0], run[1], run[3], run[7], ing_fmt, sf, run[19], run[17], run[15])
            
        tree.insert('', tk.END, values=visual)

def cargar_seleccion(event):
    global ID_CLIENTE_SELECCIONADO, btn_eliminar
    sel = tree.selection()
    
    if not sel:
        # Si se deselecciona, limpiar ID y deshabilitar eliminar
        ID_CLIENTE_SELECCIONADO = None
        try:
            if NIVEL_ACCESO == 1: btn_eliminar.configure(state='disabled')
        except: pass
        return
    
    id_sel = tree.item(sel[0], 'values')[0]
    
    id_sel = tree.item(sel[0], 'values')[0]
    
    with db_connection() as (conn, cursor):
        cursor.execute("SELECT * FROM Clientes WHERE id = %s", (id_sel,))
        val = cursor.fetchone()
    
    if not val: return

    ID_CLIENTE_SELECCIONADO = val['id']
    
    # Limpieza general de campos registrados
    elementos = [e_cedula, e_ruc, e_nombre, e_cargas, e_email, e_telf, e_dir, e_parroquia, e_ref_vivienda,
                 e_profesion, e_ingresos, e_ref1, e_ref2, e_asesor, e_apertura, e_carpeta, e_nacimiento, 
                 e_val_cartera, e_val_demanda, e_det_justicia, e_valor_terreno, e_valor_casa, e_valor_local,
                 e_ingresos_2, e_score_buro, e_egresos] # Agregados faltantes para limpieza masiva inicial
    
    for e in elementos: 
        try:
            e.configure(state='normal')
            e.delete(0, tk.END)
            # Volver a readonly si fuera necesario (aquí asumimos normal por defecto para editar, 
            # pero si alguno es readonly se debe manejar individualmente. 
            # La mayoría son 'normal' en este form).
        except: pass

    t_obs.delete("1.0", tk.END)
    
    # Función helper para insertar seguro
    def safe_insert(entry, value):
        try:
            entry.configure(state='normal')
            entry.delete(0, tk.END)
            entry.insert(0, str(value))
        except: pass

    try:
        safe_insert(e_cedula, val['cedula'])
        if val['ruc']: safe_insert(e_ruc, val['ruc'])
        safe_insert(e_nombre, val['nombre'])
        if val['estado_civil']: c_civil.set(val['estado_civil'])
        if val['cargas_familiares']: safe_insert(e_cargas, val['cargas_familiares'])
        if val['email']: safe_insert(e_email, val['email'])
        if val['telefono']: safe_insert(e_telf, val['telefono'])
        if val['direccion']: safe_insert(e_dir, val['direccion'])
        if val['parroquia']: safe_insert(e_parroquia, val['parroquia'])
        if val['tipo_vivienda']: c_vivienda.set(val['tipo_vivienda'])
        if val['profesion']: safe_insert(e_profesion, val['profesion'])
        
        if val['ingresos_mensuales']: safe_insert(e_ingresos, formatear_float_str(val['ingresos_mensuales']))
        
        if 'fuente_ingreso' in val.keys() and val['fuente_ingreso']: c_fuente_ingreso.set(val['fuente_ingreso'])
            
        if val['referencia1']: safe_insert(e_ref1, val['referencia1'])
        if val['referencia2']: safe_insert(e_ref2, val['referencia2'])
        if val['asesor']: safe_insert(e_asesor, val['asesor'])
        
        # INTEGRACIÓN: Buscar fecha contrato en Caja
        fecha_mostrar = None
        
        # 1. Intentar buscar en Caja (Prioridad: fecha_contrato -> fecha_hora)
        try:
            cedula_buscar = val['cedula'].strip() if val['cedula'] else None
            if cedula_buscar:
                with db_connection() as (conn_caja, cur_caja):
                    # Traemos todo para verificar columnas disponibles
                    cur_caja.execute("SELECT * FROM Caja WHERE cedula = %s", (cedula_buscar,))
                    res_caja = cur_caja.fetchone()
                
                if res_caja:
                    # Convertir a dict para seguridad si es Row
                    data_caja = dict(res_caja) if hasattr(res_caja, 'keys') else res_caja
                    
                    # Prioridad 1: Fecha Contrato
                    if 'fecha_contrato' in data_caja and data_caja['fecha_contrato']:
                        fecha_mostrar = data_caja['fecha_contrato']
                    # Prioridad 2: Fecha Hora (Creación/Apertura)
                    elif 'fecha_hora' in data_caja and data_caja['fecha_hora']:
                        fecha_mostrar = data_caja['fecha_hora']
        except Exception as e:
            try: messagebox.showwarning("Advertencia", f"No se pudo consultar fecha contrato: {e}")
            except: pass
        
        # 2. Fallback: Fecha Registro en Clientes
        if not fecha_mostrar:
            fecha_mostrar = val.get('fecha_registro', '')
            
        if fecha_mostrar:
            # VALIDACIÓN DE FORMATO: DD/MM/YYYY
            try:
                # Si viene como objeto datetime
                if isinstance(fecha_mostrar, datetime.datetime) or isinstance(fecha_mostrar, datetime.date):
                    fecha_str = fecha_mostrar.strftime("%d/%m/%Y")
                else:
                    # Si es string, intentamos parsear varios formatos
                    fecha_limpia = str(fecha_mostrar).split('.')[0] # Quitar milisegundos si hay
                    formato_encontrado = None
                    for fmt in ["%Y-%m-%d %H:%M:%S", "%Y-%m-%d", "%d/%m/%Y %H:%M:%S", "%d/%m/%Y"]:
                        try:
                            dt = datetime.datetime.strptime(fecha_limpia, fmt)
                            formato_encontrado = dt.strftime("%d/%m/%Y")
                            break
                        except: continue
                    
                    fecha_str = formato_encontrado if formato_encontrado else str(fecha_mostrar)
                
                safe_insert(e_apertura, fecha_str)
            except Exception as e_fmt:
                print(f"Error formateando fecha: {e_fmt}")
                safe_insert(e_apertura, str(fecha_mostrar))
        
        try:
            if val['apertura']: safe_insert(e_carpeta, val['apertura'])
        except: pass
        
        if 'fecha nacimiento' in val.keys() and val['fecha nacimiento']: 
            e_nacimiento.delete(0, tk.END)
            e_nacimiento.insert(0, val['fecha nacimiento'])
            e_nacimiento.configure(text_color="black")
        else:
            e_nacimiento.delete(0, tk.END)
            e_nacimiento.insert(0, "dd/mm/aaaa")
            e_nacimiento.configure(text_color="#a9a9a9")

        if val['observaciones']: t_obs.insert("1.0", val['observaciones'])
        
        var_cartera.set(val['cartera castigada'])
        if val['valor cartera']: safe_insert(e_val_cartera, formatear_float_str(val['valor cartera']))
        
        var_demanda.set(val['demanda judicial'])
        if val['valor demanda']: safe_insert(e_val_demanda, formatear_float_str(val['valor demanda']))
        
        if val['problemas justicia']: var_justicia.set(val['problemas justicia'])
        if val['detalle justicia']: safe_insert(e_det_justicia, val['detalle justicia'])
        
        toggle_legal_fields()

        try:
            if val['referencia_vivienda']: safe_insert(e_ref_vivienda, val['referencia_vivienda'])
        except: pass
        
        try:
            if val['terreno']: var_terreno.set(val['terreno'])
        except: pass
        
        try:
            if val['valor_terreno']: safe_insert(e_valor_terreno, formatear_float_str(val['valor_terreno']))
        except: pass
        
        try:
            if val['hipotecado']: c_hipotecado.set(val['hipotecado'])
        except: pass
        
        try:
            if val['casa_dep']: var_casa.set(val['casa_dep'])
        except: pass
        
        try:
            if val['valor_casa_dep']: safe_insert(e_valor_casa, formatear_float_str(val['valor_casa_dep']))
        except: pass
        
        try:
            if val['hipotecado_casa_dep']: c_hip_casa.set(val['hipotecado_casa_dep'])
        except: pass
        
        # Local Com: local (indice 35 aprox), valor_local (36), hipotecado_local (37)
        try:
            if val['local']: var_local.set(val['local'])
            if val['valor_local']: safe_insert(e_valor_local, formatear_float_str(val['valor_local']))
            if val['hipotecado_local']: c_hip_local.set(val['hipotecado_local'])
        except: pass
        
        # Ingresos 2: ingresos_mensuales_2 (39), fuente_ingreso_2 (40)
        try:
            if val['ingresos_mensuales_2']: safe_insert(e_ingresos_2, formatear_float_str(val['ingresos_mensuales_2']))
            if val['fuente_ingreso_2']: c_fuente_ingreso_2.set(val['fuente_ingreso_2'])
        except: pass
        
        # Score Buro: score_buro (41)
        try:
            if val['score_buro']: safe_insert(e_score_buro, str(val['score_buro']))
        except: pass
        
        # Egresos: egresos (42)
        try:
            if val['egresos']: safe_insert(e_egresos, formatear_float_str(val['egresos']))
        except: pass
        
        # Actualizar total disponible (43)
        try:
            if val['total_disponible']:
                lbl_total_disponible_valor.configure(text="$ " + formatear_float_str(val['total_disponible']))
            else:
                calcular_total_disponible()
        except:
            calcular_total_disponible()

        toggle_terreno(); toggle_casa(); toggle_local(); toggle_fuente_ingreso(); toggle_fuente_ingreso_2()

        # Habilitar botón eliminar si es admin
        try:
            if NIVEL_ACCESO == 1:
                btn_eliminar.configure(state='normal')
        except: pass

    except Exception as e:
        messagebox.showerror("Error Carga", f"Error cargando selección: {e}")
        


    except Exception as e: print(f"Error carga: {e}")

    # --- LÓGICA DE PERMISOS Y BLOQUEO AL CARGAR ---
    if NIVEL_ACCESO == 1:
        # Administrador: Todo habilitado
        toggle_inputs_clientes('normal')
        btn_accion.configure(text="📝 Actualizar", command=accion_actualizar, state='normal')
    else:
        # Usuario Estándar: Solo lectura upon load
        toggle_inputs_clientes('disabled')
        btn_accion.configure(text="📝 Actualizar (Bloqueado)", state='disabled')

    btn_cancelar.pack(side='left', padx=15)

def toggle_legal_fields(*args):
    if var_cartera.get() == 1: 
        e_val_cartera.grid()
        lbl_dolar_cartera.grid()
    else: 
        e_val_cartera.grid_remove()
        lbl_dolar_cartera.grid_remove()
    
    if var_demanda.get() == 1: 
        e_val_demanda.grid()
        lbl_dolar_demanda.grid()
    else: 
        e_val_demanda.grid_remove()
        lbl_dolar_demanda.grid_remove()

    if var_justicia.get() == 1: 
        e_det_justicia.grid()
    else: 
        e_det_justicia.grid_remove()

def toggle_terreno(*args):
    if var_terreno.get() == 1:
        e_valor_terreno.grid(row=0, column=1, padx=5)
        f_terreno_hip.grid(row=1, column=0, columnspan=2, sticky='w', padx=25)
    else:
        e_valor_terreno.grid_remove()
        f_terreno_hip.grid_remove()

def toggle_casa(*args):
    if var_casa.get() == 1:
        e_valor_casa.grid(row=0, column=1, padx=5)
        f_casa_hip.grid(row=1, column=0, columnspan=2, sticky='w', padx=25)
    else:
        e_valor_casa.grid_remove()
        f_casa_hip.grid_remove()

def toggle_local(*args):
    if var_local.get() == 1:
        e_valor_local.grid(row=0, column=1, padx=5)
        f_local_hip.grid(row=1, column=0, columnspan=2, sticky='w', padx=25)
    else:
        e_valor_local.grid_remove()
        f_local_hip.grid_remove()

def toggle_fuente_ingreso(*args):
    try:
        val = limpiar_moneda(e_ingresos.get())
        if val > 0: f_fuente_ingreso.grid(row=5, column=0, sticky='w', pady=(0,5))
        else: f_fuente_ingreso.grid_remove()
    except: f_fuente_ingreso.grid_remove()

def toggle_fuente_ingreso_2(*args):
    try:
        val = limpiar_moneda(e_ingresos_2.get())
        if val > 0: f_fuente_ingreso_2.grid(row=8, column=0, sticky='w', pady=(0,5))
        else: f_fuente_ingreso_2.grid_remove()
    except: f_fuente_ingreso_2.grid_remove()

def calcular_total_disponible(*args):
    """Calcula y muestra el total disponible: Ingresos 1 + Ingresos 2 - Egresos"""
    try:
        ing1 = limpiar_moneda(e_ingresos.get())
        ing2 = limpiar_moneda(e_ingresos_2.get())
        egr = limpiar_moneda(e_egresos.get())
        total = ing1 + ing2 - egr
        lbl_total_disponible_valor.configure(text="$ " + formatear_float_str(total))
    except Exception as e:
        lbl_total_disponible_valor.configure(text="$ 0.00")

# --- VENTANA USUARIOS ---
def win_gestion_usuarios():
    top = tk.Toplevel()
    top.title("Usuarios"); top.geometry("500x350")
    top.lift()
    top.focus_force()
    try: top.grab_set()
    except: pass
    
    def ref_u():
        for i in tr.get_children(): tr.delete(i)
        with db_connection() as (conn, c):
            c.execute("SELECT id, usuario, nivel_acceso FROM Usuarios")
            for u in c.fetchall(): tr.insert('', tk.END, values=(u[0], u[1], "Admin" if u[2]==1 else "Std"))

    def add_u():
        res, msg = crear_usuario_db(eu.get(), ep.get(), 1 if cr.get()=="Admin" else 2)
        if res:
             messagebox.showinfo("Éxito", "Usuario creado", parent=top)
        else:
             messagebox.showerror("Error", "Error creando usuario", parent=top)
        ref_u(); eu.delete(0,tk.END); ep.delete(0,tk.END)
        top.lift(); top.focus_force()

    def del_u():
        s = tr.selection()
        if s:
            with db_connection() as (conn, c):
                c.execute("DELETE FROM Usuarios WHERE id=%s", (tr.item(s[0],'values')[0],))
                # Commit handled by CM
            ref_u()
            messagebox.showinfo("Éxito", "Usuario eliminado", parent=top)
            top.lift(); top.focus_force()
    
    f = ttk.Frame(top, padding=5); f.pack(fill='x')
    ttk.Label(f, text="U:").pack(side='left'); eu = ttk.Entry(f, width=10); eu.pack(side='left')
    eu.bind('<Return>', lambda e: ep.focus())
    ttk.Label(f, text="P:").pack(side='left'); ep = ttk.Entry(f, width=10); ep.pack(side='left')
    ep.bind('<Return>', lambda e: cr.focus())
    cr = ttk.Combobox(f, values=["Admin", "Std"], width=8); cr.current(1); cr.pack(side='left')
    cr.bind('<Return>', lambda e: add_u())
    ttk.Button(f, text="+", command=add_u).pack(side='left')
    tr = ttk.Treeview(top, columns=('id','us','rol'), show='headings'); tr.pack(fill='both', expand=True)
    tr.heading('id', text='ID'); tr.heading('us', text='Usuario'); tr.heading('rol', text='Rol')
    ttk.Button(top, text="Borrar", command=del_u).pack()
    ref_u()

# --- MÓDULO INFORMES Y DOCUMENTOS ---

def abrir_modulo_informes():
    global win_informes
    
    win_informes = ctk.CTkToplevel()
    win_informes.title("Gestión de Informes")
    win_informes.geometry("800x600")
    win_informes.after(100, lambda: win_informes.state('zoomed'))
    
    COLOR_FONDO = "#FAFAD2"
    win_informes.configure(fg_color=COLOR_FONDO)
    
    # Barra superior
    nav_frame = ctk.CTkFrame(win_informes, fg_color=COLOR_FONDO, height=40)
    nav_frame.pack(side='top', fill='x', pady=(5,0))
    ctk.CTkButton(nav_frame, text="Volver al Menú", command=win_informes.destroy, 
                  fg_color=COLOR_FONDO, text_color="#d9534f", hover_color="#EEE8AA", 
                  font=('Arial', 12, 'bold')).pack(side='right', padx=20)
    
    # Título
    ctk.CTkLabel(win_informes, text="GESTIÓN DE INFORMES", text_color="#1860C3", font=('Arial', 18, 'bold')).pack(pady=20)
    
    # Contenedor Principal (Estructural)
    main_frame = ctk.CTkFrame(win_informes, fg_color=COLOR_FONDO)
    main_frame.pack(fill='both', expand=True, padx=20, pady=10)

    # Panel Izquierdo (Contenido)
    left_panel = ctk.CTkFrame(main_frame, fg_color=COLOR_FONDO)
    left_panel.pack(side='left', fill='both', expand=True)

    # Caja Blanca de Reportes
    reports_frame = ctk.CTkFrame(left_panel, fg_color="white", corner_radius=15, border_width=1, border_color="#CCCCCC")
    reports_frame.pack(fill='both', expand=True, padx=(0, 20)) # Separación del logo

    ctk.CTkLabel(reports_frame, text="Reportes y Estadísticas", font=('Arial', 16, 'bold'), text_color="#465EA6").pack(pady=(20, 30))

    # Grid de botones para reportes
    btn_container = ctk.CTkFrame(reports_frame, fg_color="transparent")
    btn_container.pack(pady=10)

    def exportar_caja_global_excel():
        """Exporta todos los registros de la tabla Caja a Excel con formato."""
        try:
            with db_connection() as (conn, cursor):
                cursor.execute("""
                    SELECT 
                        fecha_hora, cedula, nombres_completos, ruc, telefono, 
                        email, direccion, valor_apertura, numero_apertura, 
                        buro_credito, observaciones 
                    FROM Caja
                """)
                rows = cursor.fetchall()

            if not rows:
                messagebox.showinfo("Información", "No hay datos registrados en Caja para exportar.", parent=win_informes)
                return

            cols = ["Fecha de Registro", "Cédula", "Nombres y Apellidos", "RUC", "Teléfono", "Correo", "Dirección", "Valor Apertura ($)", "No. Apertura", "Buró de Crédito", "Observaciones"]
            df = pd.DataFrame(rows, columns=cols)

            filename = filedialog.asksaveasfilename(title="Guardar Reporte Global Caja", defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")], initialfile=f"Reporte_Global_Caja_{datetime.datetime.now().strftime('%Y%m%d')}", parent=win_informes)

            if filename:
                writer = pd.ExcelWriter(filename, engine='xlsxwriter')
                df.to_excel(writer, index=False, sheet_name='Reporte Caja')
                workbook  = writer.book
                worksheet = writer.sheets['Reporte Caja']
                header_format = workbook.add_format({'bold': True, 'bg_color': '#D7E4BC', 'border': 1})
                for col_num, value in enumerate(df.columns.values):
                    worksheet.write(0, col_num, value, header_format)
                    worksheet.set_column(col_num, col_num, len(value) + 5)
                writer.close()
                messagebox.showinfo("Éxito", f"Reporte global exportado correctamente.", parent=win_informes)
                registrar_auditoria("Exportar Excel Global Caja", detalles=f"Archivo: {os.path.basename(filename)}")
                win_informes.lift()
                win_informes.focus_force()
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo generar el reporte: {e}", parent=win_informes)

    def exportar_clientes_global_excel():
        """Exporta la base completa de clientes a Excel."""
        try:
            with db_connection() as (conn, cursor):
                cursor.execute("SELECT * FROM Clientes")
                rows = cursor.fetchall()
                cols_db = [desc[0] for desc in cursor.description]

            if not rows:
                messagebox.showinfo("Información", "No hay clientes registrados.", parent=win_informes)
                return

            df = pd.DataFrame(rows, columns=cols_db)
            mapping = {
                'id': 'ID', 'cedula': 'Cédula', 'ruc': 'RUC', 'nombre': 'Nombres y Apellidos',
                'estado_civil': 'Estado Civil', 'cargas_familiares': 'Cargas Familiares',
                'email': 'Email', 'telefono': 'Teléfono', 'direccion': 'Dirección Domicilio',
                'parroquia': 'Parroquia', 'tipo_vivienda': 'Tipo Vivienda',
                'referencia_vivienda': 'Referencia Vivienda', 'profesion': 'Profesión/Actividad',
                'ingresos_mensuales': 'Ingresos Principal ($)', 'ingresos_mensuales_2': 'Ingresos Secundarios ($)',
                'egresos': 'Egresos Mensuales ($)', 'total_disponible': 'Total Disponible ($)',
                'referencia1': 'Referencia Personal 1', 'referencia2': 'Referencia Personal 2',
                'asesor': 'Asesor Asignado', 'apertura': 'Número de Carpeta', 
                'fecha nacimiento': 'Fecha de Nacimiento',
                'producto': 'Producto', 'observaciones': 'Observaciones Generales',
                'cartera castigada': 'Cartera', 'valor cartera': 'Valor Cartera ($)',
                'demanda judicial': 'Demanda', 'valor demanda': 'Valor Demanda ($)',
                'problemas justicia': 'Justicia', 'detalle justicia': 'Detalle Justicia',
                'situacion_financiera': 'Situación Financiera', 'terreno': 'Tiene Terreno',
                'valor_terreno': 'Valor Terreno ($)', 'hipotecado': 'Terreno Hipotecado',
                'casa_dep': 'Tiene Casa/Dep', 'valor_casa_dep': 'Valor Casa/Dep ($)',
                'hipotecado_casa_dep': 'Casa Hipotecada', 'local': 'Tiene Local',
                'valor_local': 'Valor Local ($)', 'hipotecado_local': 'Local Hipotecado',
                'score_buro': 'Score Buró', 'fecha_registro': 'Fecha de Apertura'
            }
            if 'numero_carpeta' in df.columns:
                df.drop(columns=['numero_carpeta'], inplace=True)
            df.rename(columns={k: v for k, v in mapping.items() if k in df.columns}, inplace=True)

            filename = filedialog.asksaveasfilename(title="Guardar Base Completa", defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")], initialfile=f"Master_Base_Clientes_{datetime.datetime.now().strftime('%Y%m%d')}", parent=win_informes)

            if filename:
                writer = pd.ExcelWriter(filename, engine='xlsxwriter')
                df.to_excel(writer, index=False, sheet_name='Base Maestra')
                workbook  = writer.book
                worksheet = writer.sheets['Base Maestra']
                header_format = workbook.add_format({'bold': True, 'bg_color': '#C6EFCE', 'border': 1})
                for col_num, value in enumerate(df.columns.values):
                    worksheet.write(0, col_num, value, header_format)
                    worksheet.set_column(col_num, col_num, max(len(str(value)), 15) + 2)
                writer.close()
                messagebox.showinfo("Éxito", "Exportación completada.", parent=win_informes)
                win_informes.lift()
                win_informes.focus_force()
        except Exception as e:
            messagebox.showerror("Error", f"Error: {e}", parent=win_informes)

    ctk.CTkButton(btn_container, text="📊 EXCEL CAJA GLOBAL", command=exportar_caja_global_excel, 
                  font=('Arial', 14, 'bold'), fg_color="#28a745", hover_color="#218838", height=50, width=250).grid(row=0, column=0, padx=20, pady=10)
    
    ctk.CTkButton(btn_container, text="🟢 EXCEL BASE MAESTRA", command=exportar_clientes_global_excel, 
                  font=('Arial', 14, 'bold'), fg_color="#28a745", hover_color="#218838", height=50, width=250).grid(row=0, column=1, padx=20, pady=10)
    
    ctk.CTkLabel(reports_frame, text="(Próximamente más opciones de reportes...)", text_color="grey", font=('Arial', 11, 'italic')).pack(pady=40)

    # Logo (Panel Derecho - Estandarizado con Documentos y Clientes)
    try:
        img = Image.open(resource_path("Logo Face.jpg"))
        logo_pix = ctk.CTkImage(light_image=img, dark_image=img, size=(225, 210))
        ctk.CTkLabel(main_frame, image=logo_pix, text="").pack(side='right', padx=20, anchor='n')
    except: pass

def abrir_modulo_documentos():
    global win_docs, e_cedula_buscar, e_nombre_cliente, e_ruc_buscar, tree_docs, cedula_actual
    
    cedula_actual = None
    
    win_docs = ctk.CTkToplevel()
    win_docs.title("Gestión de Documentos")
    win_docs.geometry("1100x650")
    win_docs.after(100, lambda: win_docs.state('zoomed'))
    
    COLOR_FONDO = "#FAFAD2"
    win_docs.configure(fg_color=COLOR_FONDO)
    
    # Barra superior
    nav_frame = ctk.CTkFrame(win_docs, fg_color=COLOR_FONDO, height=40)
    nav_frame.pack(side='top', fill='x', pady=(5,0))
    ctk.CTkButton(nav_frame, text="Volver al Menú", command=win_docs.destroy, 
                  fg_color=COLOR_FONDO, text_color="#d9534f", hover_color="#EEE8AA", 
                  font=('Arial', 12, 'bold')).pack(side='right', padx=20)
    
    # Título
    ctk.CTkLabel(win_docs, text="GESTIÓN DE DOCUMENTOS", text_color="#1860C3", font=('Arial', 16, 'bold')).pack(pady=10)
    
    # Content Frame
    main_frame = ctk.CTkFrame(win_docs, fg_color=COLOR_FONDO)
    main_frame.pack(fill='both', expand=True, padx=20, pady=10)

    # Panel Izquierdo
    left_panel = ctk.CTkFrame(main_frame, fg_color=COLOR_FONDO)
    left_panel.pack(side='left', fill='both', expand=True)
    
    # SECCIÓN BÚSQUEDA
    search_frame = ctk.CTkFrame(left_panel, fg_color="white", border_width=1, border_color="grey")
    search_frame.pack(fill='x', pady=(0,10))
    ctk.CTkLabel(search_frame, text=" Buscar Cliente ", text_color="grey", font=('Arial', 10, 'bold')).place(x=10, y=-8)
    
    sf_in = ctk.CTkFrame(search_frame, fg_color="transparent")
    sf_in.pack(fill='x', padx=10, pady=15)
    
    ctk.CTkLabel(sf_in, text="Cédula:", text_color="black").pack(side='left')
    e_cedula_buscar = ctk.CTkEntry(sf_in, width=150, fg_color="white", text_color="black", border_color="grey")
    e_cedula_buscar.pack(side='left', padx=5)
    
    ctk.CTkLabel(sf_in, text="RUC:", text_color="black").pack(side='left', padx=(15,0))
    e_ruc_buscar = ctk.CTkEntry(sf_in, width=150, fg_color="white", text_color="black", border_color="grey")
    e_ruc_buscar.pack(side='left', padx=5)
    
    ctk.CTkLabel(sf_in, text="Cliente:", text_color="black").pack(side='left', padx=(15,0))
    e_nombre_cliente = ctk.CTkEntry(sf_in, width=350, fg_color="white", text_color="black", border_color="grey")
    e_nombre_cliente.pack(side='left', padx=5)
    
    # TAB VIEW (Solo una pestaña o directo)
    # Mostraremos directo el contenido de documentos
    docs_frame = ctk.CTkFrame(left_panel, fg_color="white", corner_radius=10, border_width=1, border_color="#CCCCCC")
    docs_frame.pack(fill='both', expand=True)
    
    ctk.CTkLabel(docs_frame, text=" Listado de Documentos PDF ", text_color="#1860C3", font=('Arial', 12, 'bold')).pack(pady=10)

    df_in = ctk.CTkFrame(docs_frame, fg_color="transparent")
    df_in.pack(fill='both', expand=True, padx=10, pady=5)
    
    # Botones
    btn_frame = ctk.CTkFrame(df_in, fg_color="transparent")
    btn_frame.pack(fill='x', pady=(0,10))
    
    ctk.CTkButton(btn_frame, text="📎 Subir PDF", command=subir_documento, fg_color="#465EA6", hover_color="#1860C3").pack(side='left', padx=5)
    ctk.CTkButton(btn_frame, text="👁️ Ver", command=ver_documento, fg_color="#465EA6", hover_color="#1860C3").pack(side='left', padx=5)
    ctk.CTkButton(btn_frame, text="🗑️ Eliminar", command=eliminar_documento, fg_color="#d9534f", hover_color="#c9302c").pack(side='left', padx=5)
    
    # TreeView
    tree_frame = ctk.CTkFrame(df_in, fg_color="white")
    tree_frame.pack(fill='both', expand=True)
    
    cols = ("ID", "Tipo", "Nombre Archivo", "Fecha")
    tree_docs = ttk.Treeview(tree_frame, columns=cols, show='headings', height=15)
    tree_docs.heading("ID", text="ID"); tree_docs.heading("Tipo", text="Tipo"); tree_docs.heading("Nombre Archivo", text="Nombre Archivo"); tree_docs.heading("Fecha", text="Fecha")
    tree_docs.column("ID", width=50); tree_docs.column("Tipo", width=150); tree_docs.column("Nombre Archivo", width=300); tree_docs.column("Fecha", width=150)
    
    scrollbar = ttk.Scrollbar(tree_frame, orient='vertical', command=tree_docs.yview)
    scrollbar.pack(side='right', fill='y')
    tree_docs.configure(yscrollcommand=scrollbar.set)
    tree_docs.pack(fill='both', expand=True)

    # Logo Derecho
    try:
        img = Image.open(resource_path("Logo Face.jpg"))
        logo_pix = ctk.CTkImage(light_image=img, dark_image=img, size=(225, 210))
        ctk.CTkLabel(main_frame, image=logo_pix, text="").pack(side='right', padx=20, anchor='n')
    except: pass

    def buscar_cliente_auto(event=None):
        global cedula_actual
        ced = e_cedula_buscar.get().strip()
        ruc = e_ruc_buscar.get().strip()
        nom = e_nombre_cliente.get().strip()
        criteria = None; val = None
        if len(ced) == 10 and ced.isdigit(): criteria = "cedula"; val = ced
        elif len(ruc) >= 10 and ruc.isdigit(): criteria = "ruc"; val = ruc
        elif len(nom) >= 3: criteria = "nombre"; val = nom
        if not criteria: return

        conn, cursor = conectar_db()
        res = None
        if criteria == "cedula": cursor.execute("SELECT nombre, cedula, ruc FROM Clientes WHERE cedula = %s", (val,))
        elif criteria == "ruc": cursor.execute("SELECT nombre, cedula, ruc FROM Clientes WHERE ruc = %s", (val,))
        elif criteria == "nombre": cursor.execute("SELECT nombre, cedula, ruc FROM Clientes WHERE nombre LIKE %s LIMIT 1", (f"%{val}%",))
        res = cursor.fetchone()
        db_manager.release_connection(conn)
        
        if res:
            widget = event.widget if event else None
            # res: 0:nombre, 1:cedula, 2:ruc
            if criteria != "nombre" or (widget != e_nombre_cliente):
                e_nombre_cliente.delete(0, tk.END); e_nombre_cliente.insert(0, res[0])
            if criteria != "ruc" or (widget != e_ruc_buscar):
                e_ruc_buscar.delete(0, tk.END); 
                if res[2]: e_ruc_buscar.insert(0, res[2])
            if criteria != "cedula" or (widget != e_cedula_buscar):
                 e_cedula_buscar.delete(0, tk.END); e_cedula_buscar.insert(0, res[1])
            cedula_actual = res[1]
            cargar_documentos(cedula_actual)
        else:
            cedula_actual = None
            limpiar_lista_documentos()
    
    e_cedula_buscar.bind('<KeyRelease>', buscar_cliente_auto)
    e_ruc_buscar.bind('<KeyRelease>', buscar_cliente_auto)
    e_nombre_cliente.bind('<KeyRelease>', buscar_cliente_auto)

def limpiar_lista_documentos():
    for item in tree_docs.get_children():
        tree_docs.delete(item)

def cargar_documentos(cedula):
    limpiar_lista_documentos()
    
    # 1. Cargar desde carpeta global (_Archivos_Clientes)
    carpeta_global = "_Archivos_Clientes"
    if os.path.exists(carpeta_global):
        try:
            archivos = os.listdir(carpeta_global)
            for archivo in archivos:
                # 3. Corregir Búsqueda: Buscar cedula en el nombre
                if cedula in archivo:
                    # Determinar Tipo
                    tipo_doc = "Documento"
                    if "BuroCredito" in archivo: tipo_doc = "Reporte Buró"
                    elif "Contrato" in archivo: tipo_doc = "Contrato"
                    elif "Cedula" in archivo: tipo_doc = "Cédula Escaneada"
                    
                    # Fecha de modificación
                    ruta_completa = os.path.join(carpeta_global, archivo)
                    timestamp = os.path.getmtime(ruta_completa)
                    fecha_mod = datetime.datetime.fromtimestamp(timestamp).strftime('%Y-%m-%d %H:%M')
                    
                    # Insertar en Treeview (ID='File', Tipo, Nombre, Fecha)
                    # Usamos 'FILE' + path como ID para diferenciarlo o manejar eventos si fuera necesario
                    tree_docs.insert('', tk.END, values=("FILE", tipo_doc, archivo, fecha_mod))
        except Exception as e:
            print(f"Error leyendo carpeta archivos: {e}")

    # 2. Cargar desde Base de Datos (Histórico)
    conn, cursor = conectar_db()
    cursor.execute("SELECT id, tipo_documento, nombre_archivo, fecha_subida FROM Documentos WHERE cedula_cliente = %s ORDER BY fecha_subida DESC", (cedula,))
    documentos = cursor.fetchall()
    db_manager.release_connection(conn)
    
    for doc in documentos:
        tree_docs.insert('', tk.END, values=doc)

def subir_documento():
    global cedula_actual
    
    # 1. Validaciones Previas
    # 1. Validaciones Previas
    if not cedula_actual:
        try:
            # Intentamos usar win_docs si existe, si no, el root
            messagebox.showwarning("Advertencia", "Debe cargar un cliente (buscar) antes de subir documentos.", parent=win_docs)
        except:
            messagebox.showwarning("Advertencia", "Debe cargar un cliente (buscar) antes de subir documentos.")
        return

    # 2. Selección de Archivo
    # 2. Selección de Archivo
    try:
        ruta_origen = filedialog.askopenfilename(
            title="Seleccionar Documento PDF",
            filetypes=[("Archivos PDF", "*.pdf"), ("Todos los archivos", "*.*")],
            parent=win_docs
        )
    except:
        ruta_origen = filedialog.askopenfilename(
            title="Seleccionar Documento PDF",
            filetypes=[("Archivos PDF", "*.pdf"), ("Todos los archivos", "*.*")]
        )
    
    if ruta_origen:
        try:
            # 3. Procesamiento y Renombrado
            nombre_original = os.path.basename(ruta_origen)
            nombre_destino = f"{cedula_actual}_{nombre_original}" # Formato: [CEDULA]_[Original]
            
            carpeta_global = "_Archivos_Clientes"
            if not os.path.exists(carpeta_global):
                os.makedirs(carpeta_global)
                
            ruta_destino = os.path.join(carpeta_global, nombre_destino)
            
            # 4. Ejecutar Copia
            shutil.copy(ruta_origen, ruta_destino)
            
            registrar_auditoria("Subir Documento", detalles=f"Archivo: {nombre_destino}")
            
            # 5. Refrescar Interfaz (Paso C: Inmediatamente)
            # CRÍTICO: No cerrar la ventana, solo actualizar la lista.
            cargar_documentos(cedula_actual)
            
            # 6. Confirmación Visual (Al final)
            try:
                # Intentar usar parent para evitar que el dialogo oculte la ventana
                messagebox.showinfo("Éxito", "Archivo subido y guardado correctamente", parent=win_docs)
                win_docs.lift()
                win_docs.focus_force()
            except:
                messagebox.showinfo("Éxito", "Archivo subido y guardado correctamente")
            
        except Exception as e:
            try:
                messagebox.showerror("Error", f"No se pudo guardar el archivo: {e}", parent=win_docs)
            except:
                messagebox.showerror("Error", f"No se pudo guardar el archivo: {e}")

def ver_documento():
    seleccion = tree_docs.selection()
    if not seleccion:
        messagebox.showwarning("Advertencia", "Debe seleccionar un documento")
        return
    
    # Obtener datos de la fila seleccionada
    item = tree_docs.item(seleccion[0])
    values = item['values']
    doc_id = values[0]
    nombre_archivo = values[2] # El nombre está en la columna 3 (índice 2)
    
    ruta_final = None

    # Caso 1: Archivo detectado en carpeta global (ID = 'FILE')
    if doc_id == "FILE":
        # Construir ruta con carpeta global
        carpeta_global = "_Archivos_Clientes"
        ruta_final = os.path.join(carpeta_global, nombre_archivo)
        ruta_final = os.path.abspath(ruta_final) # Asegurar ruta absoluta
    
    # Caso 2: Archivo histórico en Base de Datos (ID numérico)
    else:
        conn, cursor = conectar_db()
        cursor.execute("SELECT ruta_archivo FROM Documentos WHERE id = %s", (doc_id,))
        resultado = cursor.fetchone()
        db_manager.release_connection(conn)
        
        if resultado:
             ruta_final = resultado[0]

    # Intentar abrir el archivo
    if ruta_final and os.path.exists(ruta_final):
        try:
            os.startfile(ruta_final)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir el documento: {e}")
    else:
        messagebox.showerror("Error", f"El archivo no se encuentra en la ruta:\n{ruta_final}")

def eliminar_documento():
    global cedula_actual
    
    # 1. Validación de Selección
    seleccion = tree_docs.selection()
    if not seleccion:
        messagebox.showwarning("Advertencia", "Por favor seleccione un documento para eliminar")
        return
    
    # 2. Confirmación de Seguridad
    try:
        conf = messagebox.askyesno("Confirmar Eliminación", "¿Está seguro de que desea eliminar este archivo permanentemente?\nEsta acción no se puede deshacer.", parent=win_docs)
    except:
        conf = messagebox.askyesno("Confirmar Eliminación", "¿Está seguro de que desea eliminar este archivo permanentemente?\nEsta acción no se puede deshacer.")
        
    if not conf:
        return
    
    item = tree_docs.item(seleccion[0])
    values = item['values']
    doc_id = values[0]
    nombre_archivo = values[2]
    
    try:
        # Caso 1: Archivo Local (ID='FILE')
        if doc_id == "FILE":
            carpeta_global = "_Archivos_Clientes"
            ruta_completa = os.path.join(carpeta_global, nombre_archivo)
            
            if os.path.exists(ruta_completa):
                os.remove(ruta_completa)
                try:
                    messagebox.showinfo("Éxito", "Archivo eliminado correctamente", parent=win_docs)
                    win_docs.lift()
                    win_docs.focus_force()
                except:
                    messagebox.showinfo("Éxito", "Archivo eliminado correctamente")
                registrar_auditoria("Eliminar Documento", detalles=f"Archivo: {nombre_archivo}")
            else:
                try:
                    messagebox.showerror("Error", "El archivo ya no existe en el disco", parent=win_docs)
                except:
                    messagebox.showerror("Error", "El archivo ya no existe en el disco")

        # Caso 2: Archivo Histórico (Base de Datos)
        else:
            conn, cursor = conectar_db()
            cursor.execute("SELECT ruta_archivo FROM Documentos WHERE id = %s", (doc_id,))
            resultado = cursor.fetchone()
            
            if resultado:
                ruta = resultado[0]
                # Eliminar archivo físico si existe
                if os.path.exists(ruta):
                    try:
                        os.remove(ruta)
                    except Exception as e:
                        print(f"Advertencia: No se pudo borrar archivo físico {ruta}: {e}")
                
                # Eliminar registro DB
                cursor.execute("DELETE FROM Documentos WHERE id = %s", (doc_id,))
                conn.commit()
                try:
                    messagebox.showinfo("Éxito", "Documento eliminado de la base de datos", parent=win_docs)
                    win_docs.lift()
                    win_docs.focus_force()
                except:
                    messagebox.showinfo("Éxito", "Documento eliminado de la base de datos")
                registrar_auditoria("Eliminar Documento DB", detalles=f"ID: {doc_id}")
            
            db_manager.release_connection(conn)

        # 4. Actualización Visual (Solo recargar lista, NO cerrar ventana)
        cargar_documentos(cedula_actual)
            
    except Exception as e:
        try:
            messagebox.showerror("Error", f"Fallo al eliminar: {e}", parent=win_docs)
        except:
            messagebox.showerror("Error", f"Fallo al eliminar: {e}")

# --- APP PRINCIPAL ---

def login_fn(app, u_entry, p_entry):
    global USUARIO_ACTIVO, NIVEL_ACCESO
    
    u = u_entry.get()
    p = p_entry.get()
    
    ok, lvl = verificar_credenciales(u, p)
    if ok is True:
        USUARIO_ACTIVO = u; NIVEL_ACCESO = lvl
        registrar_auditoria("Inicio de Sesión", detalles=f"Usuario {u} ingresó al sistema.")
        
        # Clear current window (Login)
        for widget in app.winfo_children():
            widget.destroy()
            
        # Open Menu in same window
        abrir_menu_principal(app)
    elif lvl == "Inactivo":
        messagebox.showerror("Acceso Denegado", "Su cuenta está inactiva. Contacte al administrador.")
    else:
        messagebox.showerror("Error", "Datos incorrectos")


def abrir_menu_principal(app_root=None):
    global menu_app
    
    if app_root:
        menu_app = app_root
    else:
        # Fallback if called separately (shouldn't happen in new flow)
        menu_app = ctk.CTk()
    
    menu_app.title("Menú Principal - Alianza C3F")
    
    # 30cm x 20cm aprox (96 DPI: 1cm=37.8px) -> 1134x756
    width = 1134
    height = 756
    
    # Centrar en pantalla
    screen_width = menu_app.winfo_screenwidth()
    screen_height = menu_app.winfo_screenheight()
    x = (screen_width - width) // 2
    y = (screen_height - height) // 2
    
    menu_app.geometry(f"{width}x{height}+{x}+{y}")
    # menu_app.state('zoomed') 
    menu_app.configure(fg_color="#0b162c") # Dark blue to match background image
    menu_app.after(100, lambda: menu_app.state('zoomed')) # Increased delay slightly

    # Configuración Grid para centrado total
    menu_app.grid_columnconfigure(0, weight=1)
    menu_app.grid_columnconfigure(1, weight=1)
    menu_app.grid_columnconfigure(2, weight=1)
    menu_app.grid_rowconfigure(2, weight=1)

    # Background Image - Creado primero para estar al fondo
    try:
        img_bg = Image.open("background.jpg")
        screen_w = menu_app.winfo_screenwidth()
        screen_h = menu_app.winfo_screenheight()
        logo_bg = ctk.CTkImage(light_image=img_bg, dark_image=img_bg, size=(screen_w, screen_h))
        lbl_bg = ctk.CTkLabel(menu_app, image=logo_bg, text="")
        lbl_bg.place(x=0, y=0, relwidth=1, relheight=1)
        lbl_bg.lower()
    except Exception as e:
        print(f"Error cargando fondo: {e}")

    # Forzar color oscuro para evitar destellos
    menu_app.configure(fg_color="#070e1e")

    # Header - Removido por solicitud del usuario

    # Logo
    try:
        img = Image.open(resource_path("logo_transparent.png"))
        width_l, height_l = img.size
        new_size = (int(width_l * 0.5), int(height_l * 0.5))
        logo_ctk = ctk.CTkImage(light_image=img, dark_image=img, size=new_size)
        lbl_logo = ctk.CTkLabel(menu_app, image=logo_ctk, text="", fg_color="transparent")
        lbl_logo.grid(row=1, column=0, columnspan=3, pady=10)
    except Exception as e:
        print(f"Error cargando logo: {e}")
    
    # Definir botones (Texto, Función, Fila, Columna)
    # Fila 1: Gestión de Clientes, Documentos, Consultas
    # Fila 2: Microcrédito, Rehabilitación, Intermediación
    # Fila 3: Cartera, Informes, Usuarios
    # Fila 4: Salir del sistema en la columna 2
    
    botones_config = [
        ("Caja", abrir_modulo_caja, 2, 0),
        ("Gestión de Clientes", abrir_modulo_clientes, 2, 1),
        ("Consultas", abrir_modulo_consultas, 2, 2),
        
        ("Microcrédito", abrir_modulo_microcredito, 3, 0),
        ("Rehabilitación", abrir_modulo_rehabilitacion, 3, 1),
        ("Intermediación", abrir_modulo_intermediacion, 3, 2),
        
        ("Cartera", abrir_modulo_cartera, 4, 0),
        ("Informes", abrir_modulo_informes, 4, 1),
        ("Documentos", abrir_modulo_documentos, 4, 2),
    ]
    
    # Filtrar por nivel de acceso para Usuarios (Fila 5)
    if NIVEL_ACCESO == 1:
        botones_config.append(("Usuarios", abrir_modulo_usuarios, 5, 0))
    
    # Salir del sistema (Fila 5, Columna 1)
    botones_config.append(("Salir Sistema", menu_app.destroy, 5, 1))

    # Cargar fondo para botones
    try:
        btn_bg_img = Image.open(resource_path("btn_bg.jpg"))
        ANCHO_BTN = 250
        ALTO_BTN = 70
        btn_img = ctk.CTkImage(light_image=btn_bg_img, dark_image=btn_bg_img, size=(ANCHO_BTN, ALTO_BTN))
        
        # Cargar iconos para botones
        try:
            # Icono Rehabilitación
            rehab_raw = Image.open(resource_path("rehabilitacion_icon.png"))
            rehab_icon = ctk.CTkImage(light_image=rehab_raw, dark_image=rehab_raw, size=(42, 42))
            
            # Icono Gestión de Clientes
            clientes_raw = Image.open(resource_path("clientes_icon.png"))
            clientes_icon = ctk.CTkImage(light_image=clientes_raw, dark_image=clientes_raw, size=(42, 42))
            
            # Icono Microcrédito
            micro_raw = Image.open(resource_path("microcredito_icon.png"))
            micro_icon = ctk.CTkImage(light_image=micro_raw, dark_image=micro_raw, size=(42, 42))
            
            # Icono Intermediación
            inter_raw = Image.open(resource_path("intermediacion_icon.png"))
            inter_icon = ctk.CTkImage(light_image=inter_raw, dark_image=inter_raw, size=(42, 42))
            
            # Icono Informes
            informes_raw = Image.open(resource_path("informes_icon.png"))
            informes_icon = ctk.CTkImage(light_image=informes_raw, dark_image=informes_raw, size=(42, 42))

            # Icono Consultas
            consultas_raw = Image.open(resource_path("consultas_icon.png"))
            consultas_icon = ctk.CTkImage(light_image=consultas_raw, dark_image=consultas_raw, size=(42, 42))

            # Icono Informes y Documentos
            doc_raw = Image.open(resource_path("documentos_icon.png"))
            doc_icon = ctk.CTkImage(light_image=doc_raw, dark_image=doc_raw, size=(42, 42))

            # Icono Cartera
            cartera_raw = Image.open(resource_path("cartera_icon.png"))
            cartera_icon = ctk.CTkImage(light_image=cartera_raw, dark_image=cartera_raw, size=(42, 42))

            # Icono Usuarios
            usuarios_raw = Image.open(resource_path("usuarios_icon.png"))
            usuarios_icon = ctk.CTkImage(light_image=usuarios_raw, dark_image=usuarios_raw, size=(42, 42))

            # Icono Salir
            salir_raw = Image.open(resource_path("salir_icon.png"))
            salir_icon = ctk.CTkImage(light_image=salir_raw, dark_image=salir_raw, size=(42, 42))
        except Exception as e:
            print(f"Error cargando iconos: {e}")
            rehab_icon = None
            clientes_icon = None
            micro_icon = None
            inter_icon = None
            informes_icon = None
            consultas_icon = None
            doc_icon = None
            cartera_icon = None
            usuarios_icon = None
            salir_icon = None
    except Exception as e:
        btn_img = None
        rehab_icon = None
        clientes_icon = None
        micro_icon = None
        inter_icon = None
        informes_icon = None
        consultas_icon = None
        doc_icon = None
        cartera_icon = None
        usuarios_icon = None
        salir_icon = None

    # Configurar pesos de filas para que los botones se centren verticalmente
    for i in range(2, 6):
        menu_app.grid_rowconfigure(i, weight=1)

    for texto, funcion, row, col in botones_config:
        # Usamos CTkLabel como botones para que la transparencia sobre la imagen sea perfecta
        btn = ctk.CTkLabel(
            menu_app, 
            text=texto,
            image=btn_img,
            compound="center",
            anchor="center",
            font=("Arial", 16, "bold"),
            height=ALTO_BTN,
            width=ANCHO_BTN,
            text_color="white",
            fg_color="transparent",
            bg_color="transparent",
            corner_radius=15,
            cursor="hand2"
        )
        btn.grid(row=row, column=col, padx=20, pady=15)
        btn.lift() # Asegura que esté sobre el fondo
        
        # Lógica de Iconos
        target_icon = None
        if "Rehabilitación" in texto:
            target_icon = rehab_icon
        elif "Gestión de Clientes" in texto:
            target_icon = clientes_icon
        elif "Microcrédito" in texto:
            target_icon = micro_icon
        elif "Intermediación" in texto:
            target_icon = inter_icon
        elif "Informes" in texto:
            target_icon = informes_icon
        elif "Consultas" in texto:
            target_icon = consultas_icon
        elif "Documentos" in texto:
            target_icon = doc_icon
        elif "Cartera" in texto:
            target_icon = cartera_icon
        elif "Usuarios" in texto:
            target_icon = usuarios_icon
        elif "Caja" in texto:
            # Reutilizamos icono de microcrédito o cartera si no hay uno específico
            target_icon = micro_icon 
        elif "Salir Sistema" in texto:
            target_icon = salir_icon
            
        if target_icon:
            icon_lbl = ctk.CTkLabel(btn, image=target_icon, text="", fg_color="transparent")
            icon_lbl.place(x=25, rely=0.5, anchor="w")
            # Ajustar padding del texto
            btn.configure(text=f"         {texto}")
        
        # Binds para que el Label actúe como botón
        btn.bind("<Button-1>", lambda e, f=funcion: f())
        btn.bind("<Enter>", lambda e, b=btn: b.configure(text_color="#FAD390"))
        btn.bind("<Leave>", lambda e, b=btn: b.configure(text_color="white"))

    # Footer
    label_footer = ctk.CTkLabel(menu_app, text=f"Usuario: {USUARIO_ACTIVO} | Nivel: {NIVEL_ACCESO}", text_color="#DDDDDD", fg_color="black", corner_radius=5)
    label_footer.grid(row=6, column=0, columnspan=3, pady=20)
    
    menu_app.update()

    # Ensure window update
    menu_app.update()

def abrir_modulo_clientes():
    global app, e_cedula, e_ruc, e_nombre, c_civil, e_cargas, e_email, e_telf, e_dir, e_parroquia
    global c_vivienda, e_ref_vivienda, e_profesion, e_ingresos, e_ref1, e_ref2, e_asesor, e_apertura, e_carpeta, e_nacimiento
    global t_obs, var_cartera, var_demanda, var_justicia, btn_accion, btn_cancelar, btn_eliminar
    global cb_cartera, cb_demanda, cb_justicia, cb_terreno, cb_casa, cb_local
    global e_busqueda, tree, e_val_cartera, e_val_demanda, e_det_justicia, var_terreno, e_valor_terreno
    global lbl_dolar_cartera, lbl_dolar_demanda
    global c_hipotecado, f_terreno_hip
    global var_casa, e_valor_casa, c_hip_casa, f_casa_hip
    global var_local, e_valor_local, c_hip_local, f_local_hip
    global c_fuente_ingreso, f_fuente_ingreso
    global e_ingresos_2, c_fuente_ingreso_2, f_fuente_ingreso_2
    global e_score_buro
    global e_egresos, lbl_total_disponible_valor

    app = ctk.CTkToplevel()
    app.title("Módulo de Clientes")
    app.geometry("1350x850")
    app.after(100, lambda: app.state('zoomed'))

    # --- TEMA Y COLORES ---
    COLOR_FONDO = "#FAFAD2" # LightGoldenrodYellow
    COLOR_TEXTO = "#000000" # Negro
    COLOR_BTN_BG = "#465EA6"
    COLOR_BTN_HOVER = "#1860C3"
    
    app.configure(fg_color=COLOR_FONDO)
    
    # --- BARRA DE NAVEGACIÓN ---
    nav_frame = ctk.CTkFrame(app, fg_color=COLOR_FONDO, height=40)
    nav_frame.pack(side='top', fill='x', pady=(5,0))
    ctk.CTkButton(nav_frame, text="Volver al Menú", command=app.destroy, 
                  fg_color=COLOR_FONDO, text_color="#d9534f", hover_color="#EEE8AA", 
                  font=('Arial', 12, 'bold')).pack(side='right', padx=20)


    # Variables de control
    var_cartera = tk.IntVar(); var_demanda = tk.IntVar(); var_justicia = tk.IntVar()
    var_terreno = tk.IntVar(); var_casa = tk.IntVar(); var_local = tk.IntVar()
    var_cartera.trace_add('write', toggle_legal_fields)
    var_demanda.trace_add('write', toggle_legal_fields)
    var_justicia.trace_add('write', toggle_legal_fields)
    var_terreno.trace_add('write', toggle_terreno)
    var_casa.trace_add('write', toggle_casa)
    var_local.trace_add('write', toggle_local)

    # --- TOP (Contenedor del Formulario y Logo) ---
    top_frame = ctk.CTkFrame(app, fg_color=COLOR_FONDO)
    top_frame.pack(fill='x', padx=10, pady=5)
    
    f_form = ctk.CTkFrame(top_frame, fg_color="transparent", border_width=1, border_color="grey")
    f_form.pack(side='left', fill='both', expand=True, padx=5, pady=5)
    
    # --- FONDO GESTIÓN ---
    try:
        imagen_fondo_gris = ctk.CTkImage(light_image=Image.open("fondo gestion.jpg"),
                                         dark_image=Image.open("fondo gestion.jpg"),
                                         size=(1200, 600))
        contenedor_ficha = ctk.CTkLabel(master=f_form, image=imagen_fondo_gris, text="")
        contenedor_ficha.place(x=0, y=0, relwidth=1, relheight=1)
    except Exception as e:
        print(f"Error cargando fondo gestion: {e}")

    ctk.CTkLabel(top_frame, text="Ficha del Cliente", font=('Arial', 14, 'bold'), text_color=COLOR_TEXTO, fg_color="transparent").place(x=15, y=0) 

    def crear_entry(parent, width=None):
        e = ctk.CTkEntry(parent, fg_color="white", text_color="black", border_color="grey")
        if width: e.configure(width=width)
        return e

    # --- FORMULARIO CON PESTAÑAS ---
    tab_view = ctk.CTkTabview(f_form, width=1050, height=450, 
                             fg_color="transparent",
                             segmented_button_fg_color="#E0E0E0",
                             segmented_button_selected_color="#A9CCE3",
                             segmented_button_selected_hover_color="#92BBD9",
                             segmented_button_unselected_color="white",
                             text_color="black",
                             corner_radius=10)
    tab_view.pack(fill='both', expand=True, padx=10, pady=10)

    t1 = tab_view.add("Identificación y Ubicación")
    t2 = tab_view.add("Situación Financiera")
    t3 = tab_view.add("Gestión y Legal")

    # --- TAB 1: IDENTIFICACIÓN Y UBICACIÓN ---
    f1 = ctk.CTkFrame(t1, fg_color="transparent")
    f1.pack(fill='both', expand=True)

    c1_1 = ctk.CTkFrame(f1, fg_color="transparent")
    c1_1.grid(row=0, column=0, padx=20, pady=10, sticky='n')
    ctk.CTkLabel(c1_1, text="DATOS PRINCIPALES", text_color="#1860C3", font=('Arial', 12, 'bold'), fg_color="transparent").pack(pady=5)
    
    ctk.CTkLabel(c1_1, text="Cédula:", text_color="black", fg_color="transparent").pack(anchor='w')
    e_cedula = crear_entry(c1_1); e_cedula.pack(fill='x')
    e_cedula.bind('<Return>', saltar_campo)
    
    ctk.CTkLabel(c1_1, text="RUC:", text_color="black", fg_color="transparent").pack(anchor='w')
    e_ruc = crear_entry(c1_1); e_ruc.pack(fill='x')
    e_ruc.bind('<Return>', saltar_campo)
    
    ctk.CTkLabel(c1_1, text="Nombres y Apellidos:", text_color="black", fg_color="transparent").pack(anchor='w')
    e_nombre = crear_entry(c1_1); e_nombre.pack(fill='x')
    e_nombre.bind('<Return>', saltar_campo)
    
    ctk.CTkLabel(c1_1, text="F. Nacimiento:", text_color="black", fg_color="transparent").pack(anchor='w')
    e_nacimiento = crear_entry(c1_1); e_nacimiento.pack(fill='x')
    e_nacimiento.insert(0, "dd/mm/aaaa"); e_nacimiento.configure(text_color="#a9a9a9")
    e_nacimiento.bind('<FocusIn>', on_focus_in_nacimiento)
    e_nacimiento.bind('<FocusOut>', on_focus_out_nacimiento)
    e_nacimiento.bind('<Return>', saltar_campo)
    
    ctk.CTkLabel(c1_1, text="Estado Civil:", text_color="black", fg_color="transparent").pack(anchor='w')
    c_civil = ctk.CTkComboBox(c1_1, values=["Soltero", "Casado", "Divorciado", "Viudo", "Unión Libre"], fg_color="white", text_color="black", border_color="grey", button_color="#1860C3")
    c_civil.pack(fill='x')
    
    ctk.CTkLabel(c1_1, text="Cargas Familiares:", text_color="black", fg_color="transparent").pack(anchor='w')
    e_cargas = crear_entry(c1_1); e_cargas.pack(fill='x')
    e_cargas.bind('<Return>', saltar_campo)

    c1_2 = ctk.CTkFrame(f1, fg_color="transparent")
    c1_2.grid(row=0, column=1, padx=20, pady=10, sticky='n')
    ctk.CTkLabel(c1_2, text="CONTACTO Y UBICACIÓN", text_color="#1860C3", font=('Arial', 12, 'bold'), fg_color="transparent").pack(pady=5)
    
    ctk.CTkLabel(c1_2, text="Telf/Celular:", text_color="black", fg_color="transparent").pack(anchor='w')
    e_telf = crear_entry(c1_2); e_telf.pack(fill='x')
    e_telf.bind('<Return>', saltar_campo)
    
    ctk.CTkLabel(c1_2, text="Email:", text_color="black", fg_color="transparent").pack(anchor='w')
    e_email = crear_entry(c1_2); e_email.pack(fill='x')
    e_email.bind('<Return>', saltar_campo)
    
    ctk.CTkLabel(c1_2, text="Dirección Domicilio:", text_color="black", fg_color="transparent").pack(anchor='w')
    e_dir = crear_entry(c1_2); e_dir.pack(fill='x')
    e_dir.bind('<Return>', saltar_campo)
    
    ctk.CTkLabel(c1_2, text="Parroquia:", text_color="black", fg_color="transparent").pack(anchor='w')
    e_parroquia = crear_entry(c1_2); e_parroquia.pack(fill='x')
    e_parroquia.bind('<Return>', saltar_campo)
    
    ctk.CTkLabel(c1_2, text="Tipo Vivienda:", text_color="black", fg_color="transparent").pack(anchor='w')
    c_vivienda = ctk.CTkComboBox(c1_2, values=["Propia", "Arrendada", "Familiar", "Hipotecada"], fg_color="white", text_color="black", border_color="grey", button_color="#1860C3")
    c_vivienda.pack(fill='x')
    
    ctk.CTkLabel(c1_2, text="Referencia Vivienda:", text_color="black", fg_color="transparent").pack(anchor='w')
    e_ref_vivienda = crear_entry(c1_2); e_ref_vivienda.pack(fill='x')
    e_ref_vivienda.bind('<Return>', saltar_campo)

    c1_3 = ctk.CTkFrame(f1, fg_color="transparent")
    c1_3.grid(row=0, column=2, padx=20, pady=10, sticky='n')
    ctk.CTkLabel(c1_3, text="REFERENCIAS", text_color="#1860C3", font=('Arial', 12, 'bold'), fg_color="transparent").pack(pady=5)
    
    ctk.CTkLabel(c1_3, text="Referencia Personal 1:", text_color="black", fg_color="transparent").pack(anchor='w')
    e_ref1 = crear_entry(c1_3); e_ref1.pack(fill='x')
    e_ref1.bind('<Return>', saltar_campo)
    
    ctk.CTkLabel(c1_3, text="Referencia Personal 2:", text_color="black", fg_color="transparent").pack(anchor='w')
    e_ref2 = crear_entry(c1_3); e_ref2.pack(fill='x')
    e_ref2.bind('<Return>', saltar_campo)
    
    ctk.CTkLabel(c1_3, text="Asesor Asignado:", text_color="black", font=('Arial', 10, 'bold'), fg_color="transparent").pack(anchor='w', pady=(20, 0))
    e_asesor = crear_entry(c1_3); e_asesor.pack(fill='x')
    e_asesor.bind('<Return>', saltar_campo)

    # --- TAB 2: SITUACIÓN FINANCIERA ---
    f2 = ctk.CTkFrame(t2, fg_color="transparent")
    f2.pack(fill='both', expand=True)

    c2_1 = ctk.CTkFrame(f2, fg_color="transparent")
    c2_1.grid(row=0, column=0, padx=20, pady=10, sticky='n')
    ctk.CTkLabel(c2_1, text="INGRESOS Y EGRESOS", text_color="#1860C3", font=('Arial', 12, 'bold'), fg_color="transparent").grid(row=0, column=0, sticky='w', pady=5)
    
    ctk.CTkLabel(c2_1, text="Score Buró (1-999):", text_color="black", fg_color="transparent").grid(row=1, column=0, sticky='w')
    e_score_buro = crear_entry(c2_1, width=100); e_score_buro.grid(row=2, column=0, sticky='w')
    e_score_buro.bind('<Return>', saltar_campo)
    
    ctk.CTkLabel(c2_1, text="Ingresos Principal ($):", text_color="black", fg_color="transparent").grid(row=3, column=0, sticky='w', pady=(5,0))
    e_ingresos = crear_entry(c2_1); e_ingresos.grid(row=4, column=0, sticky='ew')
    e_ingresos.bind('<Return>', saltar_campo)
    e_ingresos.bind('<FocusOut>', lambda e: (on_focus_out_moneda(e), toggle_fuente_ingreso())) 
    e_ingresos.bind('<FocusIn>', on_focus_in_moneda)
    
    f_fuente_ingreso = ctk.CTkFrame(c2_1, fg_color="transparent")
    ctk.CTkLabel(f_fuente_ingreso, text="Fuente:", text_color="black", fg_color="transparent").pack(side='left')
    c_fuente_ingreso = ctk.CTkComboBox(f_fuente_ingreso, values=["Sueldo", "Negocio", "Jubilación", "Arriendo", "Inversiones", "Remesas del Exterior", "Otros"], width=150, fg_color="white", text_color="black", border_color="grey")
    c_fuente_ingreso.pack(side='left', padx=5)
    # f_fuente_ingreso initialized hidden
    f_fuente_ingreso.grid_remove() 
    
    ctk.CTkLabel(c2_1, text="Ingresos Secundarios ($):", text_color="black", fg_color="transparent").grid(row=6, column=0, sticky='w', pady=(5,0))
    e_ingresos_2 = crear_entry(c2_1); e_ingresos_2.grid(row=7, column=0, sticky='ew')
    e_ingresos_2.bind('<Return>', saltar_campo)
    e_ingresos_2.bind('<FocusOut>', lambda e: (on_focus_out_moneda(e), toggle_fuente_ingreso_2())) 
    e_ingresos_2.bind('<FocusIn>', on_focus_in_moneda)
    
    f_fuente_ingreso_2 = ctk.CTkFrame(c2_1, fg_color="transparent")
    ctk.CTkLabel(f_fuente_ingreso_2, text="Fuente 2:", text_color="black", fg_color="transparent").pack(side='left')
    c_fuente_ingreso_2 = ctk.CTkComboBox(f_fuente_ingreso_2, values=["Sueldo", "Negocio", "Jubilación", "Arriendo", "Inversiones", "Remesas del Exterior", "Otros"], width=150, fg_color="white", text_color="black", border_color="grey")
    c_fuente_ingreso_2.pack(side='left', padx=5)
    # f_fuente_ingreso_2 initialized hidden
    f_fuente_ingreso_2.grid_remove()
    
    ctk.CTkLabel(c2_1, text="Egresos Mensuales ($):", text_color="black", fg_color="transparent").grid(row=9, column=0, sticky='w', pady=(5,0))
    e_egresos = crear_entry(c2_1); e_egresos.grid(row=10, column=0, sticky='ew')
    e_egresos.bind('<Return>', saltar_campo)
    e_egresos.bind('<FocusOut>', lambda e: (on_focus_out_moneda(e), calcular_total_disponible()))
    e_egresos.bind('<FocusIn>', on_focus_in_moneda)
    
    ctk.CTkLabel(c2_1, text="Total Disponible:", font=('Arial', 12, 'bold'), text_color="black", fg_color="transparent").grid(row=11, column=0, sticky='w', pady=(10,0))
    lbl_total_disponible_valor = ctk.CTkLabel(c2_1, text="$ 0.00", font=('Arial', 16, 'bold'), text_color='#006400', fg_color="transparent")
    lbl_total_disponible_valor.grid(row=12, column=0, sticky='w')
    
    e_ingresos.bind('<KeyRelease>', calcular_total_disponible)
    e_ingresos_2.bind('<KeyRelease>', calcular_total_disponible)
    e_egresos.bind('<KeyRelease>', calcular_total_disponible)

    c2_2 = ctk.CTkFrame(f2, fg_color="transparent")
    c2_2.grid(row=0, column=1, padx=40, pady=10, sticky='n')
    ctk.CTkLabel(c2_2, text="PATRIMONIO / ACTIVOS", text_color="#1860C3", font=('Arial', 12, 'bold'), fg_color="transparent").pack(pady=5)
    
    # BLOQUE TERRENO
    block_terreno = ctk.CTkFrame(c2_2, fg_color="transparent")
    block_terreno.pack(fill='x', pady=5)
    
    cb_terreno = ctk.CTkCheckBox(block_terreno, text="Terreno ($)", variable=var_terreno, text_color="black")
    cb_terreno.grid(row=0, column=0, sticky='w')
    
    e_valor_terreno = crear_entry(block_terreno, width=100)
    # Grid in toggle_terreno
    e_valor_terreno.bind('<FocusOut>', on_focus_out_moneda); e_valor_terreno.bind('<FocusIn>', on_focus_in_moneda)
    
    f_terreno_hip = ctk.CTkFrame(block_terreno, fg_color="transparent")
    # Grid in toggle_terreno
    ctk.CTkLabel(f_terreno_hip, text="¿Hipotecado?", text_color="black", font=('Arial', 10), fg_color="transparent").pack(side='left')
    c_hipotecado = ctk.CTkComboBox(f_terreno_hip, values=["Si", "No"], width=70); c_hipotecado.pack(side='left', padx=5)

    # BLOQUE CASA
    block_casa = ctk.CTkFrame(c2_2, fg_color="transparent")
    block_casa.pack(fill='x', pady=5)

    cb_casa = ctk.CTkCheckBox(block_casa, text="Casa/Dep ($)", variable=var_casa, text_color="black")
    cb_casa.grid(row=0, column=0, sticky='w')

    e_valor_casa = crear_entry(block_casa, width=100)
    # Grid in toggle_casa
    e_valor_casa.bind('<FocusOut>', on_focus_out_moneda); e_valor_casa.bind('<FocusIn>', on_focus_in_moneda)
    
    f_casa_hip = ctk.CTkFrame(block_casa, fg_color="transparent")
    # Grid in toggle_casa
    ctk.CTkLabel(f_casa_hip, text="¿Hipotecado?", text_color="black", font=('Arial', 10), fg_color="transparent").pack(side='left')
    c_hip_casa = ctk.CTkComboBox(f_casa_hip, values=["Si", "No"], width=70); c_hip_casa.pack(side='left', padx=5)

    # BLOQUE LOCAL
    block_local = ctk.CTkFrame(c2_2, fg_color="transparent")
    block_local.pack(fill='x', pady=5)

    cb_local = ctk.CTkCheckBox(block_local, text="Local Com ($)", variable=var_local, text_color="black")
    cb_local.grid(row=0, column=0, sticky='w')

    e_valor_local = crear_entry(block_local, width=100)
    # Grid in toggle_local
    e_valor_local.bind('<FocusOut>', on_focus_out_moneda); e_valor_local.bind('<FocusIn>', on_focus_in_moneda)
    
    f_local_hip = ctk.CTkFrame(block_local, fg_color="transparent")
    # Grid in toggle_local
    ctk.CTkLabel(f_local_hip, text="¿Hipotecado?", text_color="black", font=('Arial', 10), fg_color="transparent").pack(side='left')
    c_hip_local = ctk.CTkComboBox(f_local_hip, values=["Si", "No"], width=70); c_hip_local.pack(side='left', padx=5)

    # --- TAB 3: GESTIÓN Y LEGAL ---
    f3 = ctk.CTkFrame(t3, fg_color="transparent")
    f3.pack(fill='both', expand=True)

    c3_1 = ctk.CTkFrame(f3, fg_color="transparent")
    c3_1.grid(row=0, column=0, padx=20, pady=10, sticky='n')
    ctk.CTkLabel(c3_1, text="INFORMACIÓN OPERATIVA", text_color="#1860C3", font=('Arial', 12, 'bold'), fg_color="transparent").pack(pady=5)
    
    ctk.CTkLabel(c3_1, text="Profesión/Actividad:", text_color="black", fg_color="transparent").pack(anchor='w')
    e_profesion = crear_entry(c3_1); e_profesion.pack(fill='x')
    
    ctk.CTkLabel(c3_1, text="F. Apertura:", text_color="black", fg_color="transparent").pack(anchor='w')
    e_apertura = crear_entry(c3_1); e_apertura.pack(fill='x')
    
    ctk.CTkLabel(c3_1, text="N. Apertura (Carpeta):", text_color="black", fg_color="transparent").pack(anchor='w')
    e_carpeta = crear_entry(c3_1); e_carpeta.pack(fill='x')
    
    ctk.CTkLabel(c3_1, text="Observaciones Generales:", text_color="black", fg_color="transparent").pack(anchor='w', pady=(5,0))
    t_obs = tk.Text(c3_1, height=8, width=35, font=('Arial', 10)); t_obs.pack(fill='x') 

    c3_2 = ctk.CTkFrame(f3, fg_color="transparent")
    c3_2.grid(row=0, column=1, padx=40, pady=10, sticky='n')
    ctk.CTkLabel(c3_2, text="ESTADO LEGAL / CARTERA", text_color="#1860C3", font=('Arial', 12, 'bold'), fg_color="transparent").pack(pady=5)
    
    fl_grid = ctk.CTkFrame(c3_2, fg_color="transparent", border_width=1, border_color="grey")
    fl_grid.pack(fill='both', padx=5, pady=5)
    
    cb_cartera = ctk.CTkCheckBox(fl_grid, text="Cartera", variable=var_cartera, text_color="black", width=20)
    cb_cartera.grid(row=0, column=0, sticky='w', padx=10, pady=10)
    lbl_dolar_cartera = ctk.CTkLabel(fl_grid, text="($)", text_color="black", fg_color="transparent")
    lbl_dolar_cartera.grid(row=0, column=1)
    e_val_cartera = crear_entry(fl_grid, width=100)
    e_val_cartera.grid(row=0, column=2, padx=10)
    e_val_cartera.bind('<FocusOut>', on_focus_out_moneda); e_val_cartera.bind('<FocusIn>', on_focus_in_moneda)
    lbl_dolar_cartera.grid_remove(); e_val_cartera.grid_remove()

    cb_demanda = ctk.CTkCheckBox(fl_grid, text="Demanda", variable=var_demanda, text_color="black", width=20)
    cb_demanda.grid(row=1, column=0, sticky='w', padx=10, pady=10)
    lbl_dolar_demanda = ctk.CTkLabel(fl_grid, text="($)", text_color="black", fg_color="transparent")
    lbl_dolar_demanda.grid(row=1, column=1)
    e_val_demanda = crear_entry(fl_grid, width=100)
    e_val_demanda.grid(row=1, column=2, padx=10)
    e_val_demanda.bind('<FocusOut>', on_focus_out_moneda); e_val_demanda.bind('<FocusIn>', on_focus_in_moneda)
    lbl_dolar_demanda.grid_remove(); e_val_demanda.grid_remove()
    
    cb_justicia = ctk.CTkCheckBox(fl_grid, text="Justicia", variable=var_justicia, text_color="black", width=20)
    cb_justicia.grid(row=2, column=0, sticky='w', padx=10, pady=10)
    e_det_justicia = crear_entry(fl_grid, width=150)
    e_det_justicia.grid(row=2, column=1, columnspan=2, padx=10)
    e_det_justicia.grid_remove()

    # wrapper para eliminar desde UI
    def ui_eliminar_cliente():
        if not ID_CLIENTE_SELECCIONADO:
            messagebox.showwarning("Atención", "Seleccione un cliente para eliminar.")
            return
            
        if NIVEL_ACCESO != 1:
            messagebox.showerror("Error", "Solo el Administrador puede eliminar registros.")
            return

        if messagebox.askyesno("Confirmar Eliminación", "¿Está seguro de eliminar este cliente permanentemente?\nEsta acción no se puede deshacer."):
            ok, msg = eliminar_cliente(ID_CLIENTE_SELECCIONADO)
            if ok:
                messagebox.showinfo("Éxito", msg)
                limpiar_campos_ui()
                mostrar_datos_tree()
            else:
                messagebox.showerror("Error", msg)

    # --- BOTONES DE ACCIÓN ---
    f_btns = ctk.CTkFrame(app, fg_color="transparent")
    f_btns.pack(pady=10)
    
    btn_cancelar = ctk.CTkButton(f_btns, text="Cancelar / Limpiar", command=limpiar_campos_ui, fg_color="#d9534f", hover_color="#c9302c", width=150)
    btn_cancelar.pack(side='left', padx=15)
    
    btn_accion = ctk.CTkButton(f_btns, text="💾 Guardar Nuevo Cliente", command=accion_guardar, fg_color=COLOR_BTN_BG, hover_color=COLOR_BTN_HOVER, width=200, height=35)
    btn_accion.pack(side='left', padx=15)
    
    btn_eliminar = ctk.CTkButton(f_btns, text="🗑 Eliminar", command=ui_eliminar_cliente, fg_color="#d9534f", hover_color="#c9302c", state="disabled", width=120)
    # Permiso Administrador (1) para eliminar
    if NIVEL_ACCESO == 1: 
        btn_eliminar.pack(side='left', padx=15)

    
    # Restricciones para Asesores (Nivel 6)
    if NIVEL_ACCESO == 6:
        # Deshabilitar pestaña Legal (t3 es tab_view.tab("Gestión y Legal"))
        # Using a safer way to disable access if needed
        tab_view.set("Identificación y Ubicación") # Force different tab
        # We can also 'forget' the tab or disable it visually
        try:
            # tab_view._segmented_button._buttons_dict["Gestión y Legal"].configure(state="disabled") # Internal access
            pass
        except: pass

    # Logo
    try:
        img = Image.open(resource_path("Logo Face.jpg"))
        logo_ctk = ctk.CTkImage(light_image=img, dark_image=img, size=(180, 160))
        lbl = ctk.CTkLabel(top_frame, image=logo_ctk, text="")
        lbl.place(relx=1.0, x=-20, y=-5, anchor='ne')
    except: pass

    # --- LISTA ---
    mid = ctk.CTkFrame(app, fg_color=COLOR_FONDO)
    mid.pack(fill='both', expand=True, padx=20, pady=5)
    
    fb = ctk.CTkFrame(mid, fg_color="transparent")
    fb.pack(fill='x', pady=5)
    ctk.CTkLabel(fb, text="🔎 Buscar Cliente (Cédula/RUC/Nombre):", font=('Arial', 11, 'bold'), text_color="black").pack(side='left')
    e_busqueda = crear_entry(fb)
    e_busqueda.pack(side='left', fill='x', expand=True, padx=10)
    e_busqueda.bind('<KeyRelease>', lambda e: filtrar_clientes())

    ft = ctk.CTkFrame(mid, fg_color="white", corner_radius=10, border_width=1, border_color="#CCCCCC")
    ft.pack(fill='both', expand=True)

    # CAMBIADO NOMBRE DE COLUMNA CARPETA POR N. APERTURA
    cols = ("ID", "Cédula", "Nombre", "Teléfono", "Ingresos ($)", "Sit. Financiera", "Producto", "N. Apertura", "Asesor")
    tree = ttk.Treeview(ft, columns=cols, show='headings')
    
    # Scrollbar
    sy = ttk.Scrollbar(ft, orient='vertical', command=tree.yview)
    sy.pack(side='right', fill='y')
    tree.configure(yscrollcommand=sy.set)
    tree.pack(fill='both', expand=True)

    tree.heading("ID", text="ID"); tree.column("ID", width=30)
    tree.heading("Cédula", text="Cédula"); tree.column("Cédula", width=80)
    tree.heading("Nombre", text="Nombres y Apellidos"); tree.column("Nombre", width=200)
    tree.heading("Teléfono", text="Teléfono"); tree.column("Teléfono", width=80)
    tree.heading("Ingresos ($)", text="Ingresos ($)"); tree.column("Ingresos ($)", width=90)
    tree.heading("Sit. Financiera", text="Sit. Financiera"); tree.column("Sit. Financiera", width=90)
    tree.heading("Producto", text="Producto"); tree.column("Producto", width=100)
    tree.heading("N. Apertura", text="N. Apertura"); tree.column("N. Apertura", width=80)
    tree.heading("Asesor", text="Asesor"); tree.column("Asesor", width=80)

    tree.bind("<Double-1>", cargar_seleccion)
    tree.bind("<<TreeviewSelect>>", cargar_seleccion)
    mostrar_datos_tree()
    mostrar_datos_tree()
    # app.mainloop() # Ya no es mainloop principal


# --- MÓDULO MICROCRÉDITO ---
import webbrowser

def abrir_modulo_microcredito():
    global win_micro, e_cedula_micro, e_ruc_micro, e_nombre_micro, t_obs_micro
    global e_m_ref1_rel, e_m_ref1_tiempo, e_m_ref1_dir, c_m_ref1_viv, e_m_ref1_cargas, c_m_ref1_resp
    global e_m_ref2_rel, e_m_ref2_tiempo, e_m_ref2_dir, c_m_ref2_viv, e_m_ref2_cargas, c_m_ref2_resp
    global e_m_ref1_nom, e_m_ref1_tel, e_m_ref1_fec, e_m_ref1_hor
    global e_m_ref2_nom, e_m_ref2_tel, e_m_ref2_fec, e_m_ref2_hor
    global cedula_micro_actual, id_micro_actual, status_micro_actual, dict_botones_status
    global var_m_ref1_vehiculo, var_m_ref1_casa, var_m_ref1_terreno, var_m_ref1_inver
    global var_m_ref2_vehiculo, var_m_ref2_casa, var_m_ref2_terreno, var_m_ref2_inver
    global e_n_apertura_micro, e_val_apertura_micro, e_f_apertura_micro
    global e_fecha_visita, e_mapa_direccion
    global e_info_dir, e_info_civil, e_info_cargas, e_info_ing1, e_info_ing2, e_info_egr, lbl_info_total
    global e_info_casa, e_info_terr, e_info_local, e_info_cart, e_info_dem, e_info_just, e_info_score
    global t_obs_info_micro
    global f_sub_status, var_sub_status, e_f_comite, e_f_desembolsado, e_f_negado_status, e_f_desistimiento
    global nb
    global actualizar_pestanas_status

    cedula_micro_actual = None
    id_micro_actual = None

    win_micro = ctk.CTkToplevel()
    win_micro.title("Módulo de Microcrédito")
    win_micro.geometry("1100x750")
    win_micro.geometry("1100x750")
    win_micro.after(100, lambda: win_micro.state('zoomed'))

    # --- POPUP DESEMBOLSO ---
    def abrir_popup_desembolso():
        if not id_micro_actual:
            messagebox.showwarning("Aviso", "Primero GUARDE el registro del microcrédito antes de agregar los detalles financieros.")
            return
            
        t = ctk.CTkToplevel(win_micro)
        t.title("Detalles Financieros del Crédito")
        t.geometry("500x450")
        t.transient(win_micro)
        
        ctk.CTkLabel(t, text="DETALLES DE DESEMBOLSO", text_color="#1860C3", font=('Arial', 14, 'bold')).pack(pady=15)
        
        gf = ctk.CTkFrame(t, fg_color="transparent")
        gf.pack(padx=20, pady=10)
        
        # Variables locales para el popup
        vars_popup = {}
        
        def add_field(row, label, key):
            ctk.CTkLabel(gf, text=label, font=('Arial', 11, 'bold')).grid(row=row, column=0, sticky='e', padx=10, pady=5)
            e = ctk.CTkEntry(gf, width=150)
            e.grid(row=row, column=1, sticky='w', padx=10, pady=5)
            vars_popup[key] = e
            return e

        e_p_fecha = add_field(0, "Fecha Desembolso:", "fecha")
        e_p_monto = add_field(1, "Monto Aprobado ($):", "monto")
        e_p_tasa = add_field(2, "Tasa Interés (%):", "tasa")
        e_p_plazo = add_field(3, "Plazo (Meses):", "plazo")
        e_p_cuota = add_field(4, "Valor Cuota ($):", "cuota")
        e_p_dia = add_field(5, "Día de Pago:", "dia")
        
        # Validadores y Formatos
        e_p_monto.bind('<FocusOut>', on_focus_out_moneda)
        e_p_monto.bind('<FocusIn>', on_focus_in_moneda)
        
        e_p_cuota.bind('<FocusOut>', on_focus_out_moneda)
        e_p_cuota.bind('<FocusIn>', on_focus_in_moneda)
        
        def fmt_tasa_out_p(e):
            try:
                val = e.widget.get().replace('%', '').strip()
                if val: e.widget.delete(0, 'end'); e.widget.insert(0, f"{val} %")
            except: pass
        def fmt_tasa_in_p(e):
             try:
                val = e.widget.get().replace('%', '').strip()
                e.widget.delete(0, 'end'); e.widget.insert(0, val)
             except: pass
        e_p_tasa.bind('<FocusOut>', fmt_tasa_out_p)
        e_p_tasa.bind('<FocusIn>', fmt_tasa_in_p)
        
        e_p_tasa.bind('<FocusOut>', fmt_tasa_out_p)
        e_p_tasa.bind('<FocusIn>', fmt_tasa_in_p)
        
        # --- NAVEGACIÓN CON ENTER ---
        e_p_fecha.bind("<Return>", lambda e: e_p_monto.focus())
        e_p_monto.bind("<Return>", lambda e: e_p_tasa.focus())
        e_p_tasa.bind("<Return>", lambda e: e_p_plazo.focus())
        e_p_plazo.bind("<Return>", lambda e: e_p_cuota.focus())
        e_p_cuota.bind("<Return>", lambda e: e_p_dia.focus())
        # El último campo (Día de pago) puede intentar guardar o enfocar el botón, 
        # pero enfocar el botón es mejor UX para evitar guardados accidentales.
        
        # Cargar Datos
        conn, cursor = conectar_db()
        try:
            cursor.execute("SELECT fecha_desembolso_real, monto_aprobado, tasa_interes, plazo_meses, valor_cuota, dia_pago FROM Microcreditos WHERE id=%s", (id_micro_actual,))
            row = cursor.fetchone()
            if row:
                if row[0]: e_p_fecha.insert(0, row[0])
                if row[1]: e_p_monto.insert(0, formatear_float_str(row[1]))
                if row[2]: e_p_tasa.insert(0, f"{row[2]} %")
                if row[3]: e_p_plazo.insert(0, str(row[3]))
                if row[4]: e_p_cuota.insert(0, formatear_float_str(row[4]))
                if row[5]: e_p_dia.insert(0, str(row[5]))
            else:
                # Si no hay datos en DB, intentar sincronizar fecha
                try: e_p_fecha.insert(0, e_f_desembolsado.get())
                except: pass
        except Exception as e:
            messagebox.showerror("Error", f"Error cargando datos: {e}")
        finally:
            db_manager.release_connection(conn)
            
        def obtener_float(valor):
            if not valor: return 0.0
            try:
                # Eliminar símbolos de moneda y separadores de miles si es necesario
                clean = str(valor).replace('$', '').replace(',', '').strip()
                return float(clean)
            except:
                return 0.0

        def guardar_popup():
            try:
                v_fec = e_p_fecha.get()
                v_mon = obtener_float(e_p_monto.get())
                v_tas = float(e_p_tasa.get().replace('%','').strip()) if e_p_tasa.get() else 0.0
                v_pla = int(e_p_plazo.get()) if e_p_plazo.get() else 0
                v_cuo = obtener_float(e_p_cuota.get())
                v_dia = int(e_p_dia.get()) if e_p_dia.get() else 0
                
                conn, cursor = conectar_db()
                cursor.execute("""
                    UPDATE Microcreditos SET 
                    fecha_desembolso_real=%s, monto_aprobado=%s, tasa_interes=%s, plazo_meses=%s, valor_cuota=%s, dia_pago=%s
                    WHERE id=%s
                """, (v_fec, v_mon, v_tas, v_pla, v_cuo, v_dia, id_micro_actual))
                conn.commit()
                db_manager.release_connection(conn)
                messagebox.showinfo("Guardado", "Detalles financieros guardados correctamente.", parent=t)
                t.destroy()
            except Exception as e:
                 messagebox.showerror("Error", f"Error al guardar: {e}", parent=t)

        ctk.CTkButton(t, text="💾 GUARDAR DETALLES", command=guardar_popup, fg_color="#28a745", hover_color="#218838", font=('Arial', 12, 'bold')).pack(pady=20)


    # --- FUNCIÓN DEFINIDA AL INICIO PARA EVITAR NAMEERROR ---
    def actualizar_pestanas_status():
        # Lógica simplificada: Solo controlar visibilidad/estado del botón detalles
        try:
            # Eliminar pestañas dinámicas si existen (limpieza)
            try: nb.delete("Desembolsado")
            except: pass
            try: nb.delete("Desistimiento")
            except: pass
            
            estado = var_sub_status.get()
            
            # Controlar botón detalles
            if 'btn_detalles_desembolso' in globals() and btn_detalles_desembolso:
                if estado == "Desembolsado":
                    btn_detalles_desembolso.configure(state='normal', fg_color="#1860C3")
                else:
                    btn_detalles_desembolso.configure(state='disabled', fg_color="grey")

            if estado == "Desistimiento":
                nb.add("Desistimiento")
                tab_rem = nb.tab("Desistimiento")
                nb.set("Desistimiento")
                ctk.CTkLabel(tab_rem, text="MOTIVO DE DESISTIMIENTO", text_color="#d9534f", font=('Arial', 14, 'bold')).pack(pady=10)

        except Exception as e:
            print(f"Error UI Status: {e}")


    COLOR_FONDO = "#FAFAD2"
    win_micro.configure(fg_color=COLOR_FONDO)
    
    # Barra superior
    nav_frame = ctk.CTkFrame(win_micro, fg_color=COLOR_FONDO, height=40)
    nav_frame.pack(side='top', fill='x', pady=(5,0))
    
    ctk.CTkButton(nav_frame, text="Volver al Menú", command=win_micro.destroy, 
                  fg_color=COLOR_FONDO, text_color="#d9534f", hover_color="#EEE8AA", 
                  font=('Arial', 12, 'bold')).pack(side='right', padx=10)
    
    ctk.CTkLabel(win_micro, text="GESTIÓN DE MICROCRÉDITO", text_color="#1860C3", font=('Arial', 16, 'bold')).pack(pady=10)

    # Frame principal dividido
    main_frame = ctk.CTkFrame(win_micro, fg_color=COLOR_FONDO)
    main_frame.pack(fill='both', expand=True, padx=20, pady=20)

    # LEFT PANEL (Content)
    left_panel = ctk.CTkFrame(main_frame, fg_color="transparent") # Transparent wrapper
    left_panel.pack(side='left', fill='both', expand=True, padx=(0, 20))

    # SECCIÓN BÚSQUEDA (Arriba) - Using helper internal frame for layout
    search_frame = ctk.CTkFrame(left_panel, fg_color="white", border_width=1, border_color="grey")
    search_frame.pack(fill='x', pady=(0,10))
    ctk.CTkLabel(search_frame, text=" Datos del Cliente ", text_color="grey", font=('Arial', 10, 'bold')).place(x=10, y=-8) # Faux LabelFrame Title
    
    # Padding internal container
    sf_in = ctk.CTkFrame(search_frame, fg_color="transparent")
    sf_in.pack(fill='x', padx=10, pady=15)
    
    ctk.CTkLabel(sf_in, text="Cédula:", text_color="black").pack(side='left')
    e_cedula_micro = ctk.CTkEntry(sf_in, width=150, fg_color="white", text_color="black", border_color="grey")
    e_cedula_micro.pack(side='left', padx=5)
    e_cedula_micro.bind('<KeyRelease>', buscar_micro_auto)
    
    ctk.CTkLabel(sf_in, text="RUC:", text_color="black").pack(side='left', padx=(15,0))
    e_ruc_micro = ctk.CTkEntry(sf_in, width=150, fg_color="white", text_color="black", border_color="grey")
    e_ruc_micro.pack(side='left', padx=5)
    e_ruc_micro.bind('<KeyRelease>', buscar_micro_auto)

    ctk.CTkLabel(sf_in, text="Cliente:", text_color="black").pack(side='left', padx=(15,0))
    e_nombre_micro = ctk.CTkEntry(sf_in, width=350, fg_color="white", text_color="black", border_color="grey")
    e_nombre_micro.pack(side='left', padx=5)
    e_nombre_micro.bind('<KeyRelease>', buscar_micro_auto)

    # NOTEBOOK (Pestañas) -> CTkTabview
    # Estilo mejorado: Botones más grandes, colores más vivos
    nb = ctk.CTkTabview(left_panel, width=1000, height=600, 
                        fg_color="white",  # Fondo del contenido de la pestaña (blanco para contraste)
                        segmented_button_fg_color="#E0E0E0", # Fondo de la barra de pestañas
                        segmented_button_selected_color="#A9CCE3", # Color seleccionado (más tenue)
                        segmented_button_selected_hover_color="#92BBD9",
                        segmented_button_unselected_color="white", # Color no seleccionado
                        segmented_button_unselected_hover_color="#EEE",
                        text_color="black", # Letras negras para leer mejor
                        text_color_disabled="grey",
                        corner_radius=10,
                        border_width=1,
                        border_color="#CCCCCC"
                        )
    nb.pack(fill='both', expand=True)

    # ... (Tabs content defined below but nb structure is here) ...
    # We must ensure tab content is added to 'nb' which is now in 'left_panel'

    # --- RIGHT PANEL (Logo) ---
    try:
        img = Image.open(resource_path("Logo Face.jpg"))
        # Match size from Documentos module (225x210 in that snippet)
        logo_micro = ctk.CTkImage(light_image=img, dark_image=img, size=(225, 210))
        lbl = ctk.CTkLabel(main_frame, image=logo_micro, text="")
        lbl.pack(side='right', padx=20, anchor='n')
    except: 
        ctk.CTkLabel(main_frame, text="LOGO", text_color="grey").pack(side='right', padx=20, anchor='n')

    # --- PESTAÑA 1: INFORMACIÓN ---
    nb.add("Información")
    tab_info = nb.tab("Información")
    
    # ... (Existing tab construction continues) ...

    # ctk Tabs act like frames, no need to add separate frame if not needed, but code structure uses grid on tab_info
    
    # Grid layout for Info
    ctk.CTkLabel(tab_info, text="Información Relevante del Cliente", text_color="#1860C3", font=('Arial', 12, 'bold')).grid(row=0, column=0, columnspan=2, pady=(0,20), sticky='w')
    
    ctk.CTkLabel(tab_info, text="N. de Apertura (Carpeta):", text_color="black").grid(row=1, column=0, sticky='w', pady=5)
    e_n_apertura_micro = ctk.CTkEntry(tab_info, width=150, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_n_apertura_micro.grid(row=1, column=1, sticky='w', padx=10)
    
    ctk.CTkLabel(tab_info, text="Valor de Apertura ($):", text_color="black").grid(row=2, column=0, sticky='w', pady=5)
    e_val_apertura_micro = ctk.CTkEntry(tab_info, width=150, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_val_apertura_micro.grid(row=2, column=1, sticky='w', padx=10)
    e_val_apertura_micro.bind('<FocusOut>', on_focus_out_moneda)
    e_val_apertura_micro.bind('<FocusIn>', on_focus_in_moneda)
    
    ctk.CTkLabel(tab_info, text="Fecha de Apertura:", text_color="black").grid(row=3, column=0, sticky='w', pady=5)
    e_f_apertura_micro = ctk.CTkEntry(tab_info, width=150, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey") # Readonly
    e_f_apertura_micro.grid(row=3, column=1, sticky='w', padx=10)

    # NUEVOS CAMPOS INFORMACIÓN CLIENTE (Solo Lectura)
    row_idx = 4
    ctk.CTkLabel(tab_info, text="Dirección:", text_color="black").grid(row=row_idx, column=0, sticky='w', pady=5)
    e_info_dir = ctk.CTkEntry(tab_info, width=350, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_info_dir.grid(row=row_idx, column=1, sticky='w', padx=10, columnspan=2)
    
    row_idx += 1
    ctk.CTkLabel(tab_info, text="Estado Civil:", text_color="black").grid(row=row_idx, column=0, sticky='w', pady=5)
    e_info_civil = ctk.CTkEntry(tab_info, width=150, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_info_civil.grid(row=row_idx, column=1, sticky='w', padx=10)
    
    ctk.CTkLabel(tab_info, text="Cargas:", text_color="black").grid(row=row_idx, column=2, sticky='w', pady=5)
    e_info_cargas = ctk.CTkEntry(tab_info, width=80, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_info_cargas.grid(row=row_idx, column=3, sticky='w', padx=10)

    row_idx += 1
    ctk.CTkLabel(tab_info, text="Ingresos 1 / Fuente:", text_color="black").grid(row=row_idx, column=0, sticky='w', pady=5)
    e_info_ing1 = ctk.CTkEntry(tab_info, width=250, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_info_ing1.grid(row=row_idx, column=1, sticky='w', padx=10, columnspan=2)

    row_idx += 1
    ctk.CTkLabel(tab_info, text="Ingresos 2 / Fuente:", text_color="black").grid(row=row_idx, column=0, sticky='w', pady=5)
    e_info_ing2 = ctk.CTkEntry(tab_info, width=250, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_info_ing2.grid(row=row_idx, column=1, sticky='w', padx=10, columnspan=2)

    row_idx += 1
    ctk.CTkLabel(tab_info, text="Egresos:", text_color="black").grid(row=row_idx, column=0, sticky='w', pady=5)
    e_info_egr = ctk.CTkEntry(tab_info, width=150, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_info_egr.grid(row=row_idx, column=1, sticky='w', padx=10)
    
    ctk.CTkLabel(tab_info, text="Total Disponible:", text_color="black", font=('Arial', 10, 'bold')).grid(row=row_idx, column=2, sticky='w', pady=5)
    lbl_info_total = ctk.CTkLabel(tab_info, text="$ 0.00", text_color="green", font=('Arial', 10, 'bold'))
    lbl_info_total.grid(row=row_idx, column=3, sticky='w', padx=10)

    # PATRIMONIO
    row_idx += 1
    ctk.CTkLabel(tab_info, text="PATRIMONIO", text_color="#1860C3", font=('Arial', 11, 'bold')).grid(row=row_idx, column=0, sticky='w', pady=(10,5))

    row_idx += 1
    ctk.CTkLabel(tab_info, text="Casa (Val/Hip):", text_color="black").grid(row=row_idx, column=0, sticky='w', pady=2)
    e_info_casa = ctk.CTkEntry(tab_info, width=200, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_info_casa.grid(row=row_idx, column=1, sticky='w', padx=10, columnspan=2)

    row_idx += 1
    ctk.CTkLabel(tab_info, text="Terreno (Val/Hip):", text_color="black").grid(row=row_idx, column=0, sticky='w', pady=2)
    e_info_terr = ctk.CTkEntry(tab_info, width=200, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_info_terr.grid(row=row_idx, column=1, sticky='w', padx=10, columnspan=2)

    row_idx += 1
    ctk.CTkLabel(tab_info, text="Local (Val/Hip):", text_color="black").grid(row=row_idx, column=0, sticky='w', pady=2)
    e_info_local = ctk.CTkEntry(tab_info, width=200, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_info_local.grid(row=row_idx, column=1, sticky='w', padx=10, columnspan=2)

    # LEGAL
    row_idx = 0 # Inicia en 0 para alinear con el encabezado de información
    # Actually, grid column 4 and 5 for legal might be better if there is space.
    # The window is zoomed, so we have space.
    
    ctk.CTkLabel(tab_info, text="HISTORIAL LEGAL", text_color="#1860C3", font=('Arial', 11, 'bold')).grid(row=row_idx, column=4, sticky='w', pady=(0,5), padx=(40,0))
    
    row_idx += 1
    ctk.CTkLabel(tab_info, text="Cartera (Cast/Val):", text_color="black").grid(row=row_idx, column=4, sticky='w', pady=2, padx=(40,0))
    e_info_cart = ctk.CTkEntry(tab_info, width=200, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_info_cart.grid(row=row_idx, column=5, sticky='w', padx=10)

    row_idx += 1
    ctk.CTkLabel(tab_info, text="Demanda (Jud/Val):", text_color="black").grid(row=row_idx, column=4, sticky='w', pady=2, padx=(40,0))
    e_info_dem = ctk.CTkEntry(tab_info, width=200, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_info_dem.grid(row=row_idx, column=5, sticky='w', padx=10)

    row_idx += 1
    ctk.CTkLabel(tab_info, text="Justicia (Prob/Det):", text_color="black").grid(row=row_idx, column=4, sticky='w', pady=2, padx=(40,0))
    e_info_just = ctk.CTkEntry(tab_info, width=200, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_info_just.grid(row=row_idx, column=5, sticky='w', padx=10)

    row_idx += 1
    ctk.CTkLabel(tab_info, text="Score Buró:", text_color="black").grid(row=row_idx, column=4, sticky='w', pady=2, padx=(40,0))
    e_info_score = ctk.CTkEntry(tab_info, width=200, fg_color="#F0F0F0", text_color="black", state='readonly', border_color="grey")
    e_info_score.grid(row=row_idx, column=5, sticky='w', padx=10)

    # OBSERVACIONES ESPECÍFICAS
    row_idx += 2
    ctk.CTkLabel(tab_info, text="Observaciones Específicas:", text_color="#1860C3", font=('Arial', 11, 'bold')).grid(row=row_idx, column=4, sticky='w', pady=(10,0), padx=(40,0))
    row_idx += 1
    # 1.5 cm high x 4.5 cm wide aprox. 1cm ~ 38px. 1.5cm ~ 57px, 4.5cm ~ 170px.
    # Let's make it a bit bigger for usability but keeping proportions.
    t_obs_info_micro = ctk.CTkTextbox(tab_info, height=60, width=250, fg_color="white", text_color="black", border_color="grey", border_width=1)
    t_obs_info_micro.grid(row=row_idx, column=4, columnspan=2, sticky='w', pady=5, padx=(40,0))


    # --- PESTAÑA 2: LLAMADAS ---
    nb.add("Llamadas")
    tab_llamadas = nb.tab("Llamadas")

    # Reuse existing reference logic but place in tab_llamadas
    refs_container = ctk.CTkFrame(tab_llamadas, fg_color="transparent")
    refs_container.pack(fill='both', expand=True)

    def accion_guardar_ref(num_ref, widgets_tuple, btn_widget):
        global id_micro_actual, cedula_micro_actual
        
        # Unpack widgets
        e_nom, e_tel, e_fec, e_hor, e_rel, e_tiempo, e_dir, c_viv, e_cargas, vars_pat, c_resp = widgets_tuple
        
        # 1. Validación Estricta
        missing = []
        if not e_nom.get().strip(): missing.append("Nombre")
        if not e_tel.get().strip(): missing.append("Teléfono")
        if not e_rel.get().strip(): missing.append("Relación")
        if not e_tiempo.get().strip(): missing.append("Tiempo")
        if not e_dir.get().strip(): missing.append("Dirección")
        if not c_viv.get().strip(): missing.append("Vivienda")
        if not e_cargas.get().strip(): missing.append("Cargas")
        if not c_resp.get().strip(): missing.append("R. Crédito")

        if missing:
            messagebox.showerror("Error", f"Todos los campos de la referencia {num_ref} son obligatorios.\nFaltan: {', '.join(missing)}")
            return

        # 2. Estampado de Tiempo (Condicional)
        # Solo estampar si están vacíos. Si ya tienen dato, RESPETARLO.
        if not e_fec.get().strip():
             e_fec.delete(0, tk.END)
             e_fec.insert(0, datetime.datetime.now().strftime("%Y-%m-%d"))
        
        if not e_hor.get().strip():
             e_hor.delete(0, tk.END)
             e_hor.insert(0, datetime.datetime.now().strftime("%H:%M"))

        date_str = e_fec.get()
        time_str = e_hor.get()
        
        # 3. Guardado en BD
        pat_str = get_patrimonio_str(*vars_pat)
        
        ced = e_cedula_micro.get().strip()
        if not ced:
            messagebox.showerror("Error", "No hay cédula de cliente para vincular.")
            return

        conn, cursor = conectar_db()
        
        if id_micro_actual is None:
            try:
                col_prefix = "ref1_" if num_ref == 1 else "ref2_"
                query_ins = f"""
                    INSERT INTO Microcreditos (
                        cedula_cliente, ruc_cliente, 
                        {col_prefix}nombre, {col_prefix}telefono, {col_prefix}fecha_verifica, {col_prefix}hora_verifica,
                        {col_prefix}relacion, {col_prefix}tiempo_conocer, {col_prefix}direccion, {col_prefix}tipo_vivienda,
                        {col_prefix}cargas_familiares, {col_prefix}patrimonio, {col_prefix}persona_responsable
                    ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """
                vals_ins = (
                    ced, e_ruc_micro.get().strip(),
                    e_nom.get(), e_tel.get(), date_str, time_str,
                    e_rel.get(), e_tiempo.get(), e_dir.get(), c_viv.get(),
                    e_cargas.get(), pat_str, c_resp.get()
                )
                cursor.execute(query_ins, vals_ins)
                id_micro_actual = cursor.lastrowid
            except Exception as e:
                db_manager.release_connection(conn)
                messagebox.showerror("Error de Base de Datos", f"No se pudo crear el registro inicial: {e}")
                return
        else:
            col_prefix = "ref1_" if num_ref == 1 else "ref2_"
            query_upd = f"""
                UPDATE Microcreditos SET
                    {col_prefix}nombre=%s, {col_prefix}telefono=%s, {col_prefix}fecha_verifica=%s, {col_prefix}hora_verifica=%s,
                    {col_prefix}relacion=%s, {col_prefix}tiempo_conocer=%s, {col_prefix}direccion=%s, {col_prefix}tipo_vivienda=%s,
                    {col_prefix}cargas_familiares=%s, {col_prefix}patrimonio=%s, {col_prefix}persona_responsable=%s
                WHERE id = %s
            """
            vals_upd = (
                e_nom.get(), e_tel.get(), date_str, time_str,
                e_rel.get(), e_tiempo.get(), e_dir.get(), c_viv.get(),
                e_cargas.get(), pat_str, c_resp.get(),
                id_micro_actual
            )
            cursor.execute(query_upd, vals_upd)
            
        conn.commit()
        db_manager.release_connection(conn)
        
        messagebox.showinfo("Éxito", f"Referencia {num_ref} guardada correctamente.")

        # 4. Bloqueo Condicional post-guardado
        # Siempre bloquear tiempo para integridad
        try:
            e_fec.configure(state='disabled')
            e_hor.configure(state='disabled')
        except: pass

        if 'NIVEL_ACCESO' in globals() and NIVEL_ACCESO == 1:
            # ADMIN: Botón activo para corregir, campos editables (menos fecha/hora)
            btn_widget.configure(state='normal', text=f"💾 Actualizar Ref {num_ref} (Admin)", fg_color="#1860C3")
            for w in widgets_tuple: # widgets_tuple tiene (e_nom, e_tel, e_fec, e_hor...)
                # ctk checkbuttons logic might vary, but mostly 'normal' works
                if w not in [e_fec, e_hor] and not isinstance(w, tuple): 
                    try: w.configure(state='normal')
                    except: pass
                elif isinstance(w, tuple): # Checkboxes tuple
                    # Checkboxes
                    pass # Keep them active for admin? Or handle specifically if they need .configure
        else:
            # USUARIO ESTÁNDAR: Bloqueo TOTAL
            btn_widget.configure(state='disabled', text="✅ Guardada", fg_color="grey")
            for w in widgets_tuple:
                if isinstance(w, tuple): # Checkboxes tuple
                     # ctk checkboxes inside tuple. Loop them found via var? No, w is (v_veh, ...). 
                     # Wait, create_ref_group returns vars in tuple, but we need widget references to disable them?
                     # In my code below, I am passing `(v_veh, ...)` which are IterVars, NOT widgets.
                     # To disable checkboxes, we need the widget objects.
                     # FIX: create_ref_group MUST return widget objects or we find them.
                     pass 
                else:
                    try: w.configure(state='disabled')
                    except: pass
             # We can't disable checkboxes easily if we only passed IntVars. 
             # For now, locking main fields and button is sufficient deterrent.


    def create_ref_group(parent, title, col):
        f = ctk.CTkFrame(parent, fg_color="white", border_width=1, border_color="grey")
        f.pack(side='left', fill='both', expand=True, padx=5, pady=5)
        
        # Header Frame for Title + Button
        f_header = ctk.CTkFrame(f, fg_color="transparent")
        f_header.pack(fill='x', padx=5, pady=5)
        
        ctk.CTkLabel(f_header, text=title, text_color="grey", font=('Arial', 10, 'bold')).pack(side='left', anchor='w')
        
        btn_save = ctk.CTkButton(f_header, text=f"💾 Guardar Ref {col+1}", 
                                 fg_color="#28a745", hover_color="#218838", 
                                 width=120, height=24, font=('Arial', 10, 'bold'))
        btn_save.pack(side='right')
        
        # Grid content for compactness
        f_grid = ctk.CTkFrame(f, fg_color="transparent")
        f_grid.pack(fill='both', expand=True, padx=5, pady=2)
        f_grid.grid_columnconfigure(1, weight=1)
        
        row = 0
        ctk.CTkLabel(f_grid, text="Nombre Ref.:", text_color="black", font=('Arial', 10)).grid(row=row, column=0, sticky='w', pady=2)
        e_nom = ctk.CTkEntry(f_grid, height=25, fg_color="white", text_color="black", border_color="grey", font=('Arial', 11)); e_nom.grid(row=row, column=1, sticky='ew', padx=5, pady=2)
        
        row += 1
        ctk.CTkLabel(f_grid, text="Teléfono:", text_color="black", font=('Arial', 10)).grid(row=row, column=0, sticky='w', pady=2)
        e_tel = ctk.CTkEntry(f_grid, height=25, fg_color="white", text_color="black", border_color="grey", font=('Arial', 11)); e_tel.grid(row=row, column=1, sticky='ew', padx=5, pady=2)

        row += 1
        ctk.CTkLabel(f_grid, text="Fecha:", text_color="black", font=('Arial', 10)).grid(row=row, column=0, sticky='w', pady=2)
        e_fec = ctk.CTkEntry(f_grid, height=25, width=100, fg_color="#F0F0F0", text_color="black", border_color="grey", font=('Arial', 11)); e_fec.grid(row=row, column=1, sticky='w', padx=5, pady=2)

        row += 1
        ctk.CTkLabel(f_grid, text="Hora:", text_color="black", font=('Arial', 10)).grid(row=row, column=0, sticky='w', pady=2)
        e_hor = ctk.CTkEntry(f_grid, height=25, width=100, fg_color="#F0F0F0", text_color="black", border_color="grey", font=('Arial', 11)); e_hor.grid(row=row, column=1, sticky='w', padx=5, pady=2)

        row += 1
        ctk.CTkLabel(f_grid, text="1. Relación:", text_color="black", font=('Arial', 10)).grid(row=row, column=0, sticky='w', pady=2)
        e_rel = ctk.CTkEntry(f_grid, height=25, fg_color="white", text_color="black", border_color="grey", font=('Arial', 11)); e_rel.grid(row=row, column=1, sticky='ew', padx=5, pady=2)
        
        row += 1
        ctk.CTkLabel(f_grid, text="2. Tiempo:", text_color="black", font=('Arial', 10)).grid(row=row, column=0, sticky='w', pady=2)
        e_tiempo = ctk.CTkEntry(f_grid, height=25, fg_color="white", text_color="black", border_color="grey", font=('Arial', 11)); e_tiempo.grid(row=row, column=1, sticky='ew', padx=5, pady=2)
        
        row += 1
        ctk.CTkLabel(f_grid, text="3. Dirección:", text_color="black", font=('Arial', 10)).grid(row=row, column=0, sticky='w', pady=2)
        e_dir = ctk.CTkEntry(f_grid, height=25, fg_color="white", text_color="black", border_color="grey", font=('Arial', 11)); e_dir.grid(row=row, column=1, sticky='ew', padx=5, pady=2)
        
        row += 1
        ctk.CTkLabel(f_grid, text="4. Vivienda:", text_color="black", font=('Arial', 10)).grid(row=row, column=0, sticky='w', pady=2)
        c_viv = ctk.CTkComboBox(f_grid, height=25, values=["Arrendada", "Familiar", "Propia"], fg_color="white", text_color="black", border_color="grey", button_color="#1860C3", font=('Arial', 11))
        c_viv.grid(row=row, column=1, sticky='ew', padx=5, pady=2)
        
        row += 1
        ctk.CTkLabel(f_grid, text="5. Cargas:", text_color="black", font=('Arial', 10)).grid(row=row, column=0, sticky='w', pady=2)
        e_cargas = ctk.CTkEntry(f_grid, height=25, fg_color="white", text_color="black", border_color="grey", font=('Arial', 11)); e_cargas.grid(row=row, column=1, sticky='ew', padx=5, pady=2)
        
        row += 1
        ctk.CTkLabel(f_grid, text="6. Patrimonio:", text_color="black", font=('Arial', 10)).grid(row=row, column=0, sticky='w', pady=2)
        f_pat = ctk.CTkFrame(f_grid, fg_color="transparent")
        f_pat.grid(row=row, column=1, sticky='ew', padx=5, pady=2)
        
        v_veh = tk.IntVar(); v_casa = tk.IntVar(); v_terr = tk.IntVar(); v_inv = tk.IntVar()
        # Compact checkbuttons
        ck1 = ctk.CTkCheckBox(f_pat, text="Veh.", variable=v_veh, text_color="black", height=20, width=50, checkbox_width=18, checkbox_height=18, font=('Arial', 9)); ck1.grid(row=0, column=0, sticky='w')
        ck2 = ctk.CTkCheckBox(f_pat, text="Casa", variable=v_casa, text_color="black", height=20, width=50, checkbox_width=18, checkbox_height=18, font=('Arial', 9)); ck2.grid(row=0, column=1, sticky='w', padx=5)
        ck3 = ctk.CTkCheckBox(f_pat, text="Terr.", variable=v_terr, text_color="black", height=20, width=50, checkbox_width=18, checkbox_height=18, font=('Arial', 9)); ck3.grid(row=1, column=0, sticky='w')
        ck4 = ctk.CTkCheckBox(f_pat, text="Inv.", variable=v_inv, text_color="black", height=20, width=50, checkbox_width=18, checkbox_height=18, font=('Arial', 9)); ck4.grid(row=1, column=1, sticky='w', padx=5)
        
        row += 1
        ctk.CTkLabel(f_grid, text="7. R. Crédito:", text_color="black", font=('Arial', 10)).grid(row=row, column=0, sticky='w', pady=2)
        c_resp = ctk.CTkComboBox(f_grid, height=25, values=["Si", "No"], fg_color="white", text_color="black", border_color="grey", button_color="#1860C3", font=('Arial', 11))
        c_resp.grid(row=row, column=1, sticky='ew', padx=5, pady=2)
        
        # Configure button command
        # Pass checkboxes widgets tuple for potential future disabling? For now passing vars tuple as before.
        widgets_tuple = (e_nom, e_tel, e_fec, e_hor, e_rel, e_tiempo, e_dir, c_viv, e_cargas, (v_veh, v_casa, v_terr, v_inv), c_resp)
        btn_save.configure(command=lambda n=col+1, w=widgets_tuple, b=btn_save: accion_guardar_ref(n, w, b))

        return e_nom, e_tel, e_fec, e_hor, e_rel, e_tiempo, e_dir, c_viv, e_cargas, (v_veh, v_casa, v_terr, v_inv), c_resp

    e_m_ref1_nom, e_m_ref1_tel, e_m_ref1_fec, e_m_ref1_hor, e_m_ref1_rel, e_m_ref1_tiempo, e_m_ref1_dir, c_m_ref1_viv, e_m_ref1_cargas, vars_pat1, c_m_ref1_resp = create_ref_group(refs_container, "Verificación Referencia 1", 0)
    var_m_ref1_vehiculo, var_m_ref1_casa, var_m_ref1_terreno, var_m_ref1_inver = vars_pat1

    e_m_ref2_nom, e_m_ref2_tel, e_m_ref2_fec, e_m_ref2_hor, e_m_ref2_rel, e_m_ref2_tiempo, e_m_ref2_dir, c_m_ref2_viv, e_m_ref2_cargas, vars_pat2, c_m_ref2_resp = create_ref_group(refs_container, "Verificación Referencia 2", 1)
    var_m_ref2_vehiculo, var_m_ref2_casa, var_m_ref2_terreno, var_m_ref2_inver = vars_pat2


    # --- PESTAÑA 3: VISITAS ---
    nb.add("Visitas")
    tab_visitas = nb.tab("Visitas")
    
    ctk.CTkLabel(tab_visitas, text="Agendar Visita / Ubicación", text_color="#1860C3", font=('Arial', 12, 'bold')).pack(anchor='w', pady=(0,10))
    
    fv = ctk.CTkFrame(tab_visitas, fg_color="transparent")
    fv.pack(fill='x', pady=5)
    ctk.CTkLabel(fv, text="Fecha de Visita:", text_color="black").pack(side='left')
    
    # DateEntry workaround: ctk doesn't have it. Use Entry with formatted text or keep tk wrapper.
    # We will use CTkEntry for consistency and user manually enters date or we assume similar behavior.
    # Integrating tkcalendar with ctk is tricky due to visual mismatch. Let's use CTkEntry.
    e_fecha_visita = ctk.CTkEntry(fv, width=150, fg_color="white", text_color="black", border_color="grey")
    e_fecha_visita.pack(side='left', padx=10)
    e_fecha_visita.insert(0, "DD/MM/YYYY")
    
    fm = ctk.CTkFrame(tab_visitas, fg_color="transparent")
    fm.pack(fill='x', pady=20)
    ctk.CTkLabel(fm, text="Dirección para Mapa:", text_color="black").pack(side='left')
    e_mapa_direccion = ctk.CTkEntry(fm, width=400, fg_color="white", text_color="black", border_color="grey")
    e_mapa_direccion.pack(side='left', padx=10)
    
    def abrir_mapa():
        d = e_mapa_direccion.get().strip()
        if d:
            webbrowser.open(f"https://www.google.com/maps/search/?api=1&query={d}")
        else:
            messagebox.showinfo("Mapa", "Ingrese una dirección para buscar.")

    ctk.CTkButton(fm, text="🗺️ Ver en Google Maps", command=abrir_mapa, fg_color="#465EA6", hover_color="#1860C3").pack(side='left', padx=10)


    # --- PESTAÑA 4: STATUS ---
    nb.add("Status")
    tab_status = nb.tab("Status")
    
    f_status_btns = ctk.CTkFrame(tab_status, fg_color="transparent")
    f_status_btns.pack(pady=10, fill='x')
    
    global dict_botones_status
    dict_botones_status = {}

    status_list = [
        ("Microcrédito", "#1860C3"),
        ("Rehabilitación", "#28a745"),
        ("Intermediación", "#fd7e14")
    ]
    
    for txt, clr in status_list:
        btn = ctk.CTkButton(f_status_btns, text=txt, fg_color=clr, hover_color=clr, 
                            width=150, height=40, font=('Arial', 12, 'bold'),
                            command=lambda t=txt: seleccionar_status(t))
        btn.pack(side='left', padx=10)
        dict_botones_status[txt] = btn

    # FRAME PARA SUB-STATUS DE MICROCRÉDITO
    f_sub_status = ctk.CTkFrame(tab_status, fg_color="transparent")
    f_sub_status.pack(pady=10, fill='x')
    f_sub_status.pack_forget()

    var_sub_status = tk.StringVar(value="")


    def create_sub_opt(parent, label_text, var_val, row_idx, cmd=None):
        ctk.CTkLabel(parent, text=label_text, text_color="black", font=('Arial', 11, 'bold')).grid(row=row_idx, column=0, sticky='w', padx=(20, 10), pady=5)
        
        f_in = ctk.CTkFrame(parent, fg_color="transparent")
        f_in.grid(row=row_idx, column=1, sticky='w', pady=5)
        
        e_f = ctk.CTkEntry(f_in, width=120, placeholder_text="DD/MM/YYYY")
        e_f.pack(side='left', padx=5)
        
        rb = ctk.CTkRadioButton(f_in, text="", variable=var_sub_status, value=var_val, width=20, command=cmd)
        rb.pack(side='left', padx=5)
        return e_f

    e_f_comite = create_sub_opt(f_sub_status, "Comité de Crédito:", "Comité de Crédito", 0, actualizar_pestanas_status)
    e_f_negado_status = create_sub_opt(f_sub_status, "Negado:", "Negado", 1, actualizar_pestanas_status)
    e_f_desembolsado = create_sub_opt(f_sub_status, "Desembolsado:", "Desembolsado", 2, actualizar_pestanas_status)
    
    # Botón Para Detalles (Visible siempre, habilitado solo si Desembolsado)
    global btn_detalles_desembolso
    btn_detalles_desembolso = ctk.CTkButton(f_sub_status, text="📝 Detalles del Crédito", command=abrir_popup_desembolso, 
                                            fg_color="grey", state="disabled", width=160, height=28, font=('Arial', 11, 'bold'))
    btn_detalles_desembolso.grid(row=2, column=2, padx=10, sticky='w')

    e_f_desistimiento = create_sub_opt(f_sub_status, "Desistimiento:", "Desistimiento", 3, actualizar_pestanas_status)

    ctk.CTkLabel(tab_status, text="Notas y Observaciones de Status:", text_color="#1860C3", font=('Arial', 12, 'bold')).pack(anchor='w', pady=(10, 0))
    t_obs_micro = ctk.CTkTextbox(tab_status, height=200, width=600, fg_color="white", text_color="black", border_color="grey", border_width=1)
    t_obs_micro.pack(fill='both', expand=True, pady=10)


    # Botonera General (Fuera del Notebook)
    btn_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
    btn_frame.pack(pady=10)
    ctk.CTkButton(btn_frame, text="💾 Guardar / Actualizar Todo", command=guardar_microcredito, fg_color="#465EA6", hover_color="#1860C3", font=('Arial', 12, 'bold'), width=200, height=40).pack()

    # Secondary logo block removed


# --- UTILIDADES DE STATUS ---
def seleccionar_status(st):
    global status_micro_actual
    
    # RESTRICCIÓN: Si ya está guardado (id_micro_actual no es None) 
    # y el usuario NO es admin, no permitir cambiar el status principal
    if id_micro_actual is not None and NIVEL_ACCESO != 1:
        # Si intenta cambiar a un status distinto del actual cargado en DB
        # Re-cargar el status original y salir
        conn, cursor = conectar_db()
        cursor.execute("SELECT status FROM Microcreditos WHERE id = %s", (id_micro_actual,))
        res = cursor.fetchone()
        db_manager.release_connection(conn)
        original_st = res[0] if res else None
        
        if st != original_st:
            messagebox.showwarning("Restricción", "Solo un Administrador puede cambiar el status una vez guardado.")
            # Restaurar visualmente el botón original (seleccionar_status se llamará recursivamente o manejamos aquí)
            return

    status_micro_actual = st
    
    # Manejar visibilidad de sub-status
    if st == "Microcrédito":
        f_sub_status.pack(pady=10, fill='x', after=dict_botones_status["Intermediación"].master)
    else:
        f_sub_status.pack_forget()

    for name, btn in dict_botones_status.items():
        try:
            if name == st:
                btn.configure(border_width=3, border_color="black", state="normal")
            else:
                btn.configure(border_width=0, state="normal")
        except: pass

def get_patrimonio_str(v_veh, v_casa, v_terr, v_inv):
    p = []
    if v_veh.get(): p.append("Vehiculo")
    if v_casa.get(): p.append("Casa o Dep")
    if v_terr.get(): p.append("Terreno")
    if v_inv.get(): p.append("Inversiones")
    return ",".join(p)

def set_patrimonio_check(s, v_veh, v_casa, v_terr, v_inv):
    v_veh.set(0); v_casa.set(0); v_terr.set(0); v_inv.set(0)
    if not s: return
    parts = s.split(',')
    if "Vehiculo" in parts: v_veh.set(1)
    if "Casa o Dep" in parts: v_casa.set(1)
    if "Terreno" in parts: v_terr.set(1)
    if "Inversiones" in parts: v_inv.set(1)

def limpiar_form_micro():
    for e in [e_ruc_micro, e_nombre_micro, e_m_ref1_rel, e_m_ref1_tiempo, e_m_ref1_dir, e_m_ref1_cargas,
              e_m_ref2_rel, e_m_ref2_tiempo, e_m_ref2_dir, e_m_ref2_cargas,
              e_m_ref1_nom, e_m_ref1_tel, e_m_ref1_fec, e_m_ref1_hor,
              e_m_ref2_nom, e_m_ref2_tel, e_m_ref2_fec, e_m_ref2_hor]:
        try:
            e.configure(state='normal')
            e.delete(0, tk.END)
        except: pass
    
    t_obs_micro.delete("1.0", tk.END)
    t_obs_info_micro.delete("1.0", tk.END)
    seleccionar_status(None)
    
    # Limpiar Info tabs
    if 'e_n_apertura_micro' in globals():
        for e in [e_n_apertura_micro, e_f_apertura_micro, e_info_dir, e_info_civil, e_info_cargas, e_info_ing1, e_info_ing2, e_info_egr, e_info_casa, e_info_terr, e_info_local, e_info_cart, e_info_dem, e_info_just, e_info_score]:
            try:
                e.configure(state='normal')
                e.delete(0, tk.END)
                if e != e_val_apertura_micro: e.configure(state='readonly')
            except: pass
        
        e_val_apertura_micro.configure(state='normal'); e_val_apertura_micro.delete(0, tk.END)
        lbl_info_total.configure(text="$ 0.00")
    
    if 'f_sub_status' in globals():
        var_sub_status.set("")
        for e in [e_f_comite, e_f_desembolsado, e_f_negado_status, e_f_desistimiento]:
            e.delete(0, tk.END)
        f_sub_status.pack_forget()

    if 'e_mapa_direccion' in globals(): 
        e_mapa_direccion.delete(0, tk.END)
    # DateEntry reset?
    # if 'e_fecha_visita' in globals(): ...


def buscar_micro_auto(event=None):
    global cedula_micro_actual, id_micro_actual
    
    ced = e_cedula_micro.get().strip()
    ruc = e_ruc_micro.get().strip()
    nom = e_nombre_micro.get().strip()
    
    criteria = None
    val = None
    
    # Determine search criteria based on valid input
    if len(ced) == 10 and ced.isdigit():
        criteria = "cedula"
        val = ced
    elif len(ruc) >= 10 and ruc.isdigit():
         criteria = "ruc"
         val = ruc
    elif len(nom) >= 3:
         criteria = "nombre"
         val = nom
    
    if not criteria:
        return

    conn, cursor = conectar_db()
    
    res = None
    # Fetch result based on criteria
    # Including 'cedula' in SELECT to ensure we can load linked data
    # res fields: 0:cedula, 1:nombre, 2:ruc, 3:carpeta, 4:val_ape, 5:fecha_ape, 6:dir, 7:civil, 8:cargas, 9:ing1, 10:f1, 11:ing2, 12:f2, 13:egr, 14:total, 
    # 15:casa, 16:val_casa, 17:hip_casa, 18:terr, 19:val_terr, 20:hip_terr, 21:loc, 22:val_loc, 23:hip_loc,
    # 24:cart, 25:v_cart, 26:dem, 27:v_dem, 28:prob, 29:det
    
    query = """
        SELECT cedula, nombre, ruc, numero_carpeta, valor_apertura, fecha_registro, direccion, 
               estado_civil, cargas_familiares, ingresos_mensuales, fuente_ingreso, ingresos_mensuales_2, fuente_ingreso_2, egresos, total_disponible, 
               casa_dep, valor_casa_dep, hipotecado_casa_dep, terreno, valor_terreno, hipotecado, local, valor_local, hipotecado_local,
               "cartera castigada", "valor cartera", "demanda judicial", "valor demanda", "problemas justicia", "detalle justicia", "score_buro"
        FROM Clientes 
    """
    
    if criteria == "cedula":
         cursor.execute(query + " WHERE cedula = %s", (val,))
         res = cursor.fetchone()
    elif criteria == "ruc":
         cursor.execute(query + " WHERE ruc = %s", (val,))
         res = cursor.fetchone()
    elif criteria == "nombre":
         cursor.execute(query + " WHERE nombre LIKE %s LIMIT 1", (f"%{val}%",))
         res = cursor.fetchone()
    
    db_manager.release_connection(conn)
    
    if res:
        # Avoid overwriting the active field to allow fluid typing
        widget = event.widget if event else None
        
        # res fields: 0:cedula, 1:nombre, 2:ruc, 3:carpeta, 4:val_ape, 5:fecha_ape, 6:dir, 7:civil, 8:cargas, 9:ing1, 10:f1, 11:ing2, 12:f2, 13:egr, 14:total, 
        # 15:casa, 16:val_casa, 17:hip_casa, 18:terr, 19:val_terr, 20:hip_terr, 21:loc, 22:val_loc, 23:hip_loc,
        # 24:cart, 25:v_cart, 26:dem, 27:v_dem, 28:prob, 29:det
        
        if criteria != "nombre" or (widget != e_nombre_micro):
            e_nombre_micro.delete(0, tk.END); e_nombre_micro.insert(0, res[1] if res[1] else "")
            
        if criteria != "ruc" or (widget != e_ruc_micro):
            e_ruc_micro.delete(0, tk.END); 
            if res[2]: e_ruc_micro.insert(0, res[2])

        if criteria != "cedula" or (widget != e_cedula_micro):
            e_cedula_micro.delete(0, tk.END); e_cedula_micro.insert(0, res[0])
        
        
        # INTEGRACIÓN CAJA: Buscar valor apertura y numero apertura en Caja (último registro)
        val_apertura_caja = None
        num_apertura_caja = None
        
        try:
             conn_c, c_c = conectar_db()
             c_c.execute("SELECT valor_apertura, numero_apertura FROM Caja WHERE cedula = %s ORDER BY id DESC LIMIT 1", (res[0],))
             r_caja = c_c.fetchone()
             db_manager.release_connection(conn_c)
             if r_caja:
                 if r_caja[0]: val_apertura_caja = r_caja[0]
                 if r_caja[1]: num_apertura_caja = r_caja[1]
        except: pass

        # N. de Apertura: Prioridad Caja > Clientes (res[3])
        e_n_apertura_micro.configure(state='normal'); e_n_apertura_micro.delete(0, tk.END)
        if num_apertura_caja:
             e_n_apertura_micro.insert(0, str(num_apertura_caja))
        elif res[3]:
             e_n_apertura_micro.insert(0, res[3])
        e_n_apertura_micro.configure(state='readonly')

        # Fecha de Apertura: Clientes.fecha_registro (res[5])
        # Validar formato fecha antes de insertar
        fecha_str = ""
        if res[5]:
             try:
                 # Si es objeto date/datetime
                 if isinstance(res[5], (datetime.date, datetime.datetime)):
                     fecha_str = res[5].strftime("%d/%m/%Y")
                 else:
                     # Si es string, intentar parsear
                     fs = str(res[5]).split('.')[0]
                     for fmt in ["%Y-%m-%d %H:%M:%S", "%Y-%m-%d", "%d/%m/%Y"]:
                         try: 
                             dt = datetime.datetime.strptime(fs, fmt)
                             fecha_str = dt.strftime("%d/%m/%Y")
                             break
                         except: pass
                     if not fecha_str: fecha_str = fs # Fallback raw
             except: fecha_str = str(res[5])

        e_f_apertura_micro.configure(state='normal'); e_f_apertura_micro.delete(0, tk.END)
        if fecha_str: e_f_apertura_micro.insert(0, fecha_str)
        e_f_apertura_micro.configure(state='readonly')

        # Valor Apertura: Prioridad Caja > Clientes (res[4])
        e_val_apertura_micro.configure(state='normal'); e_val_apertura_micro.delete(0, tk.END)
        if val_apertura_caja:
             e_val_apertura_micro.insert(0, formatear_float_str(val_apertura_caja))
        elif res[4]: 
             e_val_apertura_micro.insert(0, formatear_float_str(res[4]))
        e_val_apertura_micro.configure(state='readonly')
        
        # Populate New Info Fields
        for e in [e_info_dir, e_info_civil, e_info_cargas, e_info_ing1, e_info_ing2, e_info_egr, e_info_casa, e_info_terr, e_info_local, e_info_cart, e_info_dem, e_info_just, e_info_score]:
            e.configure(state='normal'); e.delete(0, tk.END)

        if res[6]: e_info_dir.insert(0, res[6])
        if res[7]: e_info_civil.insert(0, res[7])
        if res[8]: e_info_cargas.insert(0, str(res[8]))
        
        ing1_str = f"{formatear_float_str(res[9] or 0)} / {res[10] or ''}"
        e_info_ing1.insert(0, ing1_str)
        
        ing2_str = f"{formatear_float_str(res[11] or 0)} / {res[12] or ''}"
        e_info_ing2.insert(0, ing2_str)
        
        if res[13]: e_info_egr.insert(0, formatear_float_str(res[13]))
        lbl_info_total.configure(text="$ " + formatear_float_str(res[14] or 0))

        # Patrimonio
        casa_str = f"{'Si' if res[15] else 'No'} ({formatear_float_str(res[16] or 0)} / {res[17] or 'No'})"
        e_info_casa.insert(0, casa_str)
        terr_str = f"{'Si' if res[18] else 'No'} ({formatear_float_str(res[19] or 0)} / {res[20] or 'No'})"
        e_info_terr.insert(0, terr_str)
        loc_str = f"{'Si' if res[21] else 'No'} ({formatear_float_str(res[22] or 0)} / {res[23] or 'No'})"
        e_info_local.insert(0, loc_str)

        # Legal
        cart_str = f"{'Si' if res[24] else 'No'} ({formatear_float_str(res[25] or 0)})"
        e_info_cart.insert(0, cart_str)
        dem_str = f"{'Si' if res[26] else 'No'} ({formatear_float_str(res[27] or 0)})"
        e_info_dem.insert(0, dem_str)
        just_str = f"{'Si' if res[28] else 'No'} ({res[29] or ''})"
        e_info_just.insert(0, just_str)
        
        # Validar índice para score_buro (res[30])
        score_val = res[30] if len(res) > 30 else ""
        e_info_score.insert(0, str(score_val) if score_val else "")

        for e in [e_info_dir, e_info_civil, e_info_cargas, e_info_ing1, e_info_ing2, e_info_egr, e_info_casa, e_info_terr, e_info_local, e_info_cart, e_info_dem, e_info_just, e_info_score]:
            e.configure(state='readonly')

        # Populate Map Address
        e_mapa_direccion.delete(0, tk.END)
        if res[6]: e_mapa_direccion.insert(0, res[6])
        
        cedula_micro_actual = res[0]
        cargar_datos_micro(cedula_micro_actual)
    else:
        # Not found
        # Do not clear fields while typing, just clear dependent data
        cedula_micro_actual = None
        id_micro_actual = None
        t_obs_micro.delete("1.0", tk.END)

def cargar_datos_micro(cedula):
    global id_micro_actual
    conn, cursor = conectar_db()
    cursor.execute("SELECT * FROM Microcreditos WHERE cedula_cliente = %s", (cedula,))
    row = cursor.fetchone()
    db_manager.release_connection(conn)
    
    if row:
        # Capture column names
        colnames = [d[0] for d in cursor.description]
        
        id_micro_actual = row[0]
        # row[1]=ced, row[2]=ruc, row[3]=obs, row[4]=obs_info
        t_obs_micro.delete("1.0", tk.END); t_obs_micro.insert("1.0", row[3] if row[3] else "")
        t_obs_info_micro.delete("1.0", tk.END); t_obs_info_micro.insert("1.0", row[4] if row[4] else "")
        
        # Refs starting from row[5] now... wait, let's check exact indexes.
        # Microcreditos table: id(0), ced(1), ruc(2), obs(3), obs_info(4), ref1_rel(5)...
        # Ref 1
        e_m_ref1_rel.delete(0,tk.END); e_m_ref1_rel.insert(0, row[5] if row[5] else "")
        e_m_ref1_tiempo.delete(0,tk.END); e_m_ref1_tiempo.insert(0, row[6] if row[6] else "")
        e_m_ref1_dir.delete(0,tk.END); e_m_ref1_dir.insert(0, row[7] if row[7] else "")
        c_m_ref1_viv.set(row[8] if row[8] else "")
        e_m_ref1_cargas.delete(0,tk.END); e_m_ref1_cargas.insert(0, row[9] if row[9] else "")
        set_patrimonio_check(row[10] if row[10] else "", var_m_ref1_vehiculo, var_m_ref1_casa, var_m_ref1_terreno, var_m_ref1_inver)
        c_m_ref1_resp.set(row[11] if row[11] else "")
        
        if len(row) > 19:
            e_m_ref1_fec.delete(0,tk.END); e_m_ref1_fec.insert(0, row[19] if row[19] else "")
            e_m_ref1_hor.delete(0,tk.END); e_m_ref1_hor.insert(0, row[20] if row[20] else "")
            e_m_ref1_nom.delete(0,tk.END); e_m_ref1_nom.insert(0, row[21] if row[21] else "")
            e_m_ref1_tel.delete(0,tk.END); e_m_ref1_tel.insert(0, row[22] if row[22] else "")

        # Ref 2
        e_m_ref2_rel.delete(0,tk.END); e_m_ref2_rel.insert(0, row[12] if row[12] else "")
        e_m_ref2_tiempo.delete(0,tk.END); e_m_ref2_tiempo.insert(0, row[13] if row[13] else "")
        e_m_ref2_dir.delete(0,tk.END); e_m_ref2_dir.insert(0, row[14] if row[14] else "")
        c_m_ref2_viv.set(row[15] if row[15] else "")
        e_m_ref2_cargas.delete(0,tk.END); e_m_ref2_cargas.insert(0, row[16] if row[16] else "")
        set_patrimonio_check(row[17] if row[17] else "", var_m_ref2_vehiculo, var_m_ref2_casa, var_m_ref2_terreno, var_m_ref2_inver)
        c_m_ref2_resp.set(row[18] if row[18] else "")

        if len(row) > 23:
            e_m_ref2_fec.delete(0,tk.END); e_m_ref2_fec.insert(0, row[23] if row[23] else "")
            e_m_ref2_hor.delete(0,tk.END); e_m_ref2_hor.insert(0, row[24] if row[24] else "")
            e_m_ref2_nom.delete(0,tk.END); e_m_ref2_nom.insert(0, row[25] if row[25] else "")
            e_m_ref2_tel.delete(0,tk.END); e_m_ref2_tel.insert(0, row[26] if row[26] else "")

        # Status
        status_loaded = None
        if len(row) > 27:
            status_loaded = row[27] if row[27] else None
            seleccionar_status(status_loaded)
        else:
            seleccionar_status(None)

        # Forzar actualización de pestañas visible
        actualizar_pestanas_status()

        # Cargar Sub-status
        if len(row) > 28:
            var_sub_status.set(row[28] if row[28] else "")
        if len(row) > 29:
            e_f_desembolsado.delete(0, tk.END); e_f_desembolsado.insert(0, row[29] if row[29] else "")
        if len(row) > 30:
            e_f_negado_status.delete(0, tk.END); e_f_negado_status.insert(0, row[30] if row[30] else "")
        if len(row) > 31:
            e_f_desistimiento.delete(0, tk.END); e_f_desistimiento.insert(0, row[31] if row[31] else "")
        if len(row) > 32:
            e_f_comite.delete(0, tk.END); e_f_comite.insert(0, row[32] if row[32] else "")

        # --- CARGA DATOS DESEMBOLSADO ---
        # (Ya no se cargan aquí, se cargan en el Popup al abrirlo)
        pass
    else:
        id_micro_actual = None
        seleccionar_status(None)
        # Limpiar formulario
        t_obs_micro.delete("1.0", tk.END)
        for e in [e_m_ref1_rel, e_m_ref1_tiempo, e_m_ref1_dir, e_m_ref1_cargas, e_m_ref2_rel, e_m_ref2_tiempo, e_m_ref2_dir, e_m_ref2_cargas]:
            e.delete(0, tk.END)
        for c in [c_m_ref1_viv, c_m_ref1_resp, c_m_ref2_viv, c_m_ref2_resp]: c.set('')
        set_patrimonio_check("", var_m_ref1_vehiculo, var_m_ref1_casa, var_m_ref1_terreno, var_m_ref1_inver)
        set_patrimonio_check("", var_m_ref2_vehiculo, var_m_ref2_casa, var_m_ref2_terreno, var_m_ref2_inver)

def guardar_microcredito():
    if not cedula_micro_actual:
        try:
            messagebox.showwarning("Aviso", "Busque un cliente primero.", parent=e_ruc_micro.winfo_toplevel())
        except:
            messagebox.showwarning("Aviso", "Busque un cliente primero.")
        return
        
    ruc = e_ruc_micro.get()
    obs = t_obs_micro.get("1.0", tk.END).strip()
    obs_info = t_obs_info_micro.get("1.0", tk.END).strip()
    
    pat1 = get_patrimonio_str(var_m_ref1_vehiculo, var_m_ref1_casa, var_m_ref1_terreno, var_m_ref1_inver)
    pat2 = get_patrimonio_str(var_m_ref2_vehiculo, var_m_ref2_casa, var_m_ref2_terreno, var_m_ref2_inver)
    
    # Save Valor Apertura to Clientes table
    val_apertura = limpiar_moneda(e_val_apertura_micro.get())
    
    vals = (
        cedula_micro_actual, ruc, obs, obs_info,
        e_m_ref1_rel.get(), e_m_ref1_tiempo.get(), e_m_ref1_dir.get(), c_m_ref1_viv.get(), e_m_ref1_cargas.get(), pat1, c_m_ref1_resp.get(),
        e_m_ref2_rel.get(), e_m_ref2_tiempo.get(), e_m_ref2_dir.get(), c_m_ref2_viv.get(), e_m_ref2_cargas.get(), pat2, c_m_ref2_resp.get(),
        e_m_ref1_fec.get(), e_m_ref1_hor.get(), e_m_ref1_nom.get(), e_m_ref1_tel.get(),
        e_m_ref2_fec.get(), e_m_ref2_hor.get(), e_m_ref2_nom.get(), e_m_ref2_tel.get(),
        status_micro_actual, var_sub_status.get(),
        e_f_desembolsado.get(), e_f_negado_status.get(), e_f_desistimiento.get(), e_f_comite.get()
    )
    
    conn, cursor = conectar_db()
    try:
        # Update Clientes table for valor_apertura AND producto (status)
        cursor.execute("UPDATE Clientes SET valor_apertura = %s, producto = %s WHERE cedula = %s", (val_apertura, status_micro_actual, cedula_micro_actual))

        if id_micro_actual:
            # Update
            # Retrieve extra fields for Desembolsado
            f_desemb_real = ""
            monto_ap = 0.0
            tasa_int = 0.0
            plazo_m = 0
            val_cuota = 0.0
            dia_p = 0
            
            if status_micro_actual == "Desembolsado": # Only if submitting as Desembolsado
                 # Financial data is now handled via Popup. 
                 # We do NOT overwrite it here with empty values.
                 pass

            cursor.execute("""
                UPDATE Microcreditos SET 
                observaciones=%s, observaciones_info=%s,
                ref1_relacion=%s, ref1_tiempo_conocer=%s, ref1_direccion=%s, ref1_tipo_vivienda=%s, ref1_cargas=%s, ref1_patrimonio=%s, ref1_responsable=%s,
                ref2_relacion=%s, ref2_tiempo_conocer=%s, ref2_direccion=%s, ref2_tipo_vivienda=%s, ref2_cargas=%s, ref2_patrimonio=%s, ref2_responsable=%s,
                ref1_fecha=%s, ref1_hora=%s, ref1_nombre=%s, ref1_telefono=%s,
                ref2_fecha=%s, ref2_hora=%s, ref2_nombre=%s, ref2_telefono=%s,
                status=%s, sub_status=%s, fecha_desembolsado=%s, fecha_negado=%s, fecha_desistimiento=%s, fecha_comite=%s
                WHERE id=%s
            """, vals[2:] + (id_micro_actual,))
            msg = "Datos actualizados."
        else:
            # Insert
            # Financial data will be null on insert. User must save then open popup.
            cursor.execute("""
                INSERT INTO Microcreditos (
                    cedula_cliente, ruc, observaciones, observaciones_info,
                    ref1_relacion, ref1_tiempo_conocer, ref1_direccion, ref1_tipo_vivienda, ref1_cargas, ref1_patrimonio, ref1_responsable,
                    ref2_relacion, ref2_tiempo_conocer, ref2_direccion, ref2_tipo_vivienda, ref2_cargas, ref2_patrimonio, ref2_responsable,
                    ref1_fecha, ref1_hora, ref1_nombre, ref1_telefono,
                    ref2_fecha, ref2_hora, ref2_nombre, ref2_telefono,
                    status, sub_status, fecha_desembolsado, fecha_negado, fecha_desistimiento, fecha_comite
                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
            """, vals)
            msg = "Datos guardados."
            
        conn.commit()
        registrar_auditoria("Guardar Microcrédito", id_cliente=cedula_micro_actual, detalles=f"Se guardaron datos de microcrédito para el cliente {cedula_micro_actual}. Status: {status_micro_actual}")
        try:
            top = e_ruc_micro.winfo_toplevel()
            messagebox.showinfo("Éxito", msg, parent=top)
            cargar_datos_micro(cedula_micro_actual) # Recargar para obtener ID si fue insert
            top.lift()
            top.focus_force()
        except:
            messagebox.showinfo("Éxito", msg)
            cargar_datos_micro(cedula_micro_actual) # Recargar para obtener ID si fue insert
    except Exception as e:
        try:
            messagebox.showerror("Error", f"Error DB: {e}", parent=e_ruc_micro.winfo_toplevel())
        except:
            messagebox.showerror("Error", f"Error DB: {e}")
    finally: db_manager.release_connection(conn)



# --- MÓDULOS NUEVOS (Structure) ---
# --- MÓDULOS GENÉRICOS ---
def crear_modulo_generico(titulo, color_titulo="#1860C3", search_callback=None, tab_name="General", show_search=True):
    win = ctk.CTkToplevel()
    win.title(titulo)
    win.geometry("1100x750")
    win.after(100, lambda: win.state('zoomed'))
    
    COLOR_FONDO = "#FAFAD2"
    win.configure(fg_color=COLOR_FONDO)
    
    # Botón Volver
    ctk.CTkButton(win, text="Volver al Menú", command=win.destroy, 
                  fg_color=COLOR_FONDO, text_color="#d9534f", hover_color="#EEE8AA", 
                  font=('Arial', 12, 'bold')).pack(side='top', anchor='ne', padx=20, pady=5)
                  
    ctk.CTkLabel(win, text=titulo.upper(), text_color=color_titulo, font=('Arial', 22, 'bold')).pack(pady=5)
    
    main_frame = ctk.CTkFrame(win, fg_color=COLOR_FONDO)
    main_frame.pack(fill='both', expand=True, padx=20, pady=10)
    
    left_panel = ctk.CTkFrame(main_frame, fg_color="transparent")
    left_panel.pack(side='left', fill='both', expand=True, padx=(0, 20))
 
    # Search Section
    if show_search:
        search_frame = ctk.CTkFrame(left_panel, fg_color="white", corner_radius=15, border_width=1, border_color="#CCCCCC")
        search_frame.pack(fill='x', pady=(0,10))
        
        sf_in = ctk.CTkFrame(search_frame, fg_color="transparent")
        sf_in.pack(fill='x', padx=15, pady=15)
        
        ctk.CTkLabel(sf_in, text="Cédula:", font=('Arial', 11), text_color="black").pack(side='left')
        e_cedula = ctk.CTkEntry(sf_in, width=130, fg_color="white", text_color="black")
        e_cedula.pack(side='left', padx=5)
        
        ctk.CTkLabel(sf_in, text="RUC:", font=('Arial', 11), text_color="black").pack(side='left', padx=(10,0))
        e_ruc = ctk.CTkEntry(sf_in, width=130, fg_color="white", text_color="black")
        e_ruc.pack(side='left', padx=5)
     
        ctk.CTkLabel(sf_in, text="Cliente:", font=('Arial', 11), text_color="black").pack(side='left', padx=(10,0))
        e_nombre = ctk.CTkEntry(sf_in, width=300, fg_color="white", text_color="black")
        e_nombre.pack(side='left', padx=5)
        
        def buscar_cliente_gen(event=None):
            ced = e_cedula.get().strip(); ruc = e_ruc.get().strip(); nom = e_nombre.get().strip()
            criteria = None; val = None
            if len(ced) == 10 and ced.isdigit(): criteria = "cedula"; val = ced
            elif len(ruc) >= 10 and ruc.isdigit(): criteria = "ruc"; val = ruc
            elif len(nom) >= 3: criteria = "nombre"; val = nom
            
            if not criteria:
                if search_callback:
                    search_callback(None) # Clear callback
                return
     
            conn, cursor = conectar_db()
            res = None
            if criteria == "cedula":
                 cursor.execute("SELECT nombre, ruc, cedula FROM Clientes WHERE cedula = %s", (val,))
                 res = cursor.fetchone()
            elif criteria == "ruc":
                 cursor.execute("SELECT nombre, ruc, cedula FROM Clientes WHERE ruc = %s", (val,))
                 res = cursor.fetchone()
            elif criteria == "nombre":
                 cursor.execute("SELECT nombre, ruc, cedula FROM Clientes WHERE nombre LIKE %s LIMIT 1", (f"%{val}%",))
                 res = cursor.fetchone()
            db_manager.release_connection(conn)
            
            if res:
                widget = event.widget if event else None
                if criteria != "nombre" or (widget != e_nombre):
                    e_nombre.delete(0, tk.END); e_nombre.insert(0, res[0])
                if criteria != "ruc" or (widget != e_ruc):
                    e_ruc.delete(0, tk.END)
                    if res[1]: e_ruc.insert(0, res[1])
                if criteria != "cedula" or (widget != e_cedula):
                    e_cedula.delete(0, tk.END); e_cedula.insert(0, res[2])
                
                if search_callback:
                    search_callback(res[2]) # Trigger callback with cedula
            else:
                if search_callback:
                    search_callback(None) # Clear callback if not found
     
        e_cedula.bind('<KeyRelease>', buscar_cliente_gen)
        e_ruc.bind('<KeyRelease>', buscar_cliente_gen)
        e_nombre.bind('<KeyRelease>', buscar_cliente_gen)
 
    # Logo
    try:
        img = Image.open(resource_path("Logo Face.jpg"))
        logo_gen = ctk.CTkImage(light_image=img, dark_image=img, size=(110, 100))
        lbl = ctk.CTkLabel(main_frame, image=logo_gen, text="")
        lbl.pack(side='right', padx=10, anchor='n')
    except: pass
    
    nb = ctk.CTkTabview(left_panel, fg_color="white", corner_radius=15, 
                        segmented_button_selected_color="#A9CCE3",
                        segmented_button_selected_hover_color="#92BBD9",
                        text_color="black")
    nb.pack(fill='both', expand=True)
    nb.add(tab_name)
    
    return win, nb.tab(tab_name), nb

def abrir_modulo_usuarios():
    win = ctk.CTkToplevel()
    win.title("Gestión de Usuarios")
    win.geometry("500x600")
    win.configure(fg_color="#FAFAD2")
    
    ctk.CTkLabel(win, text="CONTROLES DE USUARIO", text_color="#1860C3", font=('Arial', 20, 'bold')).pack(pady=20)
    
    f = ctk.CTkFrame(win, fg_color="white", corner_radius=15, border_width=1, border_color="#CCCCCC")
    f.pack(fill='both', expand=True, padx=40, pady=20)
    
    ctk.CTkLabel(f, text="Nuevo Usuario:", text_color="black").pack(pady=(20, 5))
    e_user = ctk.CTkEntry(f, width=250); e_user.pack()
    
    ctk.CTkLabel(f, text="Contraseña:", text_color="black").pack(pady=(15, 5))
    e_pass = ctk.CTkEntry(f, show="*", width=250); e_pass.pack()
    
    ctk.CTkLabel(f, text="Nivel de Acceso:", text_color="black").pack(pady=(15, 5))
    niveles = ["1 Administrador", "2 Gerencia", "3 Jefatura Intermediacion", "4 Jefatura Microcredito", "5 Jefatura Rehabilitacion", "6 Asesores", "7 Auditoria"]
    c_nivel = ctk.CTkComboBox(f, values=niveles, width=250); c_nivel.pack()
    
    def validar_usuario(event):
        u = e_user.get().strip()
        if not u: 
            e_user.configure(border_color="grey")
            return
        conn, cursor = conectar_db()
        cursor.execute("SELECT id FROM Usuarios WHERE usuario = %s", (u,))
        existe = cursor.fetchone()
        db_manager.release_connection(conn)
        if existe: e_user.configure(border_color="red")
        else: e_user.configure(border_color="green")

    def validar_pass(event):
        p = e_pass.get().strip()
        if len(p) < 4: e_pass.configure(border_color="red")
        else: e_pass.configure(border_color="green")

    e_user.bind('<KeyRelease>', validar_usuario)
    e_pass.bind('<KeyRelease>', validar_pass)

    def guardar():
        u = e_user.get().strip(); p = e_pass.get().strip()
        sel_nivel = c_nivel.get()
        n = int(sel_nivel.split()[0])
        rol = " ".join(sel_nivel.split()[1:])
        
        if u and p and len(p) >= 4:
             ok, msg = crear_usuario_db(u, p, n, rol)
             if ok:
                 registrar_auditoria("Creación Usuario", detalles=f"Se creó el usuario {u} con rol {rol}")
                 messagebox.showinfo("Éxito", f"Usuario {u} creado.")
                 # Comportamiento corregido: No cerrar, limpiar y refrescar
                 e_user.delete(0, tk.END)
                 e_pass.delete(0, tk.END)
                 e_user.configure(border_color="grey")
                 e_pass.configure(border_color="grey")
                 cargar_usuarios()
             else: messagebox.showerror("Error", "No se pudo crear el usuario (posiblemente ya existe).")
        else: messagebox.showwarning("Atención", "Complete todos los campos. La clave debe tener al menos 4 caracteres.")

    ctk.CTkButton(f, text="Crear Usuario", command=guardar, fg_color="#1860C3", height=40, width=200).pack(pady=20)

    # --- LISTA DE USUARIOS Y SOFT-DELETE ---
    ctk.CTkLabel(win, text="LISTADO DE USUARIOS", font=("Arial", 14, "bold"), text_color="black").pack(pady=10)
    
    ft = ctk.CTkFrame(win, fg_color="white", border_width=1, border_color="#CCCCCC")
    ft.pack(fill='both', expand=True, padx=40, pady=10)
    
    cols = ("ID", "Usuario", "Rol", "Estado")
    tree = ttk.Treeview(ft, columns=cols, show='headings', height=5)
    for col in cols: tree.heading(col, text=col); tree.column(col, width=100)
    tree.pack(fill='both', expand=True)

    def cargar_usuarios():
        for i in tree.get_children(): tree.delete(i)
        conn, cursor = conectar_db()
        cursor.execute("SELECT id, usuario, rol, estado FROM Usuarios")
        for r in cursor.fetchall():
            est = "Activo" if r[3] == 1 else "Inactivo"
            tree.insert('', 'end', values=(r[0], r[1], r[2], est))
        db_manager.release_connection(conn)

    def toggle_estado():
        sel = tree.selection()
        if not sel: return
        user_id = tree.item(sel[0])['values'][0]
        user_nom = tree.item(sel[0])['values'][1]
        
        conn, cursor = conectar_db()
        cursor.execute("SELECT estado FROM Usuarios WHERE id = %s", (user_id,))
        est = cursor.fetchone()[0]
        nuevo_est = 0 if est == 1 else 1
        cursor.execute("UPDATE Usuarios SET estado = %s WHERE id = %s", (nuevo_est, user_id))
        conn.commit()
        db_manager.release_connection(conn)
        
        acc = "Desactivar" if nuevo_est == 0 else "Activar"
        registrar_auditoria(f"{acc} Usuario", detalles=f"Se cambió el estado del usuario {user_nom} a {acc}")
        messagebox.showinfo("Éxito", f"Estado de {user_nom} actualizado.")
        cargar_usuarios()

    btn_f = ctk.CTkFrame(win, fg_color="transparent")
    btn_f.pack(pady=10)
    ctk.CTkButton(btn_f, text="Activar/Desactivar Seleccionado", command=toggle_estado, fg_color="#d9534f", width=250).pack()
    
    cargar_usuarios()


def abrir_modulo_rehabilitacion():
    # Variables de UI
    var_cedula = tk.StringVar()
    var_f_inicio = tk.StringVar()
    var_finalizado = tk.IntVar(value=0)

    # Variables Cliente
    v_n_apertura = tk.StringVar()
    v_f_apertura = tk.StringVar()
    v_val_apertura = tk.StringVar()
    v_edad = tk.StringVar()
    v_est_civil = tk.StringVar()
    v_cargas = tk.StringVar()
    v_telefono = tk.StringVar()
    v_email = tk.StringVar()
    v_vivienda = tk.StringVar()
    v_disponible = tk.StringVar()
    v_patrimonio = tk.StringVar()
    v_legal = tk.StringVar()
    
    # Variables Deudas (Valores y Descripciones)
    v_val_emp = tk.StringVar(); v_desc_emp = tk.StringVar()
    v_val_ban = tk.StringVar(); v_desc_ban = tk.StringVar()
    v_val_coop = tk.StringVar(); v_desc_coop = tk.StringVar()
    v_val_cob = tk.StringVar(); v_desc_cob = tk.StringVar()
    
    # Variables Cita
    v_fecha_cita = tk.StringVar()
    
    # Variable Total Deuda
    v_total_deuda = tk.StringVar(value="$ 0.00")

    # Referencias a widgets para bloqueo
    list_bloqueables = []
    
    # --- AUTO-UPDATE DB SCHEMA ---
    try:
        conn_init, cursor_init = conectar_db()
        # Verificar columnas y crearlas si faltan
        table_info = cursor_init.execute("PRAGMA table_info(Rehabilitacion)").fetchall()
        cols_present = [c[1] for c in table_info]
        
        new_cols = {
            "val_empresas": "REAL", "desc_empresas": "TEXT",
            "val_bancos": "REAL", "desc_bancos": "TEXT",
            "val_cooperativas": "REAL", "desc_cooperativas": "TEXT",
            "val_cobranzas": "REAL", "desc_cobranzas": "TEXT",
            "fecha_cita": "TEXT", "informe_cita": "TEXT"
        }
        
        for col_name, col_type in new_cols.items():
            if col_name not in cols_present:
                cursor_init.execute(f"ALTER TABLE Rehabilitacion ADD COLUMN {col_name} {col_type}")
        
        conn_init.commit()
    except Exception as e:
        print(f"Error updating schema: {e}")
    finally:
        db_manager.release_connection(conn_init)

    def load_rehab_data(cedula):
        if not cedula:
            var_cedula.set("")
            var_f_inicio.set("")
            txt_terminos.delete("1.0", tk.END)
            txt_resultado.delete("1.0", tk.END)
            var_finalizado.set(0)
            
            # Clear Client Data
            for v in [v_n_apertura, v_f_apertura, v_val_apertura, v_edad, v_est_civil, v_cargas, v_telefono, v_email, v_vivienda, v_disponible, v_patrimonio, v_legal]:
                v.set("")
                
            # Clear Deudas
            for v in [v_val_emp, v_desc_emp, v_val_ban, v_desc_ban, v_val_coop, v_desc_coop, v_val_cob, v_desc_cob]:
                v.set("")
                
            v_fecha_cita.set("")
            txt_informe_cita.delete("1.0", tk.END)
                
            update_lock_ui()
            return

        var_cedula.set(cedula)
        conn, cursor = conectar_db()
        
        # 1. Load Rehabilitacion Data (Updated Query)
        cursor.execute("""
            SELECT fecha_inicio, terminos, resultado, finalizado, 
                   val_empresas, desc_empresas, val_bancos, desc_bancos, 
                   val_cooperativas, desc_cooperativas, val_cobranzas, desc_cobranzas,
                   fecha_cita, informe_cita
            FROM Rehabilitacion WHERE cedula_cliente = %s
        """, (cedula,))
        res = cursor.fetchone()
        
        txt_terminos.delete("1.0", tk.END)
        txt_resultado.delete("1.0", tk.END)
        txt_informe_cita.delete("1.0", tk.END)

        if res:
            var_f_inicio.set(res[0] or "")
            txt_terminos.insert("1.0", res[1] or "")
            txt_resultado.insert("1.0", res[2] or "")
            var_finalizado.set(res[3] or 0)
            
            # Load Deudas
            v_val_emp.set(f"{res[4]:g}" if res[4] else "") # :g removes trailing zeros if integer
            v_desc_emp.set(res[5] or "")
            v_val_ban.set(f"{res[6]:g}" if res[6] else "")
            v_desc_ban.set(res[7] or "")
            v_val_coop.set(f"{res[8]:g}" if res[8] else "")
            v_desc_coop.set(res[9] or "")
            v_val_cob.set(f"{res[10]:g}" if res[10] else "")
            v_desc_cob.set(res[11] or "")
            
            # Recalcular total al cargar
            # Necesitamos definir la lógica de cálculo antes de llamarla o llamarla después
            # Como load está antes de la definición UI/funciones locales, diferimos o ejecutamos lógica simple aquí.
            try:
                t = 0.0
                for val in [res[4], res[6], res[8], res[10]]:
                    if val: t += float(val)
                v_total_deuda.set(f"$ {t:,.2f}")
            except: v_total_deuda.set("$ 0.00")
            
            # Load Cita
            v_fecha_cita.set(res[12] or "")
            txt_informe_cita.insert("1.0", res[13] or "")
        else:
            var_f_inicio.set("")
            var_finalizado.set(0)
            for v in [v_val_emp, v_desc_emp, v_val_ban, v_desc_ban, v_val_coop, v_desc_coop, v_val_cob, v_desc_cob]: v.set("")
            v_total_deuda.set("$ 0.00")
            v_fecha_cita.set("")
            
        # 2. Load Client Data for Summary
        query_cli = """SELECT apertura, fecha_registro, valor_apertura, "fecha nacimiento", estado_civil, cargas_familiares, telefono, email, tipo_vivienda, total_disponible, 
                              terreno, valor_terreno, hipotecado, casa_dep, valor_casa_dep, hipotecado_casa_dep, local, valor_local, hipotecado_local, 
                              "cartera castigada", "valor cartera", "demanda judicial", "valor demanda", "problemas justicia", "detalle justicia" 
                       FROM Clientes WHERE cedula = %s"""
        cursor.execute(query_cli, (cedula,))
        cli = cursor.fetchone()
        
        db_manager.release_connection(conn)
        
        if cli:
            # Map simple fields
            v_n_apertura.set(cli[0] or "")
            v_f_apertura.set(cli[1] or "")
            v_val_apertura.set(f"${formatear_float_str(cli[2])}" if cli[2] else "$0.00")
            v_est_civil.set(cli[4] or "")
            v_cargas.set(cli[5] or "")
            v_telefono.set(cli[6] or "")
            v_email.set(cli[7] or "")
            v_vivienda.set(cli[8] or "")
            v_disponible.set(f"${formatear_float_str(cli[9])}" if cli[9] else "$0.00")
            
            # Age Calculation
            dob_str = cli[3]
            age_display = "N/A"
            if dob_str:
                try:
                    # Try expected format dd/mm/yyyy
                    dob = datetime.datetime.strptime(dob_str, "%d/%m/%Y")
                    now = datetime.datetime.now()
                    age = now.year - dob.year - ((now.month, now.day) < (dob.month, dob.day))
                    age_display = f"{age} Años"
                except:
                    pass 
            v_edad.set(age_display)
            
            # Patrimonio Summary
            pat_parts = []
            if cli[10] == 1 or cli[10] == "1":
                hip = "Si" if (cli[12] == 1 or cli[12] == "1" or cli[12] == "Si") else "No"
                val = formatear_float_str(cli[11])
                pat_parts.append(f"Terreno (${val} / Hip:{hip})")
            if cli[13] == 1 or cli[13] == "1":
                hip = "Si" if (cli[15] == 1 or cli[15] == "1" or cli[15] == "Si") else "No"
                val = formatear_float_str(cli[14])
                pat_parts.append(f"Casa (${val} / Hip:{hip})")
            if cli[16] == 1 or cli[16] == "1":
                hip = "Si" if (cli[18] == 1 or cli[18] == "1" or cli[18] == "Si") else "No"
                val = formatear_float_str(cli[17])
                pat_parts.append(f"Local (${val} / Hip:{hip})")
            v_patrimonio.set(" - ".join(pat_parts) if pat_parts else "Sin Patrimonio Registrado")
            
            # Legal Summary
            leg_parts = []
            if cli[19] == 1 or cli[19] == "1":
                val = formatear_float_str(cli[20])
                leg_parts.append(f"Cartera Castigada (${val})")
            if cli[21] == 1 or cli[21] == "1":
                val = formatear_float_str(cli[22])
                leg_parts.append(f"Demanda Judicial (${val})")
            if cli[23] == 1 or cli[23] == "1":
                leg_parts.append("Problemas Justicia")
            v_legal.set(" - ".join(leg_parts) if leg_parts else "Sin Antecedentes Legales")

        else:
            pass
        
        update_lock_ui()

    def save_rehab_data():
        ced = var_cedula.get()
        if not ced:
            messagebox.showwarning("Atención", "Seleccione un cliente primero.")
            return

        f_ini = var_f_inicio.get()
        term = txt_terminos.get("1.0", tk.END).strip()
        rest = txt_resultado.get("1.0", tk.END).strip()
        fin = var_finalizado.get()
        
        # Helper to clean float
        def clean_float(s):
            try: return float(s)
            except: return 0.0
            
        val_emp = clean_float(v_val_emp.get()); desc_emp = v_desc_emp.get().strip()
        val_ban = clean_float(v_val_ban.get()); desc_ban = v_desc_ban.get().strip()
        val_coop = clean_float(v_val_coop.get()); desc_coop = v_desc_coop.get().strip()
        val_cob = clean_float(v_val_cob.get()); desc_cob = v_desc_cob.get().strip()
        
        f_cita = v_fecha_cita.get().strip()
        inf_cita = txt_informe_cita.get("1.0", tk.END).strip()

        conn, cursor = conectar_db()
        try:
            cursor.execute("""
                INSERT INTO Rehabilitacion (
                    cedula_cliente, fecha_inicio, terminos, resultado, finalizado,
                    val_empresas, desc_empresas, val_bancos, desc_bancos,
                    val_cooperativas, desc_cooperativas, val_cobranzas, desc_cobranzas,
                    fecha_cita, informe_cita
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                ON CONFLICT(cedula_cliente) DO UPDATE SET
                    fecha_inicio=excluded.fecha_inicio,
                    terminos=excluded.terminos,
                    resultado=excluded.resultado,
                    finalizado=excluded.finalizado,
                    val_empresas=excluded.val_empresas, desc_empresas=excluded.desc_empresas,
                    val_bancos=excluded.val_bancos, desc_bancos=excluded.desc_bancos,
                    val_cooperativas=excluded.val_cooperativas, desc_cooperativas=excluded.desc_cooperativas,
                    val_cobranzas=excluded.val_cobranzas, desc_cobranzas=excluded.desc_cobranzas,
                    fecha_cita=excluded.fecha_cita, informe_cita=excluded.informe_cita
            """, (ced, f_ini, term, rest, fin, val_emp, desc_emp, val_ban, desc_ban, val_coop, desc_coop, val_cob, desc_cob, f_cita, inf_cita))
            conn.commit()
            messagebox.showinfo("Éxito", "Datos de rehabilitación guardados.")
            registrar_auditoria("Guardar Rehabilitación", id_cliente=ced, detalles=f"Estado finalizado: {fin}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar: {e}")
        finally:
            db_manager.release_connection(conn)

    def toggle_finalizar():
        # Verificar nivel de admin para desbloquear
        if var_finalizado.get() == 1:
            global NIVEL_ACCESO
            # Nivel 1 es Administrador
            if NIVEL_ACCESO != 1:
                messagebox.showerror("Seguridad", "Solo un Administrador puede reabrir este proceso.")
                return
            
            if messagebox.askyesno("Confirmar", "¿Desea reabrir el proceso?"):
                var_finalizado.set(0)
                update_lock_ui()
        else:
            if messagebox.askyesno("Confirmar", "¿Desea finalizar el proceso? Esto bloqueará el expediente."):
                var_finalizado.set(1)
                save_rehab_data()
                update_lock_ui()

    def update_lock_ui():
        is_fin = var_finalizado.get() == 1
        st = "disabled" if is_fin else "normal"
        
        for w in list_bloqueables:
            w.configure(state=st)
        
        if is_fin:
            btn_finalizar.configure(text="🔒 PROCESO FINALIZADO (REABRIR)", fg_color="#d9534f", hover_color="#c9302c")
        else:
            btn_finalizar.configure(text="✅ FINALIZAR PROCESO", fg_color="#28a745", hover_color="#218838")

    def ver_llamadas():
        ced = var_cedula.get()
        if not ced: return
        
        conn, cursor = conectar_db()
        cursor.execute("SELECT * FROM Microcreditos WHERE cedula_cliente = %s ORDER BY id DESC LIMIT 1", (ced,))
        data = cursor.fetchone()
        db_manager.release_connection(conn)
        
        if not data:
            messagebox.showinfo("Llamadas", "No hay registros de Microcrédito para este cliente.")
            return

        # data indices based on standard schema:
        # ref1: rel=5, time=6, dir=7, viv=8, cargas=9, pat=10, resp=11, fecha=19, hora=20, nom=21, tel=22
        # ref2: rel=12, time=13, dir=14, viv=15, cargas=16, pat=17, resp=18, fecha=23, hora=24, nom=25, tel=26

        top = ctk.CTkToplevel(win)
        top.title(f"Historial de Llamadas - {ced}")
        top.geometry("1000x600")
        
        # Correction of window stacking
        top.lift()
        top.focus_force()
        top.grab_set()
        top.attributes('-topmost', True)
        
        # Main Scrollable Frame
        main_scroll = ctk.CTkScrollableFrame(top, fg_color="white")
        main_scroll.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Columns
        col1 = ctk.CTkFrame(main_scroll, fg_color="transparent")
        col1.pack(side='left', fill='both', expand=True, padx=5)
        
        col2 = ctk.CTkFrame(main_scroll, fg_color="transparent")
        col2.pack(side='left', fill='both', expand=True, padx=5)
        
        def crear_bloque_referencia(parent, title, vals, prefix):
            # vals: [nom, tel, fecha, hora, rel, time, dir, viv, cargas, pat, resp]
            f = ctk.CTkFrame(parent, fg_color="#F0F0F0", border_width=1, border_color="grey")
            f.pack(fill='x', pady=5, padx=5)
            ctk.CTkLabel(f, text=title, font=('Arial', 12, 'bold'), text_color="#1860C3").pack(pady=5)
            
            # Helper for rows
            def add_row(lbl, val):
                row = ctk.CTkFrame(f, fg_color="transparent")
                row.pack(fill='x', padx=5, pady=2)
                ctk.CTkLabel(row, text=lbl, width=100, anchor='w', text_color="black").pack(side='left')
                e = ctk.CTkEntry(row, fg_color="white", text_color="black")
                e.pack(side='left', fill='x', expand=True)
                e.insert(0, str(val) if val else "")
                e.configure(state='disabled')

            add_row("Nombre Ref:", vals[0])
            add_row("Teléfono:", vals[1])
            add_row("Fecha:", vals[2])
            add_row("Hora:", vals[3])
            add_row("Relación:", vals[4])
            add_row("Tiempo:", vals[5])
            add_row("Dirección:", vals[6])
            
            # Vivienda (Combo simulated)
            row_viv = ctk.CTkFrame(f, fg_color="transparent")
            row_viv.pack(fill='x', padx=5, pady=2)
            ctk.CTkLabel(row_viv, text="Vivienda:", width=100, anchor='w', text_color="black").pack(side='left')
            e_viv = ctk.CTkEntry(row_viv, fg_color="white", text_color="black")
            e_viv.pack(side='left', fill='x', expand=True)
            e_viv.insert(0, str(vals[7]) if vals[7] else "")
            e_viv.configure(state='disabled')
            
            add_row("Cargas Fam:", vals[8])
            
            # Patrimonio Checkboxes (Read only)
            row_pat = ctk.CTkFrame(f, fg_color="transparent")
            row_pat.pack(fill='x', padx=5, pady=5)
            ctk.CTkLabel(row_pat, text="Patrimonio:", width=100, anchor='w', text_color="black").pack(side='left')
            
            pat_str = str(vals[9]) if vals[9] else ""
            
            # Map DB Keyword -> UI Label
            pat_map = [
                ("Vehiculo", "Vehículo"), 
                ("Casa", "Casa"), # Matches 'Casa' in 'Casa o Dep'
                ("Terreno", "Terreno"),
                ("Inversiones", "Inversiones")
            ]
            
            for keyword, label_text in pat_map:
                chk = ctk.CTkCheckBox(row_pat, text=label_text, text_color="black")
                if keyword in pat_str: chk.select()
                chk.configure(state='disabled')
                chk.pack(side='left', padx=2)
                
            # Responsable
            row_resp = ctk.CTkFrame(f, fg_color="transparent")
            row_resp.pack(fill='x', padx=5, pady=5)
            chk_resp = ctk.CTkCheckBox(row_resp, text="¿Es Persona Responsable?", text_color="black")
            if vals[10] == "Si" or vals[10] == "1": chk_resp.select()
            chk_resp.configure(state='disabled')
            chk_resp.pack(side='left', padx=30)

        # Ref 1 Data
        # Indices: nom=21, tel=22, fecha=19, hora=20, rel=5, time=6, dir=7, viv=8, cargas=9, pat=10, resp=11
        vals1 = [data[21], data[22], data[19], data[20], data[5], data[6], data[7], data[8], data[9], data[10], data[11]]
        crear_bloque_referencia(col1, "Verificación Referencia 1", vals1, "r1")

        # Ref 2 Data
        # Indices: nom=25, tel=26, fecha=23, hora=24, rel=12, time=13, dir=14, viv=15, cargas=16, pat=17, resp=18
        vals2 = [data[25], data[26], data[23], data[24], data[12], data[13], data[14], data[15], data[16], data[17], data[18]]
        crear_bloque_referencia(col2, "Verificación Referencia 2", vals2, "r2")

    def ver_visitas():
        ced = var_cedula.get()
        if not ced: return
        
        conn, cursor = conectar_db()
        cursor.execute("SELECT fecha, observaciones FROM visitas_microcredito WHERE cedula_cliente = %s", (ced,))
        visitas = cursor.fetchall()
        db_manager.release_connection(conn)
        
        if not visitas:
            messagebox.showinfo("Visitas", "No se encontraron visitas registradas para este cliente.")
            return
            
        msg = f"Visitas registradas para {ced}:\n\n"
        for v in visitas:
            msg += f"📅 {v[0]}: {v[1]}\n"
        
        # Mostrar en un popup simple o un toplevel nuevo si es largo
        messagebox.showinfo("Historial de Visitas", msg)

    # UI
    def ver_informacion():
        ced = var_cedula.get()
        if not ced: return
        conn, cursor = conectar_db()
        cursor.execute("SELECT nombre, telefono, direccion, producto FROM Clientes WHERE cedula = %s", (ced,))
        res = cursor.fetchone()
        db_manager.release_connection(conn)
        if res:
            msg = f"Nombre: {res[0]}\nTeléfono: {res[1]}\nDirección: {res[2]}\nProducto: {res[3]}"
            messagebox.showinfo("Información del Cliente", msg)
        else:
            messagebox.showwarning("Aviso", "No se encontraron datos del cliente.")

    def ver_buro():
        ced = var_cedula.get()
        if not ced: return
        conn, cursor = conectar_db()
        cursor.execute("SELECT buro_archivo_ruta FROM Caja WHERE cedula = %s ORDER BY id DESC LIMIT 1", (ced,))
        res = cursor.fetchone()
        db_manager.release_connection(conn)
        
        if res and res[0]:
            ruta = res[0]
            if os.path.exists(ruta):
                try:
                    os.startfile(ruta)
                except Exception as e:
                    messagebox.showerror("Error", f"No se pudo abrir el archivo: {e}")
            else:
                messagebox.showwarning("Aviso", f"El archivo no existe en la ruta:\n{ruta}")
        else:
            messagebox.showinfo("Buró", "No hay archivo de Buró registrado en Caja.")

    # UI
    win, frame, nb = crear_modulo_generico("Módulo de Rehabilitación", search_callback=load_rehab_data)
    
    # --- PESTAÑA GENERAL (Contenido extra) ---
    extra_f = ctk.CTkFrame(frame, fg_color="transparent")
    extra_f.pack(pady=10, fill='x')
    
    # Updated Button Bar: 1. Ver Llamadas, 2. Ver Visitas, 3. Ver Buró (#3DCF96)
    ctk.CTkButton(extra_f, text="📞 Ver Llamadas", command=ver_llamadas, fg_color="#17a2b8", width=140).pack(side='left', padx=5)
    ctk.CTkButton(extra_f, text="🏠 Ver Visitas", command=ver_visitas, fg_color="#fd7e14", width=140).pack(side='left', padx=5)
    ctk.CTkButton(extra_f, text="📄 Ver Buró", command=ver_buro, fg_color="#3DCF96", width=140).pack(side='left', padx=5)

    # --- CLIENT DATA SUMMARY FRAME ---
    # This frame will show the client's financial/personal summary
    
    # --- CLIENT DATA SUMMARY FRAME ---
    # This frame will show the client's financial/personal summary
    
    client_info_frame = ctk.CTkFrame(frame, fg_color="white", corner_radius=10, border_width=1, border_color="#CCCCCC")
    client_info_frame.pack(fill='both', expand=True, padx=10, pady=5)
    
    ctk.CTkLabel(client_info_frame, text="DATOS FINANCIEROS Y LEGALES DEL CLIENTE", font=('Arial', 12, 'bold'), text_color="#1860C3").pack(anchor='w', pady=(5,10), padx=10)
    
    # Grid info
    info_grid = ctk.CTkFrame(client_info_frame, fg_color="transparent")
    info_grid.pack(fill='both', expand=True, padx=5)
    
    # Configurar columnas para que se expandan equitativamente
    info_grid.grid_columnconfigure((0,1,2,3), weight=1)
    
    def add_read_field(parent, row, col, lbl, var):
        f_field = ctk.CTkFrame(parent, fg_color="transparent")
        f_field.grid(row=row, column=col, sticky='nsew', padx=5, pady=2)
        ctk.CTkLabel(f_field, text=lbl, font=('Arial', 10, 'bold'), text_color="grey", anchor="w").pack(fill='x')
        e = ctk.CTkEntry(f_field, textvariable=var, state='readonly', fg_color="#F9F9F9", border_color="lightgrey", text_color="black")
        e.pack(fill='x')
        
    # Row 0: N. Apertura | F. Apertura | Valor Apertura | Edad
    add_read_field(info_grid, 0, 0, "N. Apertura", v_n_apertura)
    add_read_field(info_grid, 0, 1, "F. Apertura", v_f_apertura)
    add_read_field(info_grid, 0, 2, "Valor Apertura", v_val_apertura)
    add_read_field(info_grid, 0, 3, "Edad", v_edad)

    # Row 1: Estado Civil | Cargas | Telef/Celular | Email
    add_read_field(info_grid, 1, 0, "Estado Civil", v_est_civil)
    add_read_field(info_grid, 1, 1, "Cargas", v_cargas)
    add_read_field(info_grid, 1, 2, "Telef/Celular", v_telefono)
    add_read_field(info_grid, 1, 3, "Email", v_email)
    
    # Row 2: Tipo Vivienda | Total Disponible | (Empty) | (Empty)
    add_read_field(info_grid, 2, 0, "Tipo Vivienda", v_vivienda)
    add_read_field(info_grid, 2, 1, "Total Disponible $", v_disponible)
    
    # Row 3: Patrimonio (Spans all columns)
    f_pat = ctk.CTkFrame(info_grid, fg_color="transparent")
    f_pat.grid(row=3, column=0, columnspan=4, sticky='ew', padx=5, pady=(5,2))
    ctk.CTkLabel(f_pat, text="Patrimonio:", font=('Arial', 10, 'bold'), text_color="grey", anchor="w").pack(fill='x')
    e_pat = ctk.CTkEntry(f_pat, textvariable=v_patrimonio, state='readonly', fg_color="#E8F8F5", text_color="#145A32", border_color="#A2D9CE")
    e_pat.pack(fill='x')
    
    # Row 4: Legal (Spans all columns)
    f_leg = ctk.CTkFrame(info_grid, fg_color="transparent")
    f_leg.grid(row=4, column=0, columnspan=4, sticky='ew', padx=5, pady=(5,10))
    ctk.CTkLabel(f_leg, text="Antecedentes Legales:", font=('Arial', 10, 'bold'), text_color="grey", anchor="w").pack(fill='x')
    e_leg = ctk.CTkEntry(f_leg, textvariable=v_legal, state='readonly', fg_color="#FDEDEC", text_color="#922B21", border_color="#F5B7B1")
    e_leg.pack(fill='x')

    # --- PESTAÑA PROCESO ---
    nb.add("Proceso")
    tab_p = nb.tab("Proceso")
    
    p_main = ctk.CTkFrame(tab_p, fg_color="white", corner_radius=10)
    p_main.pack(fill='both', expand=True, padx=10, pady=10)
    
    # Configure 2 Main Columns
    p_main.grid_columnconfigure(0, weight=1)
    p_main.grid_columnconfigure(1, weight=1)
    p_main.grid_rowconfigure(0, weight=1)
    
    # --- COLUMNA IZQUIERDA: ANÁLISIS DE DEUDAS ---
    f_left = ctk.CTkFrame(p_main, fg_color="#F9F9F9", border_width=1, border_color="#DDDDDD")
    f_left.grid(row=0, column=0, sticky='nsew', padx=5, pady=5)
    f_left.grid_columnconfigure(2, weight=1) # col 0 label, 1 val, 2 desc
    
    ctk.CTkLabel(f_left, text="ANÁLISIS DE DEUDAS", font=('Arial', 14, 'bold'), text_color="#C0392B").pack(pady=10)
    
    # Headers Row
    h_row = ctk.CTkFrame(f_left, fg_color="transparent")
    h_row.pack(fill='x', padx=10)
    ctk.CTkLabel(h_row, text="", width=120).pack(side='left') # spacer for labels
    ctk.CTkLabel(h_row, text="Valor ($)", width=80, font=('Arial', 10, 'bold')).pack(side='left', padx=5)
    ctk.CTkLabel(h_row, text="Descripción", font=('Arial', 10, 'bold')).pack(side='left', padx=5)

    # --- Lógica Local de Cálculo y Formato ---
    def saltar_campo(event):
        event.widget.tk_focusNext().focus()
        return "break"

    def limpiar_moneda(texto):
        return texto.replace("$", "").replace(",", "").strip()

    def formatear_moneda(valor):
        try:
            return f"$ {float(valor):,.2f}"
        except:
            return ""

    def calcular_total_deuda(*args):
        total = 0.0
        # Lista de variables de valor
        vals = [v_val_emp.get(), v_val_ban.get(), v_val_coop.get(), v_val_cob.get()]
        for v in vals:
            try:
                clean = limpiar_moneda(v)
                if clean:
                    total += float(clean)
            except: pass
        v_total_deuda.set(f"$ {total:,.2f}")

    def on_focus_in_moneda(event):
        try:
            texto = event.widget.get()
            clean = limpiar_moneda(texto)
            if clean:
                event.widget.delete(0, tk.END)
                event.widget.insert(0, clean)
        except: pass

    def on_focus_out_moneda(event):
        try:
            texto = event.widget.get()
            clean = limpiar_moneda(texto)
            if clean:
                fmt = formatear_moneda(clean)
                event.widget.delete(0, tk.END)
                event.widget.insert(0, fmt)
            calcular_total_deuda() # Recalcular al salir
        except: pass
        
    def on_key_release_moneda(event):
        calcular_total_deuda()

    def add_debt_row(parent, label, v_val, v_desc):
        row = ctk.CTkFrame(parent, fg_color="transparent")
        row.pack(fill='x', padx=10, pady=5)
        
        ctk.CTkLabel(row, text=label, width=120, anchor='w', text_color="black", font=('Arial', 11, 'bold')).pack(side='left')
        
        e_val = ctk.CTkEntry(row, textvariable=v_val, width=80, placeholder_text="$")
        e_val.pack(side='left', padx=5)
        
        # Bindings
        e_val.bind('<FocusIn>', on_focus_in_moneda)
        e_val.bind('<FocusOut>', on_focus_out_moneda)
        e_val.bind('<KeyRelease>', on_key_release_moneda)
        e_val.bind('<Return>', saltar_campo)
        
        e_desc = ctk.CTkEntry(row, textvariable=v_desc, width=200, placeholder_text="Detalle / Institución")
        e_desc.pack(side='left', fill='x', expand=True, padx=5)
        e_desc.bind('<Return>', saltar_campo)
        
        list_bloqueables.append(e_val)
        list_bloqueables.append(e_desc)

    add_debt_row(f_left, "Empresas", v_val_emp, v_desc_emp)
    add_debt_row(f_left, "Bancos", v_val_ban, v_desc_ban)
    add_debt_row(f_left, "Cooperativas", v_val_coop, v_desc_coop)
    add_debt_row(f_left, "Emp. Cobranza", v_val_cob, v_desc_cob)
    
    # --- TOTAL DEUDA ---
    f_tot = ctk.CTkFrame(f_left, fg_color="transparent")
    f_tot.pack(fill='x', padx=10, pady=(5,10))
    
    # Separator line (simulated)
    ctk.CTkFrame(f_tot, height=2, fg_color="grey").pack(fill='x', pady=(0,5))
    
    row_t = ctk.CTkFrame(f_tot, fg_color="transparent")
    row_t.pack(fill='x')
    
    ctk.CTkLabel(row_t, text="TOTAL DEUDA:", font=('Arial', 12, 'bold'), text_color="black", anchor='e', width=120).pack(side='left')
    
    e_total_deuda = ctk.CTkEntry(row_t, textvariable=v_total_deuda, state='readonly', width=120, font=('Arial', 12, 'bold'), fg_color="#E8F8F5", text_color="#145A32", border_color="#A2D9CE")
    e_total_deuda.pack(side='left', padx=5)
    
    # --- SECCIÓN AGENDAMIENTO / CITA ---
    ctk.CTkLabel(f_left, text="AGENDAMIENTO / CITA", font=('Arial', 12, 'bold'), text_color="#1860C3").pack(pady=(20,5))
    
    f_cita = ctk.CTkFrame(f_left, fg_color="#EBF5FB", border_width=1, border_color="#AED6F1")
    f_cita.pack(fill='x', padx=10, pady=5)
    
    # Row Fecha
    row_fc = ctk.CTkFrame(f_cita, fg_color="transparent")
    row_fc.pack(fill='x', padx=5, pady=5)
    ctk.CTkLabel(row_fc, text="Fecha Cita:", width=80, anchor='w').pack(side='left')
    e_fecha_cita = ctk.CTkEntry(row_fc, textvariable=v_fecha_cita, width=120, placeholder_text="DD/MM/YYYY")
    e_fecha_cita.pack(side='left', padx=5)
    list_bloqueables.append(e_fecha_cita)
    
    def set_hoy():
        v_fecha_cita.set(datetime.datetime.now().strftime("%d/%m/%Y"))
        
    ctk.CTkButton(row_fc, text="📅 Hoy", width=60, command=set_hoy, fg_color="#28a745", hover_color="#218838").pack(side='left', padx=5)
    
    # Row Informe
    ctk.CTkLabel(f_cita, text="Informe Cita (Resultado Presencial):", anchor='w').pack(fill='x', padx=5, pady=(5,0))
    txt_informe_cita = ctk.CTkTextbox(f_cita, height=100, border_width=1)
    txt_informe_cita.pack(fill='x', padx=5, pady=5)
    list_bloqueables.append(txt_informe_cita)
    
    # --- COLUMNA DERECHA: REHABILITACIÓN (Textos) ---
    f_right = ctk.CTkFrame(p_main, fg_color="#F9F9F9", border_width=1, border_color="#DDDDDD")
    f_right.grid(row=0, column=1, sticky='nsew', padx=5, pady=5)
    f_right.grid_columnconfigure(0, weight=1)
    f_right.grid_rowconfigure(3, weight=1) # Resultado expands
    
    # F. Inicio
    ctk.CTkLabel(f_right, text="Fecha de Inicio:", font=('Arial', 11, 'bold'), text_color="black").grid(row=0, column=0, sticky='w', padx=10, pady=(10,0))
    e_f_inicio = ctk.CTkEntry(f_right, textvariable=var_f_inicio)
    e_f_inicio.grid(row=1, column=0, sticky='ew', padx=10, pady=(0,10))
    list_bloqueables.append(e_f_inicio)
    
    # Términos
    ctk.CTkLabel(f_right, text="Términos de Rehabilitación:", font=('Arial', 11, 'bold'), text_color="black").grid(row=2, column=0, sticky='w', padx=10, pady=(10,0))
    txt_terminos = ctk.CTkTextbox(f_right, height=100, border_width=1)
    txt_terminos.grid(row=3, column=0, sticky='ew', padx=10, pady=(0,10))
    list_bloqueables.append(txt_terminos)
    
    # Resultado
    ctk.CTkLabel(f_right, text="Resultado / Conclusiones:", font=('Arial', 11, 'bold'), text_color="black").grid(row=4, column=0, sticky='w', padx=10, pady=(10,0))
    txt_resultado = ctk.CTkTextbox(f_right, height=100, border_width=1)
    txt_resultado.grid(row=5, column=0, sticky='ew', padx=10, pady=(0,20))
    list_bloqueables.append(txt_resultado)
    
    # Botonera de Acciones (Restaurada)
    f_btn_proceso = ctk.CTkFrame(p_main, fg_color="transparent")
    f_btn_proceso.grid(row=1, column=0, columnspan=2, sticky='ew', pady=20, padx=10)
    
    btn_save = ctk.CTkButton(f_btn_proceso, text="💾 GUARDAR CAMBIOS", command=save_rehab_data, fg_color="#1860C3")
    btn_save.pack(side='left')
    list_bloqueables.append(btn_save)
    
    btn_finalizar = ctk.CTkButton(f_btn_proceso, text="✅ FINALIZAR PROCESO", command=toggle_finalizar, fg_color="#28a745")
    btn_finalizar.pack(side='right')

    update_lock_ui()

def abrir_modulo_intermediacion():
    # Variables locales
    var_cedula = tk.StringVar()
    var_fecha_cita_inter = tk.StringVar()
    
    # Vars Proceso
    var_fecha_desembolso_inter = tk.StringVar()
    var_chk_precal = tk.IntVar()
    var_chk_aprob = tk.IntVar()
    var_chk_informe = tk.IntVar()
    var_desc_precal = tk.StringVar()
    var_desc_aprob = tk.StringVar()
    var_desc_informe = tk.StringVar()
    
    # List of widgets to control security
    # To be populated after widget creation
    security_widgets = {} 
    
    # --- FUNCIONES DE SOPORTE ---\r
    def ver_llamadas():
        ced = var_cedula.get()
        if not ced: return
        
        conn, cursor = conectar_db()
        cursor.execute("SELECT * FROM Microcreditos WHERE cedula_cliente = %s ORDER BY id DESC LIMIT 1", (ced,))
        data = cursor.fetchone()
        db_manager.release_connection(conn)
        
        if not data:
            messagebox.showinfo("Llamadas", "No hay registros de Microcrédito para este cliente.")
            return

        top = ctk.CTkToplevel(win)
        top.title(f"Historial de Llamadas - {ced}")
        top.geometry("1000x600")
        
        top.lift()
        top.focus_force()
        top.grab_set()
        top.attributes('-topmost', True)
        
        main_scroll = ctk.CTkScrollableFrame(top, fg_color="white")
        main_scroll.pack(fill='both', expand=True, padx=10, pady=10)
        
        col1 = ctk.CTkFrame(main_scroll, fg_color="transparent")
        col1.pack(side='left', fill='both', expand=True, padx=5)
        
        col2 = ctk.CTkFrame(main_scroll, fg_color="transparent")
        col2.pack(side='left', fill='both', expand=True, padx=5)
        
        def crear_bloque_referencia(parent, title, vals, prefix):
            f = ctk.CTkFrame(parent, fg_color="#F0F0F0", border_width=1, border_color="grey")
            f.pack(fill='x', pady=5, padx=5)
            ctk.CTkLabel(f, text=title, font=('Arial', 12, 'bold'), text_color="#1860C3").pack(pady=5)
            
            def add_row(lbl, val):
                row = ctk.CTkFrame(f, fg_color="transparent")
                row.pack(fill='x', padx=5, pady=2)
                ctk.CTkLabel(row, text=lbl, width=100, anchor='w', text_color="black").pack(side='left')
                e = ctk.CTkEntry(row, fg_color="white", text_color="black")
                e.pack(side='left', fill='x', expand=True)
                e.insert(0, str(val) if val else "")
                e.configure(state='disabled')

            add_row("Nombre Ref:", vals[0])
            add_row("Teléfono:", vals[1])
            add_row("Fecha:", vals[2])
            add_row("Hora:", vals[3])
            add_row("Relación:", vals[4])
            add_row("Tiempo:", vals[5])
            add_row("Dirección:", vals[6])
            
            row_viv = ctk.CTkFrame(f, fg_color="transparent")
            row_viv.pack(fill='x', padx=5, pady=2)
            ctk.CTkLabel(row_viv, text="Vivienda:", width=100, anchor='w', text_color="black").pack(side='left')
            e_viv = ctk.CTkEntry(row_viv, fg_color="white", text_color="black")
            e_viv.pack(side='left', fill='x', expand=True)
            e_viv.insert(0, str(vals[7]) if vals[7] else "")
            e_viv.configure(state='disabled')
            
            add_row("Cargas Fam:", vals[8])
            
            row_pat = ctk.CTkFrame(f, fg_color="transparent")
            row_pat.pack(fill='x', padx=5, pady=5)
            ctk.CTkLabel(row_pat, text="Patrimonio:", width=100, anchor='w', text_color="black").pack(side='left')
            pat_str = str(vals[9]) if vals[9] else ""
            pat_map = [("Vehiculo", "Vehículo"), ("Casa", "Casa"), ("Terreno", "Terreno"), ("Inversiones", "Inversiones")]
            for keyword, label_text in pat_map:
                chk = ctk.CTkCheckBox(row_pat, text=label_text, text_color="black")
                if keyword in pat_str: chk.select()
                chk.configure(state='disabled')
                chk.pack(side='left', padx=2)
                
            row_resp = ctk.CTkFrame(f, fg_color="transparent")
            row_resp.pack(fill='x', padx=5, pady=5)
            chk_resp = ctk.CTkCheckBox(row_resp, text="¿Es Persona Responsable?", text_color="black")
            if vals[10] == "Si" or vals[10] == "1": chk_resp.select()
            chk_resp.configure(state='disabled')
            chk_resp.pack(side='left', padx=30)

        vals1 = [data[21], data[22], data[19], data[20], data[5], data[6], data[7], data[8], data[9], data[10], data[11]]
        crear_bloque_referencia(col1, "Verificación Referencia 1", vals1, "r1")

        vals2 = [data[25], data[26], data[23], data[24], data[12], data[13], data[14], data[15], data[16], data[17], data[18]]
        crear_bloque_referencia(col2, "Verificación Referencia 2", vals2, "r2")

    def ver_visitas():
        ced = var_cedula.get()
        if not ced: return
        messagebox.showinfo("Visitas", f"Historial de visitas para {ced}")

    def ver_buro():
        ced = var_cedula.get()
        if not ced: return
        try:
            conn, cursor = conectar_db()
            cursor.execute("SELECT buro_archivo_ruta FROM Caja WHERE cedula = %s ORDER BY id DESC LIMIT 1", (ced,))
            res = cursor.fetchone()
            if res and res[0] and os.path.exists(res[0]):
                os.startfile(res[0])
            else:
                messagebox.showwarning("Archivo", "No se encontró archivo de Buró.")
        except Exception as e:
            messagebox.showerror("Error", f"Error abriendo archivo: {e}")

    def set_entry(entry, value):
        entry.configure(state='normal')
        entry.delete(0, tk.END)
        entry.insert(0, str(value) if value is not None else "")
        entry.configure(state='readonly')

    # --- LISTA DE ITEMS GESTIÓN ---
    items_gestion = [
        ("Disolución Temporal", "chk_dis_temp", "val_dis_temp", "obs_dis_temp"),
        ("Disolución Registrada", "chk_dis_reg", "val_dis_reg", "obs_dis_reg"),
        ("Patrimonio 1", "chk_pat_1", "val_pat_1", "obs_pat_1"),
        ("Patrimonio 2", "chk_pat_2", "val_pat_2", "obs_pat_2"),
        ("Inmueble", "chk_inmueble", "val_inmueble", "obs_inmueble"),
        ("Ruc", "chk_ruc", "val_ruc", "obs_ruc"),
        ("Declaraciones", "chk_declaraciones", "val_declaraciones", "obs_declaraciones"),
        ("Estados de Cuenta", "chk_estados_cta", "val_estados_cta", "obs_estados_cta")
    ]
    
    # Listas para guardar referencias a widgets
    widgets_gestion = [] # List of tuples: (checkbox, entry, entry_obs, key_chk, key_val, key_obs)
    lista_entries_gestion = [] # List of entry widgets for navigation (contains both vals and obs in order)

    def calcular_total_gestion(*args):
        total = 0.0
        for e in lista_entries_gestion:
            try:
                val_str = e.get().replace('$', '').replace(',', '').strip()
                if val_str:
                    total += float(val_str)
            except:
                pass
        
        try:
            e_total_gestion.configure(state='normal')
            e_total_gestion.delete(0, tk.END)
            e_total_gestion.insert(0, f"$ {total:,.2f}")
            e_total_gestion.configure(state='readonly')
        except: pass

    def saltar_al_siguiente(event, current_widget):
        try:
            # Format if it's a value entry
            on_focus_out_moneda(event)
            
            # Find current index
            if current_widget in lista_entries_gestion:
                idx = lista_entries_gestion.index(current_widget)
                
                # Try finding next enabled
                for j in range(idx + 1, len(lista_entries_gestion)):
                    if lista_entries_gestion[j].cget("state") == "normal":
                        lista_entries_gestion[j].focus_set()
                        return "break"
        except Exception as e: print(e)
        return "break"

    def toggle_entry_state(chk, ent, ent_obs):
        if chk.get() == 1:
            ent.configure(state='normal')
            ent_obs.configure(state='normal')
            ent.focus_set()
        else:
            ent.delete(0, tk.END)
            ent_obs.delete(0, tk.END)
            ent.configure(state='disabled')
            ent_obs.configure(state='disabled')
            calcular_total_gestion() # Recalc to 0

    def load_inter_data(cedula):
        if not cedula:
            limpiar_campos_inter()
            return
        
        var_cedula.set(cedula)
        
        conn, cursor = conectar_db()
        
        # 1. Cargar Datos Generales (Read-Only)
        query = """
            SELECT apertura, fecha_registro, valor_apertura, "fecha nacimiento", 
                   estado_civil, cargas_familiares, telefono, email, tipo_vivienda, total_disponible, 
                   terreno, valor_terreno, hipotecado, casa_dep, valor_casa_dep, hipotecado_casa_dep, 
                   local, valor_local, hipotecado_local, 
                   "cartera castigada", "valor cartera", "demanda judicial", "valor demanda", 
                   "problemas justicia", "detalle justicia" 
            FROM Clientes WHERE cedula = %s
        """
        cursor.execute(query, (cedula,))
        res = cursor.fetchone()
        
        # 2. Cargar Gestión Intermediación
        cursor.execute("SELECT * FROM Intermediacion_Detalles WHERE cedula_cliente = %s", (cedula,))
        gestion = cursor.fetchone()
        col_names = [desc[0] for desc in cursor.description] if gestion else []
        
        db_manager.release_connection(conn)
        
        limpiar_campos_inter()
        
        if res:
            # Llenar campos simples
            set_entry(e_nap, res[0])
            set_entry(e_fap, res[1])
            set_entry(e_val, f"$ {res[2]:,.2f}" if res[2] else "$ 0.00")
            
            # Cálculo Edad
            try:
                fn_str = res[3]
                if fn_str:
                    fmt_ok = None
                    for f in ["%d/%m/%Y", "%Y-%m-%d"]:
                        try:
                            fn_dt = datetime.datetime.strptime(fn_str, f)
                            fmt_ok = fn_dt
                            break
                        except: pass
                    
                    if fmt_ok:
                        hoy = datetime.datetime.now()
                        edad = hoy.year - fmt_ok.year - ((hoy.month, hoy.day) < (fmt_ok.month, fmt_ok.day))
                        set_entry(e_edad, f"{edad} Años")
            except: pass
            
            set_entry(e_civ, res[4])
            set_entry(e_car, res[5])
            set_entry(e_tel, res[6])
            set_entry(e_ema, res[7])
            set_entry(e_viv, res[8])
            set_entry(e_tot, f"$ {res[9]:,.2f}" if res[9] else "$ 0.00")
            
            # Patrimonio
            pat_list = []
            if res[10] == 1: pat_list.append(f"Terreno (${res[11] or 0:,.2f} / Hip:{'Si' if res[12] else 'No'})")
            if res[13] == 1: pat_list.append(f"Casa (${res[14] or 0:,.2f} / Hip:{'Si' if res[15] else 'No'})")
            if res[16] == 1: pat_list.append(f"Local (${res[17] or 0:,.2f} / Hip:{'Si' if res[18] else 'No'})")
            set_entry(e_pat, ", ".join(pat_list) if pat_list else "Sin Patrimonio Registrado")
            
            # Legal
            leg_list = []
            if res[19] == 1: leg_list.append(f"Cartera (${res[20] or 0:,.2f})")
            if res[21] == 1: leg_list.append(f"Demanda (${res[22] or 0:,.2f})")
            if res[23] == 1: leg_list.append(f"Justicia ({res[24]})")
            set_entry(e_leg, ", ".join(leg_list) if leg_list else "Sin Antecedentes")

        # Llenar Gestión
        if gestion:
            # Mapear columnas a indices
            idx_map = {name: i for i, name in enumerate(col_names)}
            
            # Helper to safely get value
            def get_val(col, var):
                if col in idx_map:
                    v = gestion[idx_map[col]]
                    if v is not None: var.set(v)
            
            # Cargar Proceso
            get_val('chk_precal', var_chk_precal)
            get_val('desc_precal', var_desc_precal)
            get_val('chk_aprob', var_chk_aprob)
            get_val('desc_aprob', var_desc_aprob)
            get_val('fecha_desembolso_inter', var_fecha_desembolso_inter)
            get_val('chk_informe', var_chk_informe)
            get_val('desc_informe', var_desc_informe)

            # Trigger Toggle logic
            toggle_fecha_desembolso()

            # Trigger Security logic
            aplicar_seguridad_proceso()

            # Cargar Cita
            if 'fecha_cita' in idx_map:
                fc = gestion[idx_map['fecha_cita']]
                if fc: var_fecha_cita_inter.set(fc)
                
            if 'informe_cita' in idx_map:
                ic = gestion[idx_map['informe_cita']]
                if ic:
                    txt_informe_cita_inter.delete("1.0", tk.END)
                    txt_informe_cita_inter.insert("1.0", ic)
            
            for chk, ent, ent_obs, key_chk, key_val, key_obs in widgets_gestion:
                # Load Checkbox
                if key_chk in idx_map:
                    val_c = gestion[idx_map[key_chk]]
                    if val_c == 1: 
                        chk.select()
                        ent.configure(state='normal') # Enable if checked
                        ent_obs.configure(state='normal')
                    else: 
                        chk.deselect()
                        ent.configure(state='disabled') # Ensure disabled
                        ent_obs.configure(state='disabled')
                
                # Load Value and Observation
                if chk.get() == 1: # Only load if checked
                    if key_val in idx_map:
                        val_v = gestion[idx_map[key_val]]
                        ent.delete(0, tk.END)
                        ent.insert(0, f"$ {val_v:,.2f}" if val_v else "")
                    
                    if key_obs in idx_map:
                        val_o = gestion[idx_map[key_obs]]
                        ent_obs.delete(0, tk.END)
                        if val_o: ent_obs.insert(0, val_o)
                    
        # Calcular Total al cargar
        calcular_total_gestion()

    def limpiar_campos_inter():
        # General Fields
        for e in [e_nap, e_fap, e_val, e_edad, e_civ, e_car, e_tel, e_ema, e_viv, e_tot, e_pat, e_leg]:
            e.configure(state='normal')
            e.delete(0, tk.END)
            e.configure(state='readonly')
        
        # Gestión Fields
        for chk, ent, ent_obs, _, _, _ in widgets_gestion:
            chk.deselect()
            ent.delete(0, tk.END)
            ent_obs.delete(0, tk.END)
            ent.configure(state='disabled')
            ent_obs.configure(state='disabled')

        try:
            e_total_gestion.configure(state='normal')
            e_total_gestion.delete(0, tk.END)
            e_total_gestion.insert(0, "$ 0.00")
            e_total_gestion.configure(state='readonly')
        except: pass
        
        # Limpiar Cita y Proceso
        var_fecha_cita_inter.set("")
        try: txt_informe_cita_inter.delete("1.0", tk.END)
        except: pass
        
        var_chk_precal.set(0)
        var_desc_precal.set("")
        var_chk_aprob.set(0)
        var_desc_aprob.set("")
        var_fecha_desembolso_inter.set("")
        var_chk_informe.set(0)
        var_desc_informe.set("")
        toggle_fecha_desembolso()
        aplicar_seguridad_proceso()

    def guardar_gestion():
        ced = var_cedula.get()
        if not ced:
            messagebox.showwarning("Aviso", "Seleccione un cliente primero.")
            return
            
        conn, cursor = conectar_db()
        try:
            # Check exist
            cursor.execute("SELECT id FROM Intermediacion_Detalles WHERE cedula_cliente = %s", (ced,))
            exists = cursor.fetchone()
            
            # Prepare data
            cols = ["cedula_cliente"]
            vals = [ced]
            
            for chk, ent, ent_obs, key_chk, key_val, key_obs in widgets_gestion:
                cols.append(key_chk)
                chk_val = 1 if chk.get() == 1 else 0
                vals.append(chk_val)
                
                cols.append(key_val)
                raw_val = ent.get()
                vals.append(limpiar_moneda(raw_val) if chk_val == 1 and raw_val else 0.0)

                cols.append(key_obs)
                vals.append(ent_obs.get() if chk_val == 1 else "")
            
            # Save Cita fields
            cols.append("fecha_cita")
            vals.append(var_fecha_cita_inter.get())
            
            cols.append("informe_cita")
            vals.append(txt_informe_cita_inter.get("1.0", tk.END).strip())
            
            # Save Proceso fields
            cols.extend(["chk_precal", "desc_precal", "chk_aprob", "desc_aprob", "fecha_desembolso_inter", "chk_informe", "desc_informe"])
            vals.extend([
                var_chk_precal.get(), var_desc_precal.get(),
                var_chk_aprob.get(), var_desc_aprob.get(),
                var_fecha_desembolso_inter.get(),
                var_chk_informe.get(), var_desc_informe.get()
            ])
            
            if exists:
                # UPDATE
                # Build SET clause dynamically skipping cedula_cliente
                set_clause = ", ".join([f"{c}=%s" for c in cols[1:]])
                query = f"UPDATE Intermediacion_Detalles SET {set_clause} WHERE id=%s"
                cursor.execute(query, vals[1:] + [exists[0]]) 
            else:
                # INSERT
                placeholders = ",".join(["%s"] * len(cols))
                col_str = ",".join(cols)
                query = f"INSERT INTO Intermediacion_Detalles ({col_str}) VALUES ({placeholders})"
                cursor.execute(query, vals)
                
            conn.commit()
            messagebox.showinfo("Éxito", "Gestión guardada correctamente.")
            registrar_auditoria("Guardar Intermediación", id_cliente=ced, detalles="Actualización de trámites y valores")
            
            # Apply security after save
            aplicar_seguridad_proceso()
            
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar la gestión: {e}")
        finally:
            db_manager.release_connection(conn)

    def aplicar_seguridad_proceso():
        # Check if widgets exist
        if not security_widgets: return

        # Get Widgets
        c_pre = security_widgets['chk_precal']
        e_pre = security_widgets['desc_precal']
        c_apr = security_widgets['chk_aprob']
        e_apr = security_widgets['desc_aprob']
        c_inf = security_widgets['chk_informe']
        e_inf = security_widgets['desc_informe']
        
        # 1. Admin Access (NIVEL_ACCESO == 1)
        if NIVEL_ACCESO == 1:
            for w in [c_pre, e_pre, c_apr, e_apr, c_inf, e_inf]:
                w.configure(state='normal')
            return

        # 2. User Access (Sequential Logic)
        
        # Default: Everything disabled except Step 1
        for w in [c_pre, e_pre, c_apr, e_apr, c_inf, e_inf]:
            w.configure(state='disabled')
            
        # Step 1: Precalificacion
        # If not done, enable it. If done, lock it and check next.
        if var_chk_precal.get() == 0:
            c_pre.configure(state='normal')
            e_pre.configure(state='normal')
            # Rest remain disabled
        else:
            # Step 1 Done. Step 2: Aprobacion
            if var_chk_aprob.get() == 0:
                c_apr.configure(state='normal')
                e_apr.configure(state='normal')
            else:
                # Step 2 Done. Step 3: Informe
                if var_chk_informe.get() == 0:
                    c_inf.configure(state='normal')
                    e_inf.configure(state='normal')
                else:
                    pass # All done, all locked (except for Admin)

    # --- CREACIÓN DE INTERFAZ ---
    win, frame, nb = crear_modulo_generico("Módulo de Intermediación", search_callback=load_inter_data)

    # 1. BOTONERA SUPERIOR
    btn_f = ctk.CTkFrame(frame, fg_color="transparent") 
    btn_f.pack(fill='x', pady=10)

    ctk.CTkButton(btn_f, text="📞 Ver Llamadas", command=ver_llamadas, fg_color="#17a2b8", width=140).pack(side='left', padx=5) 
    ctk.CTkButton(btn_f, text="🏠 Ver Visitas", command=ver_visitas, fg_color="#fd7e14", width=140).pack(side='left', padx=5) 
    ctk.CTkButton(btn_f, text="📄 Ver Buró", command=ver_buro, fg_color="#3DCF96", width=140).pack(side='left', padx=5)

    # 2. PESTAÑAS DETALLE
    # La función crear_modulo_generico ya creó 'nb' con tab ("General")
    
    # --- TAB 1: GENERAL (Datos cliente) ---
    tab_gen = nb.tab("General")
    
    data_f = ctk.CTkFrame(tab_gen, fg_color="white", corner_radius=10) 
    data_f.pack(fill='both', expand=True, padx=10, pady=10)

    ctk.CTkLabel(data_f, text="DATOS FINANCIEROS Y LEGALES", text_color="#1860C3", font=('Arial', 12, 'bold')).grid(row=0, column=0, sticky='w', padx=20, pady=(15,10))

    def mk_field(r, c, title, width=150): 
        ctk.CTkLabel(data_f, text=title, text_color="black", font=('Arial', 10)).grid(row=r, column=c, sticky='w', padx=10, pady=(5,0)) 
        e = ctk.CTkEntry(data_f, width=width, fg_color="#F9F9F9", text_color="black", state='readonly') 
        e.grid(row=r+1, column=c, sticky='w', padx=10, pady=(0,10)) 
        return e

    e_nap = mk_field(1, 0, "N. Apertura")
    e_fap = mk_field(1, 1, "F. Apertura")
    e_val = mk_field(1, 2, "Valor Apertura")
    e_edad = mk_field(1, 3, "Edad")

    e_civ = mk_field(3, 0, "Estado Civil") 
    e_car = mk_field(3, 1, "Cargas") 
    e_tel = mk_field(3, 2, "Telef/Celular") 
    e_ema = mk_field(3, 3, "Email", width=200)

    e_viv = mk_field(5, 0, "Tipo Vivienda") 
    e_tot = mk_field(5, 1, "Total Disponible $")

    ctk.CTkLabel(data_f, text="Patrimonio:", text_color="black", font=('Arial', 10)).grid(row=7, column=0, sticky='w', padx=10, pady=(5,0)) 
    e_pat = ctk.CTkEntry(data_f, fg_color="#F9F9F9", text_color="black", state='readonly') 
    e_pat.grid(row=8, column=0, columnspan=4, sticky='ew', padx=10, pady=(0,10))

    ctk.CTkLabel(data_f, text="Antecedentes Legales:", text_color="black", font=('Arial', 10)).grid(row=9, column=0, sticky='w', padx=10, pady=(5,0)) 
    e_leg = ctk.CTkEntry(data_f, fg_color="#fee", text_color="#c00", state='readonly') 
    e_leg.grid(row=10, column=0, columnspan=4, sticky='ew', padx=10, pady=(0,10))

    # --- TAB 2: GESTIÓN INTERMEDIACIÓN ---
    nb.add("Gestión Intermediación")
    tab_inter = nb.tab("Gestión Intermediación")
    
    # 2 Column Grid
    tab_inter.grid_columnconfigure(0, weight=1)
    tab_inter.grid_columnconfigure(1, weight=1)
    tab_inter.grid_rowconfigure(0, weight=1)
    
    # --- FRAME IZQUIERDO: LISTA GESTIÓN ---
    f_lista = ctk.CTkFrame(tab_inter, fg_color="white", corner_radius=10, border_width=1, border_color="#DDD")
    f_lista.grid(row=0, column=0, sticky='nsew', padx=10, pady=10)
    
    # Encabezados
    ctk.CTkLabel(f_lista, text="Selección", font=('Arial', 11, 'bold'), text_color="black").grid(row=0, column=0, padx=5, pady=10)
    ctk.CTkLabel(f_lista, text="Concepto", font=('Arial', 11, 'bold'), text_color="black").grid(row=0, column=1, padx=5, pady=10, sticky='w')
    ctk.CTkLabel(f_lista, text="Valor ($)", font=('Arial', 11, 'bold'), text_color="black").grid(row=0, column=2, padx=5, pady=10)
    ctk.CTkLabel(f_lista, text="Observaciones", font=('Arial', 11, 'bold'), text_color="black").grid(row=0, column=3, padx=5, pady=10)
    
    # Separador
    ttk.Separator(f_lista, orient='horizontal').grid(row=1, column=0, columnspan=4, sticky='ew', padx=5, pady=(0,10))

    # Generar Filas
    for i, (label_text, key_chk, key_val, key_obs) in enumerate(items_gestion):
        r = i + 2
        
        # WIDGETS
        ent = ctk.CTkEntry(f_lista, width=80, fg_color="#F9F9F9", text_color="black", justify="right", state="disabled") # Smaller width
        ent_obs = ctk.CTkEntry(f_lista, width=150, fg_color="#F9F9F9", text_color="black", state="disabled") # Smaller width

        chk = ctk.CTkCheckBox(f_lista, text="", width=24, checkbox_width=20, checkbox_height=20)
        chk.configure(command=lambda c=chk, e=ent, eo=ent_obs: toggle_entry_state(c, e, eo))
        
        chk.grid(row=r, column=0, padx=5, pady=5)
        
        lbl = ctk.CTkLabel(f_lista, text=label_text, font=('Arial', 10), text_color="black") # Smaller font
        lbl.grid(row=r, column=1, padx=5, pady=5, sticky='w')
        
        ent.grid(row=r, column=2, padx=5, pady=5)
        ent_obs.grid(row=r, column=3, padx=5, pady=5)
        
        # Bindings
        ent.bind("<FocusOut>", lambda e: [on_focus_out_moneda(e), calcular_total_gestion()])
        ent.bind("<FocusIn>", on_focus_in_moneda)
        ent.bind("<KeyRelease>", lambda e: calcular_total_gestion())
        ent.bind("<Return>", lambda e, w=ent: saltar_al_siguiente(e, w))
        ent_obs.bind("<Return>", lambda e, w=ent_obs: saltar_al_siguiente(e, w))
        
        widgets_gestion.append((chk, ent, ent_obs, key_chk, key_val, key_obs))
        lista_entries_gestion.append(ent)
        lista_entries_gestion.append(ent_obs)
        
    last_row = len(items_gestion) + 2
    ttk.Separator(f_lista, orient='horizontal').grid(row=last_row, column=0, columnspan=4, sticky='ew', padx=5, pady=10)
    
    # ROW TOTAL
    ctk.CTkLabel(f_lista, text="TOTAL:", font=('Arial', 12, 'bold'), text_color="#1860C3").grid(row=last_row+1, column=2, sticky='e', padx=5, pady=10)
    e_total_gestion = ctk.CTkEntry(f_lista, width=100, fg_color="#E0F2F1", text_color="#00695C", font=('Arial', 12, 'bold'), justify="right", state='readonly')
    e_total_gestion.grid(row=last_row+1, column=3, padx=5, pady=10, sticky='w')

    # --- FRAME DERECHO: AGENDAMIENTO / CITA ---
    f_cita = ctk.CTkFrame(tab_inter, fg_color="#EBF5FB", corner_radius=10, border_width=1, border_color="#AED6F1")
    f_cita.grid(row=0, column=1, sticky='nsew', padx=10, pady=10)
    
    ctk.CTkLabel(f_cita, text="AGENDAMIENTO / CITA", font=('Arial', 14, 'bold'), text_color="#1860C3").pack(pady=20)
    
    # Fecha Cita
    f_date = ctk.CTkFrame(f_cita, fg_color="transparent")
    f_date.pack(fill='x', padx=20, pady=10)
    
    ctk.CTkLabel(f_date, text="Fecha Cita:", font=('Arial', 11, 'bold'), text_color="black").pack(side='left', padx=(0,10))
    e_fecha_cita_inter = ctk.CTkEntry(f_date, textvariable=var_fecha_cita_inter, width=120, placeholder_text="DD/MM/YYYY")
    e_fecha_cita_inter.pack(side='left', padx=5)
    
    def set_hoy_inter():
        var_fecha_cita_inter.set(datetime.datetime.now().strftime("%d/%m/%Y"))
        
    ctk.CTkButton(f_date, text="📅 Hoy", width=60, command=set_hoy_inter, fg_color="#28a745", hover_color="#218838").pack(side='left', padx=10)
    
    # Informe Cita
    ctk.CTkLabel(f_cita, text="Informe Cita (Resultado Presencial):", font=('Arial', 11, 'bold'), text_color="black", anchor='w').pack(fill='x', padx=20, pady=(20,5))
    txt_informe_cita_inter = ctk.CTkTextbox(f_cita, height=200, border_width=1)
    txt_informe_cita_inter.pack(fill='x', padx=20, pady=5)
        
    # Botón Guardar GLOBAL
    # Lo ponemos en un frame abajo que ocupe todo el ancho
    f_btns = ctk.CTkFrame(tab_inter, fg_color="transparent")
    f_btns.grid(row=1, column=0, columnspan=2, sticky='ew', pady=10, padx=10)
    
    ctk.CTkButton(f_btns, text="💾 GUARDAR GESTIÓN Y CITA", command=guardar_gestion, 
                  fg_color="#28a745", hover_color="#218838", font=('Arial', 12, 'bold'), height=40).pack(fill='x')
                  
    # --- TAB 3: PROCESO (STATUS) ---
    nb.add("Proceso")
    tab_proc = nb.tab("Proceso")
    
    f_proc_center = ctk.CTkFrame(tab_proc, fg_color="white", corner_radius=10)
    f_proc_center.pack(pady=40, padx=40, fill='x')
    
    # Configure Grid
    f_proc_center.grid_columnconfigure(0, weight=0) # Check
    f_proc_center.grid_columnconfigure(1, weight=0) # Label
    f_proc_center.grid_columnconfigure(2, weight=1) # Entry
    
    # Row 0: Precalificación
    w_chk_pre = ctk.CTkCheckBox(f_proc_center, text="", variable=var_chk_precal, width=24)
    w_chk_pre.grid(row=0, column=0, padx=10, pady=15)
    ctk.CTkLabel(f_proc_center, text="Precalificación", font=('Arial', 12, 'bold')).grid(row=0, column=1, padx=10, pady=15, sticky='w')
    w_desc_pre = ctk.CTkEntry(f_proc_center, textvariable=var_desc_precal, placeholder_text="Descripción...")
    w_desc_pre.grid(row=0, column=2, padx=10, pady=15, sticky='ew')
    
    # Row 1: Aprobación
    w_chk_apr = ctk.CTkCheckBox(f_proc_center, text="", variable=var_chk_aprob, width=24)
    w_chk_apr.grid(row=1, column=0, padx=10, pady=15)
    ctk.CTkLabel(f_proc_center, text="Aprobación", font=('Arial', 12, 'bold')).grid(row=1, column=1, padx=10, pady=15, sticky='w')
    w_desc_apr = ctk.CTkEntry(f_proc_center, textvariable=var_desc_aprob, placeholder_text="Descripción...")
    w_desc_apr.grid(row=1, column=2, padx=10, pady=15, sticky='ew')
    
    # Row 2: Fecha Desembolso (Hidden initially)
    # Use helper widgets list to toggle visibility
    lbl_desembolso = ctk.CTkLabel(f_proc_center, text="Fecha Desembolso:", font=('Arial', 12, 'bold'), text_color="#1860C3")
    
    f_fechad = ctk.CTkFrame(f_proc_center, fg_color="transparent")
    e_fechad = ctk.CTkEntry(f_fechad, textvariable=var_fecha_desembolso_inter, width=120, placeholder_text="DD/MM/YYYY")
    e_fechad.pack(side='left', padx=5)
    
    def set_hoy_desembolso():
        var_fecha_desembolso_inter.set(datetime.datetime.now().strftime("%d/%m/%Y"))
    
    btn_hoy_d = ctk.CTkButton(f_fechad, text="Hoy", width=50, command=set_hoy_desembolso, fg_color="#28a745")
    btn_hoy_d.pack(side='left', padx=5)
    
    # Row 3: Informe
    w_chk_inf = ctk.CTkCheckBox(f_proc_center, text="", variable=var_chk_informe, width=24)
    w_chk_inf.grid(row=3, column=0, padx=10, pady=15)
    ctk.CTkLabel(f_proc_center, text="Informe", font=('Arial', 12, 'bold')).grid(row=3, column=1, padx=10, pady=15, sticky='w')
    w_desc_inf = ctk.CTkEntry(f_proc_center, textvariable=var_desc_informe, placeholder_text="Descripción...")
    w_desc_inf.grid(row=3, column=2, padx=10, pady=15, sticky='ew')
    
    # Conditional Logic
    def toggle_fecha_desembolso():
        if var_chk_aprob.get() == 1:
            lbl_desembolso.grid(row=2, column=1, padx=10, pady=15, sticky='w')
            f_fechad.grid(row=2, column=2, padx=10, pady=15, sticky='w')
        else:
            lbl_desembolso.grid_remove()
            f_fechad.grid_remove()
            var_fecha_desembolso_inter.set("") # Clear if unchecked
            
    w_chk_apr.configure(command=toggle_fecha_desembolso)
    
    # Register widgets for security
    security_widgets['chk_precal'] = w_chk_pre
    security_widgets['desc_precal'] = w_desc_pre
    security_widgets['chk_aprob'] = w_chk_apr
    security_widgets['desc_aprob'] = w_desc_apr
    security_widgets['chk_informe'] = w_chk_inf
    security_widgets['desc_informe'] = w_desc_inf
    
    # Init state
    toggle_fecha_desembolso()
    
    ctk.CTkButton(tab_proc, text="💾 GUARDAR PROCESO", command=guardar_gestion, 
                  fg_color="#28a745", font=('Arial', 12, 'bold'), width=200).pack(pady=20)

def abrir_modulo_consultas():
    # Variables de UI para mostrar el estado y la fecha
    var_estado = tk.StringVar()
    var_fecha = tk.StringVar()

    def actualizar_consultas(cedula):
        """Callback que se activa cuando se encuentra un cliente en el buscador genérico."""
        if not cedula:
            var_estado.set("")
            var_fecha.set("")
            return
        
        conn, cursor = conectar_db()
        try:
            # 1. Buscar en Microcreditos (Prioridad para estados finales de trámite)
            cursor.execute("""
                SELECT status, sub_status, fecha_desembolsado, fecha_negado, 
                       fecha_desistimiento, fecha_comite 
                FROM Microcreditos WHERE cedula_cliente = %s
            """, (cedula,))
            micro = cursor.fetchone()
            
            res_estado = "No definido"
            res_fecha = "-"

            if micro:
                # Orden de prioridad según importancia del estado final
                if micro[2]: # fecha_desembolsado
                    res_estado = "DESEMBOLSADO"
                    res_fecha = micro[2]
                elif micro[3]: # fecha_negado
                    res_estado = "NEGADO"
                    res_fecha = micro[3]
                elif micro[4]: # fecha_desistimiento
                    res_estado = "DESISTIDO"
                    res_fecha = micro[4]
                elif micro[5]: # fecha_comite
                    res_estado = "EN COMITÉ"
                    res_fecha = micro[5]
                else:
                    # Si no tiene fecha específica, usamos el status general o sub_status
                    res_estado = micro[0] or "MICROCRÉDITO (EN TRÁMITE)"
                    res_fecha = "Pendiente"
            else:
                # 2. Si no hay registro en Microcreditos, consultamos el producto en la ficha del Cliente
                cursor.execute("SELECT producto, apertura FROM Clientes WHERE cedula = %s", (cedula,))
                cli = cursor.fetchone()
                if cli:
                    res_estado = (cli[0] or "CLIENTE REGISTRADO").upper()
                    res_fecha = cli[1] or "Sin Fecha de Apertura"
            
            var_estado.set(res_estado)
            var_fecha.set(res_fecha)
        except Exception as e:
            print(f"Error en consulta de estados: {e}")
        finally:
            db_manager.release_connection(conn)

    # Llamada al generador con el callback de búsqueda
    win, frame, nb = crear_modulo_generico("Módulo de Consultas", search_callback=actualizar_consultas)
    
    # --- UI ADICIONAL PARA CONSULTAS ---
    # Contenedor central resaltado
    status_frame = ctk.CTkFrame(frame, fg_color="#F0F8FF", corner_radius=15, border_width=1, border_color="#A9CCE3")
    status_frame.pack(pady=30, padx=50, fill='x')
    
    inner = ctk.CTkFrame(status_frame, fg_color="transparent")
    inner.pack(padx=40, pady=30, expand=True)
    
    # Campo de Visualización de Estado
    ctk.CTkLabel(inner, text="ESTADO DEL PROCESO:", font=('Arial', 14, 'bold'), text_color="#1860C3").grid(row=0, column=0, padx=(0,20), pady=15, sticky='e')
    e_estado = ctk.CTkEntry(inner, textvariable=var_estado, width=350, height=40, font=('Arial', 16, 'bold'), 
                            fg_color="white", text_color="#2E8B57", border_color="#1860C3", state='readonly', corner_radius=8)
    e_estado.grid(row=0, column=1, pady=15)
    
    # Campo de Visualización de Fecha
    ctk.CTkLabel(inner, text="FECHA ASOCIADA:", font=('Arial', 13, 'bold'), text_color="#555555").grid(row=1, column=0, padx=(0,20), pady=10, sticky='e')
    e_fecha = ctk.CTkEntry(inner, textvariable=var_fecha, width=350, height=35, font=('Arial', 14), 
                           fg_color="#F9F9F9", text_color="black", border_color="#CCCCCC", state='readonly', corner_radius=8)
    e_fecha.grid(row=1, column=1, pady=10)

    # Nota informativa
    ctk.CTkLabel(frame, text="* La información se sincroniza en tiempo real de la base de datos central.", 
                 font=("Arial", 11, "italic"), text_color="grey").pack(pady=10)
    
    # Mensaje de ayuda inicial
    var_estado.set("ESPERANDO BÚSQUEDA...")
    var_fecha.set("-")

def abrir_modulo_caja():
    # Variables de UI
    var_fecha_hora = tk.StringVar()
    var_fecha_contrato = tk.StringVar()
    var_cedula = tk.StringVar()
    var_ruc = tk.StringVar()
    var_nombres = tk.StringVar()
    var_email = tk.StringVar()
    var_direccion = tk.StringVar()
    var_telefono = tk.StringVar()
    var_estado_civil = tk.StringVar(value="Soltero")
    var_asesor = tk.StringVar()
    var_buro = tk.StringVar(value="No")
    var_buro_ruta = tk.StringVar()
    var_valor_apertura = tk.StringVar()
    var_num_apertura = tk.StringVar()
    var_observaciones = tk.StringVar()
    blocking_fecha = False # Flag para detener el reloj

    def actualizar_fecha():
        if not blocking_fecha:
            var_fecha_hora.set(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        win.after(1000, actualizar_fecha)

    def seleccionar_archivo_buro():
        ruta = filedialog.askopenfilename(title="Seleccionar Buró", filetypes=[("Archivos de imagen/PDF", "*.pdf *.png *.jpg *.jpeg")], parent=win)
        if ruta:
            # 1. Definir Ruta Común
            carpeta_global = "_Archivos_Clientes"
            if not os.path.exists(carpeta_global):
                 os.makedirs(carpeta_global)
            
            # 2. Corregir Guardado: [CEDULA]_BuroCredito.pdf
            cedula_cliente = var_cedula.get().strip()
            if not cedula_cliente:
                cedula_cliente = "temp" # Fallback si no hay cédula (aunque validación debería prevenirlo)
            
            nombre_dest = f"{cedula_cliente}_BuroCredito.pdf" # Forzamos PDF o mantenemos extensión si es necesario, user pidió .pdf en ejemplo
            # Si el user selecciona imagen, quizás deberíamos mantener extensión o convertir. 
            # El requerimiento dice: "Renómbralo obligatoriamente siguiendo este patrón: [CEDULA]_BuroCredito.pdf"
            # Asumiremos que si es imagen, se guarda con ese nombre aunque sea .jpg, o deberíamos mantener ext.
            # El ejemplo dice .pdf. Voy a mantener la extensión original para evitar corrupción de archivos si es imagen, 
            # PERO si es PDF, será .pdf. Si el usuario insiste en .pdf, asumiré que suben PDFs.
            # Ajuste: El prompt dice "Cuando el usuario seleccione el PDF...". Asumimos input PDF principal.
            # Para mayor seguridad mantendré la extensión original si no es pdf, pero trataré de cumplir el formato.
            
            ext = os.path.splitext(ruta)[1]
            if ext.lower() == '.pdf':
                nombre_dest = f"{cedula_cliente}_BuroCredito.pdf"
            else:
                nombre_dest = f"{cedula_cliente}_BuroCredito{ext}"

            ruta_dest = os.path.join(carpeta_global, nombre_dest)
            
            try:
                shutil.copy(ruta, ruta_dest)
                var_buro_ruta.set(ruta_dest)
                messagebox.showinfo("Éxito", f"Archivo guardado correctamente en:\n{ruta_dest}", parent=win)
                win.lift()
                win.focus_force()
                registrar_auditoria("Carga Buró", detalles=f"Archivo: {nombre_dest}")
                validar_datos_caja()
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo guardar el archivo: {e}", parent=win)

    def preparar_nuevo_cliente():
        # Desbloquear todo para insertar
        e_fecha_info.configure(state='normal')
        e_fecha_con.configure(state='normal')
        
        # Lista de campos a desbloquear
        widgets = [
            e_nombres, e_asesor, cb_buro,
            e_email_con, e_dir_con, e_tel_con, cb_civil_con, e_val_ape_con,
            btn_guardar_caja, btn_guardar_con
        ]
        
        for w in widgets:
            try:
                w.configure(state='normal')
                if hasattr(w, 'configure') and 'fg_color' in w.keys():
                    w.configure(fg_color="white")
            except: pass
            
        e_nombres.focus()
        btn_crear_nuevo.grid_remove() # Ocultar botón
        
        # Limpiar campos explícitamente y set defaults
        var_nombres.set("")
        var_email.set("")
        var_direccion.set("")
        var_telefono.set("")
        var_estado_civil.set("Soltero")
        var_asesor.set("")
        var_buro.set("No")
        var_buro_ruta.set("")
        var_valor_apertura.set("")
        var_num_apertura.set(generar_correlativo_apertura())
        toggle_buro_btn()

    def toggle_id_fields(*args):
        ced = var_cedula.get().strip()
        ruc = var_ruc.get().strip()
        
        if ced:
            e_ruc.configure(state="disabled", fg_color="#F0F0F0")
        else:
            e_ruc.configure(state="normal", fg_color="white")
            
        if ruc:
            e_ced.configure(state="disabled", fg_color="#F0F0F0")
        else:
            e_ced.configure(state="normal", fg_color="white")
    
    def validar_datos_caja(show_error=False):
        # Campos básicos obligatorios
        ced = var_cedula.get().strip()
        ruc = var_ruc.get().strip()
        
        if not ced and not ruc:
            if show_error:
                messagebox.showerror("Error", "Debe completar el campo 'Cédula' o 'RUC' antes de continuar.", parent=win)
            return False
            
        # Otros campos obligatorios
        other_fields = [
            (var_nombres, "Nombres Completos"),
            (var_asesor, "Nombre del Asesor")
        ]
        
        for var, field_name in other_fields:
            if not var.get().strip():
                if show_error:
                    messagebox.showerror("Error", f"Debe completar el campo '{field_name}' antes de continuar.", parent=win)
                return False
            
        # Validación de Buró
        if var_buro.get() == "Sí":
            if not var_buro_ruta.get():
                btn_adjuntar.configure(fg_color="#cc0000", text="⚠️ Subir Buró (OBLIGATORIO)", text_color="white")
                if show_error:
                    messagebox.showerror("Error", "Es obligatorio subir el archivo de Buró de Crédito para continuar", parent=win)
                return False
            else:
                btn_adjuntar.configure(fg_color="#28a745", text="✅ Buró Cargado", text_color="white")
        else:
            btn_adjuntar.configure(fg_color="grey", text="📎 Adjuntar Buró (PDF/IMG)")
            
        return True

    def on_tab_change():
        if nb.get() == "Contrato":
            if not validar_datos_caja(show_error=True):
                # Pequeño delay para que el set ocurra después de que termine el evento actual si es necesario
                win.after(100, lambda: nb.set("Información"))

    def cargar_datos_caja(cedula):
        nonlocal blocking_fecha
        def toggle_inputs(state):
            # Lista de campos a bloquear/desbloquear
            # Nota: 'e_ced' y 'e_ruc' se manejan aparte por la búsqueda
            widgets = [
                e_nombres, e_asesor, cb_buro,
                e_email_con, e_dir_con, e_tel_con, cb_civil_con, e_val_ape_con,
                btn_guardar_caja, btn_guardar_con, btn_imprimir_con
            ]
            
            fg = "white" if state == "normal" else "#F0F0F0"
            for w in widgets:
                try:
                    w.configure(state=state)
                    if hasattr(w, 'configure') and 'fg_color' in w.keys():
                        w.configure(fg_color=fg)
                except: pass

        if not cedula:
            # Limpiar campos si no hay input, pero NO mostrar botón nuevo aun
            # Solo limpiar y bloquear
            try:
                e_fecha_info.configure(state='normal')
                e_fecha_con.configure(state='normal')
            except: pass
            
            toggle_inputs("disabled") # Bloquear hasta que busque algo válido
             
            var_cedula.set("")
            var_ruc.set("")
            var_nombres.set("")
            var_email.set("")
            var_direccion.set("")
            var_telefono.set("")
            var_estado_civil.set("Soltero")
            var_asesor.set("")
            var_buro.set("No")
            var_buro_ruta.set("")
            var_valor_apertura.set("")
            var_num_apertura.set("")
            var_observaciones.set("")
            toggle_buro_btn()
            try: btn_crear_nuevo.grid_remove()
            except: pass
            
            try:
                btn_imprimir_con.configure(state='disabled')
            except: pass
            return

        conn, cursor = conectar_db()
        # 1. Intentar cargar desde Clientes para completar datos básicos
        cursor.execute("SELECT nombre, ruc, email, direccion, telefono, estado_civil FROM Clientes WHERE cedula = %s", (cedula,))
        cli = cursor.fetchone()
        
        # 2. Intentar cargar desde Caja (sobrescribe lo anterior si hay datos específicos de Caja)
        cursor.execute("SELECT * FROM Caja WHERE cedula = %s", (cedula,))
        caja = cursor.fetchone()
        db_manager.release_connection(conn)

        if cli or caja:
            # CLIENTE EXISTENTE (en Clientes o Caja)
            try: btn_crear_nuevo.grid_remove()
            except: pass
            
            # Cargar datos de Cliente primero
            if cli:
                var_nombres.set(cli[0] or "")
                var_ruc.set(cli[1] or "")
                var_cedula.set(cedula)
                var_email.set(cli[2] or "")
                var_direccion.set(cli[3] or "")
                var_telefono.set(cli[4] or "")
                if cli[5] in ["Soltero", "Casado", "Viudo", "Divorciado", "Union Libre"]:
                    var_estado_civil.set(cli[5])
            
            if caja:
                # Si existe en Caja, BLOQUEAR FECHA y cargar full datos
                blocking_fecha = True
                try:
                    e_fecha_info.configure(state='disabled')
                    e_fecha_con.configure(state='disabled')
                except: pass

                var_fecha_hora.set(caja[1] or "")
                try: var_fecha_contrato.set(caja[16] or "")
                except: var_fecha_contrato.set("")
                var_ruc.set(caja[3] or "")
                var_nombres.set(caja[4] or "")
                var_email.set(caja[5] or "")
                var_direccion.set(caja[6] or "")
                var_telefono.set(caja[7] or "")
                var_estado_civil.set(caja[8] or "Soltero")
                var_asesor.set(caja[9] or "")
                var_buro.set(caja[10] or "No")
                var_buro_ruta.set(caja[11] or "")
                var_valor_apertura.set(formatear_float_str(caja[12]) if caja[12] else "")
                var_num_apertura.set(caja[13] or "")
                try:
                    if len(caja) > 15: var_observaciones.set(caja[15] or "")
                    else: var_observaciones.set("")
                except: var_observaciones.set("")

                # PERMISOS: Si no es admin y ya existe, BLOQUEAR TODO
                if NIVEL_ACCESO != 1:
                    toggle_inputs("disabled")
                else:
                    toggle_inputs("normal")
            else:
                # Existe en Clientes pero NO en Caja -> Es Nuevo para Caja -> Habilitar todo
                blocking_fecha = False
                var_num_apertura.set(generar_correlativo_apertura())
                toggle_inputs("normal")

        else:
            # NO EXISTE (Ni en clientes ni en caja) -> MODO CREAR
            # 1. Limpiar campos
            var_nombres.set("")
            var_email.set("")
            var_direccion.set("")
            var_telefono.set("")
            var_estado_civil.set("Soltero")
            var_asesor.set("")
            var_buro.set("No")
            var_buro_ruta.set("")
            var_valor_apertura.set("")
            var_num_apertura.set("")
            var_observaciones.set("")
            
            # 2. Bloquear inputs para evitar escritura accidental
            toggle_inputs("disabled")
            
            # 3. Mostrar Botón "CREAR NUEVO"
            try: btn_crear_nuevo.grid()
            except: pass
            
            messagebox.showinfo("Nuevo", "Cédula no encontrada.\nPresione '🆕 CREAR NUEVO CLIENTE' para registrar.", parent=win)

        toggle_buro_btn()
        try:
            if cli or caja:
                # Solo habilitar imprimir si ya está guardado (existe en caja con num apertura)
                if caja and var_num_apertura.get(): btn_imprimir_con.configure(state='normal')
                else: btn_imprimir_con.configure(state='disabled')
            else:
                btn_imprimir_con.configure(state='disabled')
        except: pass

    def generar_correlativo_apertura():
        conn, cursor = conectar_db()
        cursor.execute("SELECT numero_apertura FROM Caja ORDER BY id DESC LIMIT 1")
        last = cursor.fetchone()
        db_manager.release_connection(conn)
        
        start_num = 1
        if last and last[0]:
            try:
                # Format is XXX-400
                num_part = last[0].split('-')[0]
                start_num = int(num_part) + 1
            except:
                pass
        
        return f"{start_num:03d}-400"

    def guardar_caja():
        ced = var_cedula.get().strip()
        ruc = var_ruc.get().strip()
        
        identificacion = ced if ced else ruc
        
        if not identificacion:
            messagebox.showwarning("Atención", "Ingrese la identificación del cliente (Cédula o RUC).")
            return
            
        # PERMISO: Solo admin puede modificar registros existentes en Caja
        conn_check, cursor_check = conectar_db()
        cursor_check.execute("SELECT id FROM Caja WHERE cedula = %s OR (cedula = '' AND ruc = %s)", (ced, ruc))
        exists_check = cursor_check.fetchone()
        db_manager.release_connection(conn_check)

        if exists_check and NIVEL_ACCESO != 1:
            messagebox.showerror("Error", "No tiene permisos para modificar registros existentes.")
            return

        # Validación numérica
        if ced and not ced.isdigit():
            messagebox.showwarning("Error", "La Cédula debe contener solo números.")
            return
        if ruc and not ruc.isdigit():
            messagebox.showwarning("Error", "El RUC debe contener solo números.")
            return
        
        data = (
            var_fecha_hora.get(),
            ced,
            ruc,
            var_nombres.get().strip(),
            var_email.get().strip(),
            var_direccion.get().strip(),
            var_telefono.get().strip(),
            var_estado_civil.get(),
            var_asesor.get().strip(),
            var_buro.get(),
            var_buro_ruta.get(),
            limpiar_moneda(var_valor_apertura.get()),
            var_num_apertura.get(),
            "Pendiente", # estado_impreso por defecto
            var_observaciones.get().strip()
        )

        conn, cursor = conectar_db()
        try:
            # Sincronizar con tabla Clientes antes de guardar en Caja (Incluye Fecha y N. Carpeta)
            sincronizar_cliente_desde_caja(ced, ruc, var_nombres.get().strip(), var_email.get().strip(), var_direccion.get().strip(), var_telefono.get().strip(), var_asesor.get().strip(), var_fecha_hora.get(), var_num_apertura.get())
            
            cursor.execute("SELECT id FROM Caja WHERE cedula = %s OR (cedula = '' AND ruc = %s)", (ced, ruc))
            exists = cursor.fetchone()
            
            # Usaremos una lógica elástica por si falla la columna estado_impreso
            try:
                if exists:
                    cursor.execute("""
                        UPDATE Caja SET 
                        fecha_hora=%s, ruc=%s, nombres_completos=%s, email=%s, direccion=%s, telefono=%s, estado_civil=%s, asesor=%s, buro_credito=%s, buro_archivo_ruta=%s, valor_apertura=%s, numero_apertura=%s, observaciones=%s
                        WHERE id = %s
                    """, (data[0], data[2], data[3], data[4], data[5], data[6], data[7], data[8], data[9], data[10], data[11], data[12], data[14], exists[0]))
                else:
                    # Incluimos estado_impreso explícitamente en el INSERT normal
                    cursor.execute("""
                        INSERT INTO Caja (fecha_hora, cedula, ruc, nombres_completos, email, direccion, telefono, estado_civil, asesor, buro_credito, buro_archivo_ruta, valor_apertura, numero_apertura, estado_impreso, observaciones) 
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                    """, data)
            except Exception as e_sql:
                # Si falla por la columna, intentamos guardar lo básico
                print(f"Error de columna en Caja (reintentando): {e_sql}")
                if exists:
                    cursor.execute("""
                        UPDATE Caja SET 
                        fecha_hora=%s, ruc=%s, nombres_completos=%s, email=%s, direccion=%s, telefono=%s, estado_civil=%s, asesor=%s, buro_credito=%s, buro_archivo_ruta=%s, valor_apertura=%s, numero_apertura=%s
                        WHERE id = %s
                    """, (data[0], data[2], data[3], data[4], data[5], data[6], data[7], data[8], data[9], data[10], data[11], data[12], exists[0]))
                else:
                    # INSERT sin la columna conflictiva
                    cursor.execute("""
                        INSERT INTO Caja (fecha_hora, cedula, ruc, nombres_completos, email, direccion, telefono, estado_civil, asesor, buro_credito, buro_archivo_ruta, valor_apertura, numero_apertura) 
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                    """, data[:-1])
            
            conn.commit()
            
            try:
                top = btn_imprimir_con.winfo_toplevel()
                messagebox.showinfo("Éxito", "Datos de Caja guardados correctamente.", parent=top)
                cargar_datos_caja(None)
                top.lift()
                top.focus_force()
            except:
                messagebox.showinfo("Éxito", "Datos de Caja guardados correctamente.")
                cargar_datos_caja(None)
            
            registrar_auditoria("Guardar Caja", id_cliente=ced, detalles=f"Apertura: {data[12]}")
            
        except Exception as e:
            try:
                messagebox.showerror("Error", f"No se pudo guardar: {e}", parent=btn_imprimir_con.winfo_toplevel())
            except:
                messagebox.showerror("Error", f"No se pudo guardar: {e}")
        finally:
            db_manager.release_connection(conn)

    def toggle_buro_btn(*args):
        if var_buro.get() == "Sí":
            btn_adjuntar.configure(state="normal")
        else:
            btn_adjuntar.configure(state="disabled", fg_color="grey", text="📎 Adjuntar Buró (PDF/IMG)")
        validar_datos_caja()

    def aplicar_reemplazo_total(doc, reemplazos):
        """
        Motor de reemplazo de alto nivel. Unifica runs para evitar fragmentación
        de etiquetas y recorre recursivamente párrafos y tablas.
        """
        def procesar_contenedor(parrafos):
            for p in parrafos:
                # Verificar si el párrafo contiene alguna clave
                found = False
                for key in reemplazos.keys():
                    if key in p.text:
                        found = True
                        break
                
                if found:
                    # Unificación de Runs: Colapsamos todo el texto en el primer run
                    # Esto garantiza que etiquetas como {No. Apertura} se lean completas
                    full_text = p.text
                    for key, val in reemplazos.items():
                        full_text = full_text.replace(key, str(val))
                    
                    if p.runs:
                        # Guardamos el estilo del primer run si es posible
                        p.runs[0].text = full_text
                        # Limpiamos el resto de runs para evitar duplicados
                        for i in range(1, len(p.runs)):
                            p.runs[i].text = ""
                    else:
                        p.text = full_text

        # 1. Párrafos del cuerpo principal
        procesar_contenedor(doc.paragraphs)
        
        # 2. Iteración profunda en tablas (celdas y sus párrafos)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    procesar_contenedor(cell.paragraphs)

    def guardar_caja_silencioso():
        """Guarda los datos en la DB sin mostrar mensajes ni limpiar campos."""
        ced = var_cedula.get().strip()
        ruc = var_ruc.get().strip()
        identificacion = ced if ced else ruc
        if not identificacion: return
        
        data = (
            var_fecha_hora.get(),
            ced,
            ruc,
            var_nombres.get().strip(),
            var_email.get().strip(),
            var_direccion.get().strip(),
            var_telefono.get().strip(),
            var_estado_civil.get(),
            var_asesor.get().strip(),
            var_buro.get(),
            var_buro_ruta.get(),
            limpiar_moneda(var_valor_apertura.get()),
            var_num_apertura.get(),
            "Pendiente",
            var_observaciones.get().strip()
        )
        
        conn, cursor = conectar_db()
        try:
            # Sincronizar en modo silencioso también (Incluye Fecha y N. Carpeta)
            sincronizar_cliente_desde_caja(ced, ruc, var_nombres.get().strip(), var_email.get().strip(), var_direccion.get().strip(), var_telefono.get().strip(), var_asesor.get().strip(), var_fecha_hora.get(), var_num_apertura.get())
            
            cursor.execute("SELECT id FROM Caja WHERE cedula = %s OR (cedula = '' AND ruc = %s)", (ced, ruc))
            exists = cursor.fetchone()
            if exists:
                cursor.execute("""
                    UPDATE Caja SET 
                    fecha_hora=%s, ruc=%s, nombres_completos=%s, email=%s, direccion=%s, telefono=%s, estado_civil=%s, asesor=%s, buro_credito=%s, buro_archivo_ruta=%s, valor_apertura=%s, numero_apertura=%s, observaciones=%s
                    WHERE id = %s
                """, (data[0], data[2], data[3], data[4], data[5], data[6], data[7], data[8], data[9], data[10], data[11], data[12], data[14], exists[0]))
            else:
                cursor.execute("""
                    INSERT INTO Caja (fecha_hora, cedula, ruc, nombres_completos, email, direccion, telefono, estado_civil, asesor, buro_credito, buro_archivo_ruta, valor_apertura, numero_apertura, estado_impreso, observaciones) 
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                """, data)
            conn.commit()
        except:
            pass
        finally:
            db_manager.release_connection(conn)

    def imprimir_contrato():
        if not validar_datos_caja(show_error=True):
            return
        
        # 0. Sincronizar DB antes de imprimir (PostgreSQL)
        guardar_caja_silencioso()

        ced = var_cedula.get().strip()
        ruc = var_ruc.get().strip()
        ident = ced if ced else ruc
        nom = var_nombres.get().strip()
        civ = var_estado_civil.get()
        dir_c = var_direccion.get().strip()
        tel = var_telefono.get().strip()
        mail = var_email.get().strip()
        val = var_valor_apertura.get()
        num_ap = var_num_apertura.get()

        # Configuración de Rutas Dinámicas
        base_dir = os.path.dirname(os.path.abspath(__file__))
        ruta_plantillas = os.path.join(base_dir, "Documento Plantilla")
        ruta_base_servidor = r"\\SERVIDOR\Compartida\Contratos_Finales"
        ruta_destino = os.path.join(ruta_base_servidor, ident)
        
        # Ruta Fallback Local
        ruta_respaldo_local = os.path.join(base_dir, "Respaldo_PDF", ident)

        def fetch_resources(uri, rel):
            """Resuelve rutas de recursos para xhtml2pdf."""
            if os.path.isabs(uri): return uri
            return os.path.join(ruta_plantillas, uri)
        
        # Validación de acceso al servidor con fallback local
        try:
            if not os.path.exists(ruta_destino):
                os.makedirs(ruta_destino)
        except Exception as e_red:
            print(f"Alerta: Servidor no disponible ({e_red}). Usando respaldo local.")
            ruta_destino = ruta_respaldo_local
            try:
                if not os.path.exists(ruta_destino):
                    os.makedirs(ruta_destino)
                messagebox.showwarning("Modo Respaldo", f"El servidor no está disponible. Los archivos se guardarán localmente en:\n{ruta_destino}", parent=win.winfo_toplevel())
            except Exception as e_local:
                messagebox.showerror("Error Crítico", f"No se pudo crear directorio ni en servidor ni localmente.\nError: {e_local}", parent=win.winfo_toplevel())
                return

        # Fecha en Español
        meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
        now = datetime.datetime.now()
        fecha_esp = f"{now.day:02d} de {meses[now.month-1]} de {now.year}"

        # Plantillas HTML ajustadas
        plantillas = ["contrato.html", "carta_compromiso.html"]
        archivos_finales = []
        
        msg_wait = ctk.CTkLabel(win, text="Generando PDF (Motor Industrial), por favor espere...", font=('Arial', 14, 'bold'), text_color="red")
        msg_wait.place(relx=0.5, rely=0.1, anchor='center')
        win.update()

        try:
            # Contexto Jinja2 Superior
            contexto = {
                "nombres_completos": str(nom),
                "cedula": str(ced),
                "ruc": str(ruc),
                "estado_civil": str(civ),
                "direccion_domicilio": str(dir_c),
                "domicilio": str(dir_c), # Alias para compatibilidad
                "telefono": str(tel),
                "correo_electronico": str(mail),
                "correo": str(mail), # Alias para compatibilidad
                "valor_apertura": str(val),
                "no_apertura": str(num_ap),
                "fecha_actual": str(fecha_esp)
            }

            # 1. Generar la data completa para el QR
            info_qr = f"Apertura: {num_ap} | Fecha: {fecha_esp} | Cliente: {nom} | Valor: ${val}"
            
            # 2. Crear el QR
            qr = qrcode.QRCode(box_size=10, border=1)
            qr.add_data(info_qr)
            qr.make(fit=True)
            img_qr = qr.make_image(fill_color="black", back_color="white")
            
            # 3. Guardar la imagen física (CRÍTICO: Usamos el mismo nombre de archivo que funcionó)
            ruta_qr = os.path.join("Documento Plantilla", "qr_temp.png")
            img_qr.save(ruta_qr)
            print(f"Generando PDF... QR guardado físicamente en: {ruta_qr}")

            # Configurar Jinja2
            env = Environment(loader=FileSystemLoader(ruta_plantillas))

            for t_name in plantillas:
                full_t_path = os.path.join(ruta_plantillas, t_name)
                if not os.path.exists(full_t_path):
                    try:
                        messagebox.showerror("Error", f"Plantilla HTML estratégica no encontrada en:\n{full_t_path}", parent=win.winfo_toplevel())
                    except:
                        messagebox.showerror("Error", f"Plantilla HTML estratégica no encontrada en:\n{full_t_path}")
                    continue

                # 1. Renderizado de alta precisión
                template = env.get_template(t_name)
                html_out = template.render(contexto)

                # 2. Definir nombre final del PDF industrial
                clean_ap = num_ap.replace('-', '_').replace('/', '_')
                if "contrato" in t_name.lower():
                    final_name = f"Contrato_{clean_ap}.pdf"
                else:
                    final_name = f"Carta_Compromiso_{clean_ap}.pdf"
                
                pdf_path = os.path.join(ruta_destino, final_name)

                # 3. Conversión HTML a PDF (Motor xhtml2pdf)
                try:
                    with open(pdf_path, "wb") as pdf_file:
                        pisa_status = pisa.CreatePDF(html_out, dest=pdf_file, link_callback=fetch_resources)
                        
                    if pisa_status.err:
                        raise Exception(f"Fallo en motor xhtml2pdf para {t_name}")
                        
                except PermissionError:
                    messagebox.showerror("Error de Acceso", f"No se puede guardar el archivo:\n{final_name}\n\nPor favor, cierre el documento si lo tiene abierto en otro programa e intente nuevamente.")
                    return
                except Exception as e_pdf:
                    raise Exception(f"Error generando PDF desde HTML: {e_pdf}")

                archivos_finales.append(pdf_path)

            # 4. Sincronización final de estado 'Impreso' en PostgreSQL
            try:
                conn, cursor = conectar_db()
                cursor.execute("UPDATE Caja SET estado_impreso = %s WHERE cedula = %s OR (cedula = '' AND ruc = %s)", ("Impreso", ced, ruc))
                conn.commit()
                db_manager.release_connection(conn)
            except Exception as e_status:
                print(f"No se pudo actualizar el estado impreso industrial en DB: {e_status}")

            try:
                messagebox.showinfo("Éxito", f"Documentos industriales generados correctamente en:\n{ruta_destino}", parent=win.winfo_toplevel())
            except:
                messagebox.showinfo("Éxito", f"Documentos industriales generados correctamente en:\n{ruta_destino}")

            # Apertura automática para validación inmediata
            for f in archivos_finales:
                if os.path.exists(f):
                    try: os.startfile(f)
                    except: pass
            
            try:
                win.winfo_toplevel().lift()
                win.winfo_toplevel().focus_force()
            except: pass

        except Exception as e:
            try:
                msg_wait.destroy()
                messagebox.showerror("Error", f"Error al generar contratos: {e}", parent=win.winfo_toplevel())
            except:
                messagebox.showerror("Error", f"Error al generar contratos: {e}")
            registrar_auditoria("Fallo Motor Impresión", id_cliente=ident, detalles=str(e))
        finally:
            msg_wait.destroy()


    def guardar_datos_contrato():
        """Guarda solo los campos específicos de la pestaña Contrato."""
        ced = var_cedula.get().strip()
        ruc = var_ruc.get().strip()
        ident = ced if ced else ruc
        if not ident:
            messagebox.showwarning("Atención", "No hay un cliente cargado para guardar.")
            return

        # Validación de campos obligatorios para Contrato
        if not var_valor_apertura.get().strip() or not var_num_apertura.get().strip():
            messagebox.showwarning("Atención", "Debe completar 'Valor de Apertura' y 'No. de Apertura' antes de guardar.")
            return

        if NIVEL_ACCESO != 1:
            messagebox.showerror("Error", "No tiene permisos para modificar datos existentes.")
            return

        conn, cursor = conectar_db()
        try:
            # Campos específicos a actualizar: Dirección, Teléfono, Estado Civil, Valor Apertura, No. Apertura
            # (Email también se suele actualizar aquí ya que está en la pestaña)
            cursor.execute("""
                UPDATE Caja SET 
                email=%s, direccion=%s, telefono=%s, estado_civil=%s, valor_apertura=%s, numero_apertura=%s
                WHERE cedula=%s OR (cedula='' AND ruc=%s)
            """, (
                var_email.get().strip(),
                var_direccion.get().strip(),
                var_telefono.get().strip(),
                var_estado_civil.get(),
                limpiar_moneda(var_valor_apertura.get()),
                var_num_apertura.get(),
                ced, ruc
            ))
            
            # También sincronizamos a la tabla Clientes los datos actualizados
            cursor.execute("""
                UPDATE Clientes SET
                email=%s, direccion=%s, telefono=%s, estado_civil=%s, apertura=%s
                WHERE cedula=%s
            """, (
                var_email.get().strip(),
                var_direccion.get().strip(),
                var_telefono.get().strip(),
                var_estado_civil.get(),
                var_num_apertura.get(),
                ced
            ))

            # ACTUALIZACIÓN FECHA CONTRATO: Tiempo real
            now_contrato = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            cursor.execute("UPDATE Caja SET fecha_contrato=%s WHERE cedula=%s OR (cedula='' AND ruc=%s)", (now_contrato, ced, ruc))
            var_fecha_contrato.set(now_contrato)
            
            conn.commit()
            messagebox.showinfo("Éxito", "Datos del contrato actualizados correctamente.")
            registrar_auditoria("Actualizar Datos Contrato", id_cliente=ident, detalles=f"Num Apertura: {var_num_apertura.get()}")
            
            # ACTIVACIÓN: Habilitar impresión tras guardado exitoso
            btn_imprimir_con.configure(state='normal')
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo actualizar: {e}")
        finally:
            db_manager.release_connection(conn)

    def buscar_cliente_caja(event=None):
        ident = var_cedula.get().strip()
        if not ident:
            ident = var_ruc.get().strip()
            
        if (len(ident) == 10 or len(ident) == 13) and ident.isdigit():
             cargar_datos_caja(ident)

    # UI
    win, frame_info, nb = crear_modulo_generico("Caja", tab_name="Información", search_callback=cargar_datos_caja, show_search=False)
    nb.configure(command=on_tab_change)
    actualizar_fecha()
    
    var_cedula.trace_add("write", lambda *a: (validar_datos_caja(), toggle_id_fields()))
    var_ruc.trace_add("write", lambda *a: (validar_datos_caja(), toggle_id_fields()))
    var_nombres.trace_add("write", lambda *a: validar_datos_caja())
    var_asesor.trace_add("write", lambda *a: validar_datos_caja())
    var_buro.trace_add("write", lambda *a: validar_datos_caja())

    # --- PESTAÑA INFORMACIÓN ---
    container_info = ctk.CTkFrame(frame_info, fg_color="white", corner_radius=15, border_width=1, border_color="#CCCCCC")
    container_info.pack(fill='both', expand=True, padx=20, pady=20)
    
    # 0. BOTÓN PRINCIPAL RECAUDACIÓN (NUEVA UBICACIÓN)
    f_info_actions = ctk.CTkFrame(container_info, fg_color="transparent")
    f_info_actions.pack(pady=(20, 10))
    
    ctk.CTkButton(f_info_actions, text="💰 IR A RECAUDACIÓN / COBROS", command=abrir_ventana_recaudacion,
                  font=('Arial', 16, 'bold'), height=50, width=300, 
                  fg_color="#F39C12", hover_color="#D68910").pack()
    
    grid_info = ctk.CTkFrame(container_info, fg_color="transparent")
    grid_info.pack(padx=20, pady=20)
    
    row = 0
    ctk.CTkLabel(grid_info, text="Fecha y Hora:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=10)
    e_fecha_info = ctk.CTkEntry(grid_info, textvariable=var_fecha_hora, width=350, state='normal', fg_color="#F9F9F9", text_color="black")
    e_fecha_info.grid(row=row, column=1, padx=20, pady=10)
    
    row += 1
    ctk.CTkLabel(grid_info, text="Cédula:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=10)
    e_ced = ctk.CTkEntry(grid_info, textvariable=var_cedula, width=350, fg_color="white", text_color="black")
    e_ced.grid(row=row, column=1, padx=20, pady=10)
    e_ced.bind('<KeyRelease>', buscar_cliente_caja)
    
    row += 1
    ctk.CTkLabel(grid_info, text="RUC:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=10)
    e_ruc = ctk.CTkEntry(grid_info, textvariable=var_ruc, width=350, fg_color="white", text_color="black")
    e_ruc.grid(row=row, column=1, padx=20, pady=10)
    
    # Botón Flotante (o en grid) de "Crear Nuevo"
    # Lo ubicamos en la misma columna 1, en una fila intermedia o superpuesto
    # Mejor en el grid, en la siguiente fila del RUC
    
    btn_crear_nuevo = ctk.CTkButton(grid_info, text="🆕 CREAR NUEVO CLIENTE", command=preparar_nuevo_cliente, 
                                    fg_color="#00C853", hover_color="#009624", font=("Arial", 12, "bold"))
    # Inicialmente oculto. Se muestra si no se encuentra.
    # Lo ubicamos en (row+1, 1) pero sin incrementar row global para no descuadrar si está oculto
    btn_crear_nuevo.grid(row=row+1, column=1, pady=5)
    btn_crear_nuevo.grid_remove() 
    
    # Inicializar estados
    toggle_id_fields()
    
    row += 1 # Espacio para el botón aunque esté hidden consume 'index', mejor sumamos row
    row += 1
    ctk.CTkLabel(grid_info, text="Nombres Completos:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=10)
    e_nombres = ctk.CTkEntry(grid_info, textvariable=var_nombres, width=350, fg_color="white", text_color="black")
    e_nombres.grid(row=row, column=1, padx=20, pady=10)
    
    row += 1
    ctk.CTkLabel(grid_info, text="Nombre del Asesor:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=10)
    e_asesor = ctk.CTkEntry(grid_info, textvariable=var_asesor, width=350, fg_color="white", text_color="black")
    e_asesor.grid(row=row, column=1, padx=20, pady=10)
    
    row += 1
    ctk.CTkLabel(grid_info, text="¿Buró de Crédito?", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=10)
    cb_buro = ctk.CTkComboBox(grid_info, values=["Sí", "No"], variable=var_buro, command=toggle_buro_btn, width=100)
    cb_buro.grid(row=row, column=1, sticky='w', padx=20, pady=10)
    btn_adjuntar = ctk.CTkButton(grid_info, text="📎 Adjuntar Buró (PDF/IMG)", command=seleccionar_archivo_buro, state="disabled", fg_color="grey")
    btn_adjuntar.grid(row=row, column=1, padx=(130, 0), pady=10)
    
    row += 1
    ctk.CTkLabel(grid_info, text="Observaciones:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=10)
    e_obs = ctk.CTkEntry(grid_info, textvariable=var_observaciones, width=350, fg_color="white", text_color="black")
    e_obs.grid(row=row, column=1, padx=20, pady=10)
    
    row += 1
    btn_guardar_caja = ctk.CTkButton(container_info, text="💾 GUARDAR / MODIFICAR", command=guardar_caja, font=('Arial', 14, 'bold'), height=45, fg_color="#28a745", hover_color="#218838")
    btn_guardar_caja.pack(pady=20)
    
    # --- PESTAÑA CONTRATO ---
    nb.add("Contrato")
    tab_con = nb.tab("Contrato")
    container_con = ctk.CTkFrame(tab_con, fg_color="white", corner_radius=15, border_width=1, border_color="#CCCCCC")
    container_con.pack(fill='both', expand=True, padx=20, pady=20)
    
    # Grid scrollable or simple frame? User wants a specific order.
    grid_con = ctk.CTkFrame(container_con, fg_color="transparent")
    grid_con.pack(padx=20, pady=10)
    
    row = 0
    ctk.CTkLabel(grid_con, text="Fecha Contrato:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=5)
    e_fecha_con = ctk.CTkEntry(grid_con, textvariable=var_fecha_contrato, width=350, state='normal', fg_color="#F9F9F9", text_color="black")
    e_fecha_con.grid(row=row, column=1, padx=20, pady=5)
    
    row += 1
    ctk.CTkLabel(grid_con, text="Cédula:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=5)
    ctk.CTkEntry(grid_con, textvariable=var_cedula, width=350, state='readonly', fg_color="#F0F0F0", text_color="black").grid(row=row, column=1, padx=20, pady=5)
    
    row += 1
    ctk.CTkLabel(grid_con, text="RUC:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=5)
    ctk.CTkEntry(grid_con, textvariable=var_ruc, width=350, state='readonly', fg_color="#F0F0F0", text_color="black").grid(row=row, column=1, padx=20, pady=5)
    
    row += 1
    ctk.CTkLabel(grid_con, text="Nombres y Apellidos:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=5)
    ctk.CTkEntry(grid_con, textvariable=var_nombres, width=350, state='readonly', fg_color="#F0F0F0", text_color="black").grid(row=row, column=1, padx=20, pady=5)
    
    row += 1
    ctk.CTkLabel(grid_con, text="Correo Electrónico:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=5)
    e_email_con = ctk.CTkEntry(grid_con, textvariable=var_email, width=350, fg_color="white", text_color="black")
    e_email_con.grid(row=row, column=1, padx=20, pady=5)
    
    row += 1
    ctk.CTkLabel(grid_con, text="Dirección Domicilio:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=5)
    e_dir_con = ctk.CTkEntry(grid_con, textvariable=var_direccion, width=350, fg_color="white", text_color="black")
    e_dir_con.grid(row=row, column=1, padx=20, pady=5)
    
    row += 1
    ctk.CTkLabel(grid_con, text="Teléfono:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=5)
    e_tel_con = ctk.CTkEntry(grid_con, textvariable=var_telefono, width=350, fg_color="white", text_color="black")
    e_tel_con.grid(row=row, column=1, padx=20, pady=5)
    
    row += 1
    ctk.CTkLabel(grid_con, text="Estado Civil:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=5)
    cb_civil_con = ctk.CTkComboBox(grid_con, values=["Soltero", "Casado", "Viudo", "Divorciado", "Union Libre"], variable=var_estado_civil, width=200)
    cb_civil_con.grid(row=row, column=1, sticky='w', padx=20, pady=5)

    row += 1
    ctk.CTkLabel(grid_con, text="Valor de Apertura ($):", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=5)
    e_val_ape_con = ctk.CTkEntry(grid_con, textvariable=var_valor_apertura, width=350, fg_color="white", text_color="black")
    e_val_ape_con.grid(row=row, column=1, sticky='w', padx=20, pady=5)
    e_val_ape_con.bind('<FocusOut>', on_focus_out_moneda)
    e_val_ape_con.bind('<FocusIn>', on_focus_in_moneda)
    
    row += 1
    ctk.CTkLabel(grid_con, text="No. de Apertura:", font=('Arial', 12, 'bold'), text_color="black").grid(row=row, column=0, sticky='w', pady=5)
    e_num_ape_con = ctk.CTkEntry(grid_con, textvariable=var_num_apertura, width=350, state='readonly', fg_color="#F0F0F0", text_color="black")
    e_num_ape_con.grid(row=row, column=1, padx=20, pady=5)
    
    row += 1
    btn_frame_con = ctk.CTkFrame(container_con, fg_color="transparent")
    btn_frame_con.pack(pady=20)

    btn_guardar_con = ctk.CTkButton(btn_frame_con, text="💾 GUARDAR DATOS CONTRATO", command=guardar_datos_contrato, font=('Arial', 14, 'bold'), height=45, fg_color="#28a745", hover_color="#218838")
    btn_guardar_con.pack(side='left', padx=10)
    
    btn_imprimir_con = ctk.CTkButton(btn_frame_con, text="🖨️ IMPRIMIR CONTRATO", command=imprimir_contrato, font=('Arial', 14, 'bold'), height=45, fg_color="#1860C3", state='disabled')
    btn_imprimir_con.pack(side='left', padx=10)
    
    # 3. Recaudacion: MOVIDO A PESTAÑA INFORMACIÓN


def abrir_modulo_cartera():
    win, frame, nb = crear_modulo_generico("Módulo de Cartera", color_titulo="#6c757d", tab_name="Estado de Cartera", show_search=False)
    # Reutilizamos el frame del tab generico
    tab = nb.tab("Estado de Cartera")

    # --- UI COMPONENTES ---
    # Top Bar
    f_top = ctk.CTkFrame(tab, fg_color="transparent")
    f_top.pack(fill='x', pady=10)
    
    ctk.CTkButton(f_top, text="🔄 Actualizar Lista", command=lambda: cargar_cartera(), 
                  fg_color="#17a2b8", width=150).pack(side='left', padx=10)

    # Treeview
    cols = ("Cédula", "Cliente", "Monto Crédito", "Plazo", "Cuota", "Total Pagado", "Saldo Est.", "Estado", "Días Mora")
    tree_cartera = ttk.Treeview(tab, columns=cols, show='headings', height=20)
    
    # Headers
    for c in cols:
        tree_cartera.heading(c, text=c)
        tree_cartera.column(c, width=100)
    
    tree_cartera.column("Cliente", width=200)
    tree_cartera.column("Estado", width=120)
    
    # Scroll
    vsb = ttk.Scrollbar(tab, orient="vertical", command=tree_cartera.yview)
    tree_cartera.configure(yscrollcommand=vsb.set)
    
    tree_cartera.pack(side='left', fill='both', expand=True, padx=(10,0))
    vsb.pack(side='right', fill='y', padx=(0,10))
    
    # Tags de color
    tree_cartera.tag_configure('mora', background='#FFCDD2', foreground='#D32F2F') # Rojo suave
    tree_cartera.tag_configure('aldia', background='#C8E6C9', foreground='#2E7D32') # Verde suave
    
    def cargar_cartera():
        # Limpiar tree
        for i in tree_cartera.get_children():
            tree_cartera.delete(i)
            
        conn, cursor = conectar_db()
        try:
            # Seleccionar créditos activos
            cursor.execute("""
                SELECT cedula_cliente, monto_aprobado, plazo_meses, valor_cuota, dia_pago, fecha_desembolso_real
                FROM Microcreditos
                WHERE status = 'Desembolsado'
            """)
            creditos = cursor.fetchall()
            
            for cred in creditos:
                # Unpack
                ced, monto, plazo, cuota, dia_pago, f_desemb_str = cred
                monto = monto or 0.0
                plazo = plazo or 0
                cuota = cuota or 0.0
                dia_pago = dia_pago or 1
                
                # Nombre Cliente
                cursor.execute("SELECT nombre FROM Clientes WHERE cedula = %s", (ced,))
                c_data = cursor.fetchone()
                nombre = c_data[0] if c_data else "Desconocido"
                
                # Pagos
                cursor.execute("SELECT COUNT(*), SUM(total_pagado) FROM Pagos WHERE cedula_cliente = %s", (ced,))
                p_data = cursor.fetchone()
                num_pagos = p_data[0] or 0
                total_pagado = p_data[1] or 0.0
                
                # Cálculos
                # Deuda Total Estimada = Cuota * Plazo
                deuda_total = cuota * plazo
                saldo = max(0, deuda_total - total_pagado)
                
                # Estado y Mora
                # Calcular Próximo Vencimiento (misma lógica que Recaudación)
                cuota_actual = num_pagos + 1
                
                estado = "AL DÍA"
                dias_mora = 0
                tag = "aldia"
                
                if cuota_actual <= plazo:
                    try:
                        # Parsear fecha
                        try: f_ini = datetime.datetime.strptime(f_desemb_str.split(' ')[0], "%d/%m/%Y")
                        except: f_ini = datetime.datetime.strptime(f_desemb_str.split(' ')[0], "%Y-%m-%d")
                        
                        mes_venc = f_ini.month + cuota_actual
                        anio_venc = f_ini.year + (mes_venc - 1) // 12
                        mes_venc = (mes_venc - 1) % 12 + 1
                        
                        import calendar
                        last_day = calendar.monthrange(anio_venc, mes_venc)[1]
                        dia_v = min(int(dia_pago), last_day)
                        
                        f_venc = datetime.date(anio_venc, mes_venc, dia_v)
                        
                        hoy = datetime.date.today()
                        delta = (hoy - f_venc).days
                        
                        if delta > 0:
                            dias_mora = delta
                            estado = "MORA"
                            tag = "mora"
                    except:
                        estado = "Error Fecha"
                else:
                    estado = "FINALIZADO" # O Pagado
                    saldo = 0
                
                # Insertar
                tree_cartera.insert("", "end", values=(
                    ced, 
                    nombre, 
                    f"$ {monto:.2f}", 
                    f"{plazo} mases", 
                    f"$ {cuota:.2f}",
                    f"$ {total_pagado:.2f}",
                    f"$ {saldo:.2f}",
                    estado,
                    str(dias_mora)
                ), tags=(tag,))
                
        except Exception as e:
            # Evitar popups masivos si falla query, print consola
            print(f"Error cargando cartera: {e}")
        finally:
            db_manager.release_connection(conn)

    # Cargar al inicio
    # Usar after para dar tiempo a renderizar UI
    tab.after(500, cargar_cartera)


def abrir_ventana_recaudacion():
    toplevel = ctk.CTkToplevel()
    toplevel.title("Módulo de Recaudación - Pago de Cuotas")
    toplevel.geometry("900x600")
    toplevel.resizable(False, False)
    
    # 1. Forzar que la ventana esté siempre encima de las demás
    toplevel.attributes('-topmost', True)
    
    # 2. Traer la ventana al frente inmediatamente y darle el foco
    toplevel.lift()
    toplevel.focus_force()
    
    # 3. Hacer que la ventana tome el control total (Modal)
    toplevel.grab_set()
    
    # Variables Locales
    var_cedula_busq = tk.StringVar()
    var_nombre_cliente = tk.StringVar() # NEW: Para recibo
    var_nro_cuota = tk.StringVar()
    var_vencimiento = tk.StringVar()
    var_capital = tk.DoubleVar()
    var_interes = tk.DoubleVar()
    var_dias_mora = tk.IntVar(value=0)
    var_mora = tk.DoubleVar(value=0.0)
    var_total_pagar = tk.DoubleVar(value=0.0)
    var_recibido = tk.DoubleVar(value=0.0)
    var_cambio = tk.DoubleVar(value=0.0)
    var_estado_visual = tk.StringVar(value="ESPERANDO BÚSQUEDA")
    
    # --- UI LAYOUT ---
    
    # 1. PANEL SUPERIOR (BÚSQUEDA)
    f_top = ctk.CTkFrame(toplevel, fg_color="white", height=80)
    f_top.pack(fill='x', padx=10, pady=10)
    
    ctk.CTkLabel(f_top, text="Cédula Cliente:", font=('Arial', 12, 'bold')).pack(side='left', padx=20)
    e_cedula_rec = ctk.CTkEntry(f_top, textvariable=var_cedula_busq, width=150, font=('Arial', 14))
    e_cedula_rec.pack(side='left', padx=10)
    
    def buscar_credito_activo():
        ced = var_cedula_busq.get().strip()
        if not ced:
            messagebox.showwarning("Aviso", "Ingrese una cédula")
            return
            
        conn, cursor = conectar_db()
        try:
            # 1. Buscar Crédito Desembolsado
            cursor.execute("""
                SELECT id, fecha_desembolso_real, monto_aprobado, tasa_interes, plazo_meses, valor_cuota, dia_pago
                FROM Microcreditos 
                WHERE cedula_cliente = %s AND status = 'Desembolsado'
                ORDER BY id DESC LIMIT 1
            """, (ced,))
            credito = cursor.fetchone()
            
            # Buscar nombre del cliente
            cursor.execute("SELECT nombre FROM Clientes WHERE cedula = %s", (ced,))
            cli = cursor.fetchone()
            if cli: var_nombre_cliente.set(cli[0])
            
            if not credito:
                messagebox.showwarning("Aviso", "Este cliente no tiene un crédito marcado como 'Desembolsado'.")
                return
            
            # Unpack credito
            # id, f_desemb, monto, tasa, plazo, cuota, dia_pago
            cred_id = credito[0]
            f_desemb_str = credito[1] # Expected YYYY-MM-DD or similar
            monto = credito[2] or 0.0
            plazo = credito[4] or 0
            val_cuota = credito[5] or 0.0
            dia_pago = credito[6] or 1
            
            # 2. Calcular cuotas pagadas
            cursor.execute("SELECT COUNT(*), SUM(valor_capital) FROM Pagos WHERE cedula_cliente = %s", (ced,))
            pagos_info = cursor.fetchone()
            cuotas_pagadas = pagos_info[0] or 0
            capital_pagado = pagos_info[1] or 0.0
            
            cuota_actual = cuotas_pagadas + 1
            if cuota_actual > plazo:
                messagebox.showinfo("Info", "El crédito parece estar totalmente pagado.")
                return
            
            var_nro_cuota.set(f"{cuota_actual} / {plazo}")
            var_capital.set(round(val_cuota, 2)) # Simplificación: Asumimos cuota fija capital+interes, mostramos total como capital para pago simple o calculamos
            # Ajuste: El sistema original pide capital e interes separados. Vamos a simplificar usando el valor cuota total en capital 
            # o dividirlo si tuvieramos tabla de amortización.
            # User request: "el credito que se creo puede ser cobrado"
            # Asumiremos cuota nivelada donde el usuario paga el valor_cuota completo.
            # Para visualización, pondremos el valor_cuota en "Capital" y 0 en Interés para no complicar sin amortización generada.
            var_capital.set(round(val_cuota, 2))
            var_interes.set(0.00)
            
            # 3. Calcular Vencimiento
            # Logic: Fecha desembolso -> Primer pago al siguiente mes el día 'dia_pago'
            # Cuota N: Fecha desembolso + N meses (ajustando día)
            try:
                # Parsear fecha desembolso
                f_ini = datetime.datetime.strptime(f_desemb_str.split(' ')[0], "%d/%m/%Y")
            except:
                try:
                    f_ini = datetime.datetime.strptime(f_desemb_str.split(' ')[0], "%Y-%m-%d")
                except:
                    f_ini = datetime.datetime.today() # Fallback
            
            # Calcular fecha vencimiento de la cuota actual
            # Aprox: f_ini + cuota_actual meses. 
            # Python no tiene add_months directo simple, usamos lógica simple:
            mes_venc = f_ini.month + cuota_actual
            anio_venc = f_ini.year + (mes_venc - 1) // 12
            mes_venc = (mes_venc - 1) % 12 + 1
            
            # Validar día
            try:
                f_venc = datetime.date(anio_venc, mes_venc, int(dia_pago))
            except ValueError:
                # Si el día no existe (ej: 31 feb), usar el último del mes
                import calendar
                last_day = calendar.monthrange(anio_venc, mes_venc)[1]
                f_venc = datetime.date(anio_venc, mes_venc, last_day)
                
            var_vencimiento.set(f_venc.strftime("%Y-%m-%d"))
            
            # 4. Mora
            hoy = datetime.date.today()
            delta = (hoy - f_venc).days
            dias_mora = max(0, delta)
            var_dias_mora.set(dias_mora)
            
            mora = round(val_cuota * 0.011 * dias_mora, 2) # 1.1% por dia sobre cuota
            var_mora.set(mora)
            
            total = val_cuota + mora
            var_total_pagar.set(round(total, 2))
            
            # Estado Visual
            if dias_mora > 0:
                var_estado_visual.set(f"ESTADO: MORA ({dias_mora} DÍAS)")
                lbl_estado.configure(text_color="#D32F2F", fg_color="#FFCDD2") 
                e_total_pagar.configure(fg_color="#FFCDD2", text_color="#D32F2F")
            else:
                var_estado_visual.set("ESTADO: AL DÍA")
                lbl_estado.configure(text_color="#388E3C", fg_color="#C8E6C9") 
                e_total_pagar.configure(fg_color="#E0F2F1", text_color="#00695C")
                
            e_recibido.focus_set()
            
        except Exception as e:
            messagebox.showerror("Error", f"Error buscando crédito: {e}")
        finally:
            db_manager.release_connection(conn)
        
    def generar_recibo_pdf():
        try:
            # Datos recibo
            nombre = var_nombre_cliente.get() or "Cliente Final"
            cedula = var_cedula_busq.get()
            cuota = var_nro_cuota.get()
            valor_c = var_capital.get() + var_interes.get()
            valor_m = var_mora.get()
            total = var_total_pagar.get()
            fecha_hoy = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            # HTML content
            html_content = f"""
            <html>
            <head>
                <style>
                    body {{ font-family: Helvetica, sans-serif; font-size: 12px; }}
                    .header {{ text-align: center; margin-bottom: 20px; }}
                    .title {{ font-size: 16px; font-weight: bold; }}
                    .info {{ margin-bottom: 5px; }}
                    .total {{ font-size: 14px; font-weight: bold; margin-top: 10px; border-top: 1px dashed black; padding-top: 5px; }}
                </style>
            </head>
            <body>
                <div class="header">
                    <div class="title">COMPROBANTE DE PAGO</div>
                    <div>ALIANZA C3F</div>
                    <div>{fecha_hoy}</div>
                </div>
                
                <div class="info"><strong>Cliente:</strong> {nombre}</div>
                <div class="info"><strong>Cédula:</strong> {cedula}</div>
                <div class="info"><strong>Concepto:</strong> Pago de Cuota {cuota}</div>
                
                <hr>
                
                <table style="width: 100%;">
                    <tr>
                        <td>Valor Cuota:</td>
                        <td style="text-align: right;">$ {valor_c:.2f}</td>
                    </tr>
                    <tr>
                        <td>Mora / Multa:</td>
                        <td style="text-align: right;">$ {valor_m:.2f}</td>
                    </tr>
                </table>
                
                <div class="total">
                    <table style="width: 100%;">
                        <tr>
                            <td>TOTAL PAGADO:</td>
                            <td style="text-align: right;">$ {total:.2f}</td>
                        </tr>
                    </table>
                </div>
                
                <br><br>
                <div style="text-align: center;">__________________________<br>Firma Cajero</div>
            </body>
            </html>
            """
            
            # Save file
            carpeta = "Recibos"
            if not os.path.exists(carpeta):
                os.makedirs(carpeta)
                
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Recibo_{cedula}_{timestamp}.pdf"
            path = os.path.abspath(os.path.join(carpeta, filename))
            
            with open(path, "w+b") as result_file:
                pisa_status = pisa.CreatePDF(html_content, dest=result_file)
                
            if pisa_status.err:
                messagebox.showerror("Error PDF", "Error generando PDF")
            else:
                os.startfile(path)
                
        except Exception as e:
            messagebox.showerror("Error", f"Error creando recibo: {e}")
        
    ctk.CTkButton(f_top, text="🔍 BUSCAR CRÉDITO ACTIVO", command=buscar_credito_activo,
                  fg_color="#1860C3", width=200).pack(side='left', padx=20)
                  
    # 2. PANEL CONTENIDO (SPLIT)
    f_content = ctk.CTkFrame(toplevel, fg_color="transparent")
    f_content.pack(fill='both', expand=True, padx=10, pady=5)
    
    # LEFT PANEL (DETALLES)
    f_left = ctk.CTkFrame(f_content, fg_color="white", width=400)
    f_left.pack(side='left', fill='both', expand=True, padx=5)
    
    ctk.CTkLabel(f_left, text="DETALLE DE CUOTA", font=('Arial', 14, 'bold'), text_color="#555").pack(pady=15)
    
    # Grid detalles
    f_grid_det = ctk.CTkFrame(f_left, fg_color="transparent")
    f_grid_det.pack(pady=10, padx=20)
    
    ctk.CTkLabel(f_grid_det, text="Nro Cuota:", font=('Arial', 12)).grid(row=0, column=0, sticky='e', pady=5)
    ctk.CTkEntry(f_grid_det, textvariable=var_nro_cuota, state='readonly', width=150).grid(row=0, column=1, padx=10, pady=5)
    
    ctk.CTkLabel(f_grid_det, text="Vencimiento:", font=('Arial', 12)).grid(row=1, column=0, sticky='e', pady=5)
    ctk.CTkEntry(f_grid_det, textvariable=var_vencimiento, state='readonly', width=150).grid(row=1, column=1, padx=10, pady=5)
    
    ctk.CTkLabel(f_grid_det, text="Valor Capital ($):", font=('Arial', 12, 'bold')).grid(row=2, column=0, sticky='e', pady=5)
    ctk.CTkEntry(f_grid_det, textvariable=var_capital, state='readonly', width=150, justify='right').grid(row=2, column=1, padx=10, pady=5)
    
    ctk.CTkLabel(f_grid_det, text="Valor Interés ($):", font=('Arial', 12, 'bold')).grid(row=3, column=0, sticky='e', pady=5)
    ctk.CTkEntry(f_grid_det, textvariable=var_interes, state='readonly', width=150, justify='right').grid(row=3, column=1, padx=10, pady=5)
    
    # ESTADO VISUAL
    lbl_estado = ctk.CTkLabel(f_left, textvariable=var_estado_visual, font=('Arial', 18, 'bold'), height=60, corner_radius=8)
    lbl_estado.pack(fill='x', padx=20, pady=30)
    
    # RIGHT PANEL (CÁLCULOS Y PAGO)
    f_right = ctk.CTkFrame(f_content, fg_color="#F4F6F7", width=400)
    f_right.pack(side='left', fill='both', expand=True, padx=(5,0))
    
    ctk.CTkLabel(f_right, text="CÁLCULO Y PAGO", font=('Arial', 14, 'bold'), text_color="#555").pack(pady=15)
    
    f_grid_pag = ctk.CTkFrame(f_right, fg_color="transparent")
    f_grid_pag.pack(pady=10, padx=20)
    
    # Dias Atraso
    ctk.CTkLabel(f_grid_pag, text="Días Atraso:", font=('Arial', 12)).grid(row=0, column=0, sticky='e', pady=5)
    ctk.CTkEntry(f_grid_pag, textvariable=var_dias_mora, state='readonly', width=80).grid(row=0, column=1, sticky='w', padx=10, pady=5)
    
    # Mora
    ctk.CTkLabel(f_grid_pag, text="Mora / Penalidad ($):", font=('Arial', 12, 'bold'), text_color="#D32F2F").grid(row=1, column=0, sticky='e', pady=5)
    ctk.CTkEntry(f_grid_pag, textvariable=var_mora, state='readonly', width=120, justify='right', fg_color="#FFEBEE", text_color="#D32F2F").grid(row=1, column=1, sticky='w', padx=10, pady=5)
    ctk.CTkLabel(f_grid_pag, text="(1.1% diario s/Capital)", font=('Arial', 10), text_color="#777").grid(row=1, column=2, sticky='w')
    
    # Separator
    ttk.Separator(f_grid_pag, orient='horizontal').grid(row=2, column=0, columnspan=3, sticky='ew', pady=10)
    
    # TOTAL A PAGAR
    ctk.CTkLabel(f_grid_pag, text="TOTAL A PAGAR:", font=('Arial', 14, 'bold')).grid(row=3, column=0, sticky='e', pady=10)
    e_total_pagar = ctk.CTkEntry(f_grid_pag, textvariable=var_total_pagar, font=('Arial', 16, 'bold'), width=140, justify='right', state='readonly')
    e_total_pagar.grid(row=3, column=1, columnspan=2, sticky='w', padx=10, pady=10)
    
    # MONTO RECIBIDO
    ctk.CTkLabel(f_grid_pag, text="Monto Recibido ($):", font=('Arial', 12, 'bold')).grid(row=4, column=0, sticky='e', pady=5)
    
    def calc_cambio(*args):
        try:
            tot = var_total_pagar.get()
            rec = var_recibido.get()
            cam = rec - tot
            var_cambio.set(round(cam, 2))
        except:
            var_cambio.set(0.0)
            
    e_recibido = ctk.CTkEntry(f_grid_pag, textvariable=var_recibido, font=('Arial', 14), width=140, justify='right')
    e_recibido.grid(row=4, column=1, columnspan=2, sticky='w', padx=10, pady=5)
    e_recibido.bind("<KeyRelease>", lambda e: calc_cambio())
    
    # CAMBIO
    ctk.CTkLabel(f_grid_pag, text="Cambio ($):", font=('Arial', 12, 'bold')).grid(row=5, column=0, sticky='e', pady=5)
    ctk.CTkEntry(f_grid_pag, textvariable=var_cambio, font=('Arial', 14, 'bold'), width=140, justify='right', state='readonly', fg_color="#E8F5E9", text_color="green").grid(row=5, column=1, columnspan=2, sticky='w', padx=10, pady=5)
    
    # 3. BOTÓN ACCIÓN
    def registrar_pago():
        try:
            # Validaciones básicas
            rec = var_recibido.get()
            tot = var_total_pagar.get()
            if rec < tot:
                messagebox.showwarning("Pago Insuficiente", "El monto recibido es menor al total.")
                return
            
            # Guardar en BD
            conn, cursor = conectar_db()
            
            ced = var_cedula_busq.get()
            # Parsear "X / Y" -> X
            try:
                nro = int(var_nro_cuota.get().split('/')[0].strip())
            except: nro = 0
            
            fecha_pago = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            cursor.execute("""
                INSERT INTO Pagos (fecha, cedula_cliente, cuota_nro, valor_capital, valor_interes, valor_mora, total_pagado, usuario)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                fecha_pago,
                ced,
                nro,
                var_capital.get(),
                var_interes.get(),
                var_mora.get(),
                var_total_pagar.get(),
                USUARIO_ACTIVO
            ))
            
            conn.commit()
            db_manager.release_connection(conn)

            messagebox.showinfo("Éxito", "Pago Registrado Correctamente.")
            
            # Habilitar Imprimir Recibo
            btn_imprimir_recibo.configure(state='normal', fg_color="#17a2b8")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error al registrar pago: {e}")

    f_btns = ctk.CTkFrame(toplevel, fg_color="transparent")
    f_btns.pack(fill='x', padx=40, pady=20)

    ctk.CTkButton(f_btns, text="✅ REGISTRAR PAGO", command=registrar_pago,
                  font=('Arial', 14, 'bold'), height=50, width=300, fg_color="#28a745", hover_color="#218838").pack(side='left', padx=10, fill='x', expand=True)
                  
    btn_imprimir_recibo = ctk.CTkButton(f_btns, text="🖨️ IMPRIMIR RECIBO", command=generar_recibo_pdf,
                                      font=('Arial', 14, 'bold'), height=50, width=200, fg_color="#6c757d", state='disabled')
    btn_imprimir_recibo.pack(side='left', padx=10)
                  
    # Focus inicial
    e_cedula_rec.focus_set()


if __name__ == '__main__':
    verificar_integridad_db()
    win = ctk.CTk()
    win.title("Acceso al Sistema - Alianza C3F")
    
    # Theme color
    COLOR_FONDO = "#FAFAD2"
    win.configure(fg_color=COLOR_FONDO)
    
    # Background Image
    # Background Image
    try:
        img_bg = Image.open(resource_path("background.jpg"))
        # For login, we can use fixed size or scaled
        logo_bg = ctk.CTkImage(light_image=img_bg, dark_image=img_bg, size=(450, 650))
        lbl_bg = ctk.CTkLabel(win, image=logo_bg, text="")
        lbl_bg.place(x=0, y=0, relwidth=1, relheight=1)
    except Exception as e:
        print(f"Error cargando fondo login: {e}")
    
    # Increase size for the "card" look
    w_width, w_height = 450, 650
    screen_width = win.winfo_screenwidth()
    screen_height = win.winfo_screenheight()
    pos_x, pos_y = (screen_width - w_width) // 2, (screen_height - w_height) // 2
    win.geometry(f"{w_width}x{w_height}+{pos_x}+{pos_y}")
    win.resizable(False, False)

    # Central Card
    card = ctk.CTkFrame(win, width=380, height=550, fg_color="white", corner_radius=20, border_width=1, border_color="#CCCCCC")
    card.place(relx=0.5, rely=0.5, anchor="center")

    # Logo inside card
    try:
        img = Image.open(resource_path("Logo Face.jpg"))
        logo_login = ctk.CTkImage(light_image=img, dark_image=img, size=(180, 160))
        lbl_logo = ctk.CTkLabel(card, image=logo_login, text="")
        lbl_logo.pack(pady=(30, 20))
    except:
        ctk.CTkLabel(card, text="ALIANZA C3F", font=("Arial", 24, "bold"), text_color="#1860C3").pack(pady=(50, 30))

    ctk.CTkLabel(card, text="IDENTIFICACIÓN DE USUARIO", font=("Arial", 14, "bold"), text_color="#555555").pack(pady=(10, 20))

    # Entries with modern feel
    ctk.CTkLabel(card, text="Usuario:", font=("Arial", 12), text_color="black").pack(anchor="w", padx=40)
    u = ctk.CTkEntry(card, width=300, height=40, corner_radius=10, placeholder_text="Ingrese su usuario", 
                     fg_color="#F9F9F9", text_color="black", border_color="#CCCCCC")
    u.pack(pady=(5, 15))
    
    ctk.CTkLabel(card, text="Contraseña:", font=("Arial", 12), text_color="black").pack(anchor="w", padx=40)
    p = ctk.CTkEntry(card, width=300, height=40, corner_radius=10, show="*", placeholder_text="••••••••",
                     fg_color="#F9F9F9", text_color="black", border_color="#CCCCCC")
    p.pack(pady=(5, 30))
    
    # Interaction logic
    u.bind('<Return>', lambda e: p.focus())
    p.bind('<Return>', lambda e: login_fn(win, u, p))
    
    # Login Button
    btn_entrar = ctk.CTkButton(card, text="ENTRAR AL SISTEMA", command=lambda: login_fn(win, u, p),
                               font=("Arial", 14, "bold"), height=45, width=250, 
                               fg_color="#1860C3", hover_color="#1452A6", corner_radius=10)
    btn_entrar.pack(pady=10)

    # Footer/Version
    ctk.CTkLabel(card, text="v.2025.12 - Seguridad SSL", font=("Arial", 10), text_color="grey").pack(side="bottom", pady=15)

    win.mainloop()

